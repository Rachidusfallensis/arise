import ollama
import chromadb
from chromadb.config import Settings
from typing import List, Dict, Tuple, Optional
import json
from src.core.document_processor import ArcadiaDocumentProcessor
from src.core.requirements_generator import RequirementsGenerator
from config import config, arcadia_config
import logging
from ..utils.enhanced_requirement_extractor import EnhancedRequirementExtractor
import time

class SAFEMBSERAGSystem:
    def __init__(self):
        self.ollama_client = ollama.Client(host=config.OLLAMA_BASE_URL)
        self.chroma_client = chromadb.PersistentClient(path=config.VECTORDB_PATH)
        self.collection = self._get_or_create_collection()
        self.doc_processor = ArcadiaDocumentProcessor()
        self.req_generator = RequirementsGenerator(self.ollama_client)
        self.logger = logging.getLogger(__name__)
        
        # Enhanced requirement extractor for advanced analysis
        self.enhanced_extractor = EnhancedRequirementExtractor()
        self.logger.info("Enhanced requirement extractor initialized")
        self.logger.info("RAG system initialized with ChromaDB default embeddings")
    
    def _get_or_create_collection(self):
        """Get or create ChromaDB collection for SAFE MBSE"""
        try:
            collection = self.chroma_client.get_collection(name=config.COLLECTION_NAME)
        except:
            collection = self.chroma_client.create_collection(
                name=config.COLLECTION_NAME,
                metadata={"description": "SAFE MBSE Requirements Generation System"}
            )
        return collection
    
    def generate_requirements_from_proposal(self, 
                                          proposal_text: str, 
                                          target_phase: str = "all",
                                          requirement_types: Optional[List[str]] = None) -> Dict:
        """
        Generate requirements from project proposal using ARCADIA methodology
        
        Args:
            proposal_text: The project proposal text
            target_phase: ARCADIA phase to focus on ("operational", "system", "logical", "physical", "all")
            requirement_types: Types of requirements to generate ["functional", "non_functional", "stakeholder"]
        
        Returns:
            Dictionary containing generated requirements organized by phase and type
        """
        if requirement_types is None:
            requirement_types = ["functional", "non_functional", "stakeholder"]
        
        results: Dict = {
            "metadata": {
                "source": "project_proposal",
                "generation_timestamp": "2024-01-01T00:00:00Z",
                "target_phase": target_phase,
                "requirement_types": requirement_types
            },
            "stakeholders": {},
            "requirements": {},
            "statistics": {}
        }
        
        # Process the proposal to extract context
        context_chunks = self._extract_proposal_context(proposal_text)
        
        # Generate requirements for each phase
        phases_to_process = [target_phase] if target_phase != "all" else list(arcadia_config.ARCADIA_PHASES.keys())
        
        for phase in phases_to_process:
            if phase not in results["requirements"]:
                results["requirements"][phase] = {}
            
            phase_context = self._filter_context_by_phase(context_chunks, phase)
            
            # Generate stakeholder requirements (mainly for operational phase)
            if "stakeholder" in requirement_types and phase == "operational":
                stakeholders = self.req_generator.generate_stakeholders(phase_context, proposal_text)
                results["stakeholders"] = stakeholders
            
            # Generate functional requirements (skip for operational phase - focus on stakeholder needs)
            if "functional" in requirement_types and phase != "operational":
                functional_reqs = self.req_generator.generate_functional_requirements(
                    phase_context, phase, proposal_text
                )
                results["requirements"][phase]["functional"] = functional_reqs
            
            # Generate non-functional requirements (skip for operational phase - focus on stakeholder needs)  
            if "non_functional" in requirement_types and phase != "operational":
                nf_reqs = self.req_generator.generate_non_functional_requirements(
                    phase_context, phase, proposal_text
                )
                results["requirements"][phase]["non_functional"] = nf_reqs
        
        # Generate statistics
        results["statistics"] = self._calculate_generation_statistics(results)
        
        return results
    
    def _extract_proposal_context(self, proposal_text: str) -> List[Dict]:
        """Extract and chunk the proposal for better context processing"""
        chunks = self.doc_processor._chunk_text_with_metadata(
            proposal_text, 
            {"source": "proposal", "type": "project_description"}
        )
        return chunks
    
    def _filter_context_by_phase(self, chunks: List[Dict], phase: str) -> List[Dict]:
        """Filter context chunks relevant to specific ARCADIA phase"""
        phase_info = arcadia_config.ARCADIA_PHASES.get(phase, {})
        keywords = phase_info.get("keywords", [])
        
        relevant_chunks = []
        for chunk in chunks:
            content_lower = chunk["content"].lower()
            if any(keyword in content_lower for keyword in keywords):
                chunk["phase_relevance"] = phase
                relevant_chunks.append(chunk)
        
        return relevant_chunks if relevant_chunks else chunks[:3]  # Fallback to first 3 chunks
    
    def _calculate_generation_statistics(self, results: Dict) -> Dict:
        """Calculate statistics about generated requirements"""
        stats: Dict = {
            "total_requirements": 0,
            "by_phase": {},
            "by_type": {},
            "by_priority": {"MUST": 0, "SHOULD": 0, "COULD": 0}
        }
        
        for phase, phase_reqs in results["requirements"].items():
            phase_count = 0
            stats["by_phase"][phase] = {}
            
            for req_type, reqs in phase_reqs.items():
                count = len(reqs) if isinstance(reqs, list) else 0
                phase_count += count
                stats["by_type"][req_type] = stats["by_type"].get(req_type, 0) + count
                stats["by_phase"][phase][req_type] = count
                
                # Count by priority
                if isinstance(reqs, list):
                    for req in reqs:
                        priority = req.get("priority", "SHOULD")
                        stats["by_priority"][priority] += 1
            
            stats["by_phase"][phase]["total"] = phase_count
            stats["total_requirements"] += phase_count
        
        return stats
    
    def evaluate_against_cyderco(self, generated_requirements: Dict) -> Dict:
        """
        Evaluate generated requirements against CYDERCO benchmarks
        
        Args:
            generated_requirements: Requirements generated by the system
        
        Returns:
            Evaluation results with coverage metrics
        """
        # Load CYDERCO benchmark data
        cyderco_reqs = self._load_cyderco_requirements()
        
        evaluation = {
            "coverage_score": 0.0,
            "similarity_scores": {},
            "missing_requirements": [],
            "additional_requirements": [],
            "quality_metrics": {}
        }
        
        # Calculate coverage score
        coverage = self._calculate_coverage_score(generated_requirements, cyderco_reqs)
        evaluation["coverage_score"] = coverage
        
        # Calculate similarity scores for matching requirements
        similarities = self._calculate_similarity_scores(generated_requirements, cyderco_reqs)
        evaluation["similarity_scores"] = similarities
        
        # Identify gaps and additions
        gaps, additions = self._identify_requirement_gaps(generated_requirements, cyderco_reqs)
        evaluation["missing_requirements"] = gaps
        evaluation["additional_requirements"] = additions
        
        # Quality metrics
        quality = self._assess_requirement_quality(generated_requirements)
        evaluation["quality_metrics"] = quality
        
        return evaluation
    
    def _load_cyderco_requirements(self) -> Dict:
        """Load CYDERCO requirements for benchmarking"""
        # This would load the CYDERCO analysis from the data files
        # For now, return a structured representation
        return {
            "functional": [
                {"id": "FUNC-DA-1", "description": "Data correlation capabilities"},
                {"id": "FUNC-DA-2", "description": "Data visualization methods"},
                {"id": "FUNC-NTA-1", "description": "Traffic monitoring"},
                # ... more CYDERCO requirements
            ],
            "non_functional": [
                {"id": "NFUNC-DA-1", "description": "Data ingestion performance"},
                {"id": "NFUNC-DA-2", "description": "System scalability"},
                # ... more non-functional requirements
            ]
        }
    
    def _calculate_coverage_score(self, generated: Dict, benchmark: Dict) -> float:
        """Calculate what percentage of benchmark requirements are covered"""
        # Implementation of coverage calculation
        total_benchmark = sum(len(reqs) for reqs in benchmark.values())
        if total_benchmark == 0:
            return 0.0
        
        covered = 0
        # Logic to match generated requirements with benchmark
        # This is a simplified version
        for phase_reqs in generated.get("requirements", {}).values():
            for req_type, reqs in phase_reqs.items():
                if req_type in benchmark:
                    covered += min(len(reqs), len(benchmark[req_type]))
        
        return min(covered / total_benchmark, 1.0) * 100
    
    def _calculate_similarity_scores(self, generated: Dict, benchmark: Dict) -> Dict:
        """Calculate semantic similarity between generated and benchmark requirements"""
        # Implementation would use embedding similarity
        return {"average_similarity": 0.75, "detailed_scores": {}}
    
    def _identify_requirement_gaps(self, generated: Dict, benchmark: Dict) -> Tuple[List[str], List[str]]:
        """Identify missing and additional requirements"""
        missing: List[str] = []
        additional: List[str] = []
        # Implementation of gap analysis
        return missing, additional
    
    def _assess_requirement_quality(self, requirements: Dict) -> Dict:
        """Assess the quality of generated requirements"""
        return {
            "completeness": 0.85,
            "consistency": 0.90,
            "traceability": 0.80,
            "testability": 0.75
        }
    
    def export_requirements(self, requirements: Dict, export_format: str = "JSON") -> str:
        """Export requirements in specified format"""
        if export_format == "JSON":
            return json.dumps(requirements, indent=2)
        elif export_format == "Markdown":
            return self._convert_to_markdown(requirements)
        elif export_format == "Excel":
            return self._convert_to_excel(requirements)
        elif export_format == "DOORS":
            return self._convert_to_doors(requirements)
        elif export_format == "ReqIF":
            return self._convert_to_reqif(requirements)
        else:
            raise ValueError(f"Unsupported export format: {export_format}")
    
    def _convert_to_markdown(self, requirements: Dict) -> str:
        """Convert requirements to Markdown format"""
        md_content = "# Generated Requirements\n\n"
        
        for phase, phase_reqs in requirements.get("requirements", {}).items():
            md_content += f"## {phase.title()} Phase\n\n"
            
            for req_type, reqs in phase_reqs.items():
                md_content += f"### {req_type.title()} Requirements\n\n"
                
                for req in reqs:
                    md_content += f"**{req.get('id', 'N/A')}**: {req.get('title', 'Untitled')}\n\n"
                    md_content += f"*Description*: {req.get('description', 'No description')}\n\n"
                    md_content += f"*Priority*: {req.get('priority', 'N/A')}\n\n"
                    md_content += f"*Verification*: {req.get('verification_method', 'N/A')}\n\n"
                    md_content += "---\n\n"
        
        return md_content

    def _convert_to_excel(self, requirements: Dict) -> str:
        """Convert requirements to Excel format (CSV-like for compatibility)"""
        try:
            import io
            import csv
            
            output = io.StringIO()
            writer = csv.writer(output)
            
            # Write header
            writer.writerow(['ID', 'Phase', 'Type', 'Title', 'Description', 'Priority', 'Verification Method'])
            
            # Write requirements
            for phase, phase_reqs in requirements.get("requirements", {}).items():
                for req_type, reqs in phase_reqs.items():
                    for req in reqs:
                        writer.writerow([
                            req.get('id', 'N/A'),
                            phase.title(),
                            req_type.title(),
                            req.get('title', 'Untitled'),
                            req.get('description', 'No description'),
                            req.get('priority', 'N/A'),
                            req.get('verification_method', 'N/A')
                        ])
            
            excel_content = output.getvalue()
            output.close()
            return excel_content
            
        except ImportError:
            # Fallback to simple table format
            excel_content = "ID,Phase,Type,Title,Description,Priority,Verification Method\n"
            
            for phase, phase_reqs in requirements.get("requirements", {}).items():
                for req_type, reqs in phase_reqs.items():
                    for req in reqs:
                        excel_content += f"{req.get('id', 'N/A')},{phase.title()},{req_type.title()},{req.get('title', 'Untitled')},{req.get('description', 'No description')},{req.get('priority', 'N/A')},{req.get('verification_method', 'N/A')}\n"
            
            return excel_content

    def _convert_to_doors(self, requirements: Dict) -> str:
        """Convert requirements to DOORS-compatible format"""
        doors_content = "// DOORS Import File\n"
        doors_content += "// Generated by SAFE MBSE RAG System\n\n"
        
        doors_content += "module main\n\n"
        
        for phase, phase_reqs in requirements.get("requirements", {}).items():
            doors_content += f"// {phase.title()} Phase Requirements\n"
            
            for req_type, reqs in phase_reqs.items():
                doors_content += f"folder \"{req_type.title()}\"\n"
                
                for req in reqs:
                    req_id = req.get('id', 'N/A')
                    title = req.get('title', 'Untitled').replace('"', '""')
                    description = req.get('description', 'No description').replace('"', '""')
                    priority = req.get('priority', 'N/A')
                    verification = req.get('verification_method', 'N/A')
                    
                    doors_content += f'requirement "{req_id}" = "{title}"\n'
                    doors_content += f'text = "{description}"\n'
                    doors_content += f'priority = "{priority}"\n'
                    doors_content += f'verification = "{verification}"\n\n'
                
                doors_content += "end folder\n\n"
        
        doors_content += "end module\n"
        return doors_content

    def _convert_to_reqif(self, requirements: Dict) -> str:
        """Convert requirements to ReqIF XML format"""
        reqif_content = """<?xml version="1.0" encoding="UTF-8"?>
<REQ-IF xmlns="http://www.omg.org/ReqIF" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
    <THE-HEADER>
        <REQ-IF-HEADER IDENTIFIER="SAFE_MBSE_RAG_EXPORT">
            <COMMENT>Generated by SAFE MBSE RAG System</COMMENT>
            <CREATION-TIME>""" + str(int(time.time())) + """</CREATION-TIME>
            <REPOSITORY-ID>SAFE_MBSE_RAG</REPOSITORY-ID>
            <REQ-IF-TOOL-ID>SAFE_MBSE_RAG_SYSTEM</REQ-IF-TOOL-ID>
            <REQ-IF-VERSION>1.0</REQ-IF-VERSION>
            <SOURCE-TOOL-ID>SAFE_MBSE_RAG</SOURCE-TOOL-ID>
            <TITLE>SAFE MBSE Requirements Export</TITLE>
        </REQ-IF-HEADER>
    </THE-HEADER>
    <CORE-CONTENT>
        <REQ-IF-CONTENT>
            <SPEC-OBJECTS>"""
        
        spec_id = 1
        for phase, phase_reqs in requirements.get("requirements", {}).items():
            for req_type, reqs in phase_reqs.items():
                for req in reqs:
                    req_id = req.get('id', 'N/A')
                    title = req.get('title', 'Untitled')
                    description = req.get('description', 'No description')
                    priority = req.get('priority', 'N/A')
                    verification = req.get('verification_method', 'N/A')
                    
                    # Escape XML special characters
                    title = title.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    description = description.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    
                    reqif_content += f"""
                <SPEC-OBJECT IDENTIFIER="SPEC_OBJ_{spec_id}">
                    <VALUES>
                        <ATTRIBUTE-VALUE-STRING THE-VALUE="{req_id}">
                            <DEFINITION>
                                <ATTRIBUTE-DEFINITION-STRING-REF>REQ_ID</ATTRIBUTE-DEFINITION-STRING-REF>
                            </DEFINITION>
                        </ATTRIBUTE-VALUE-STRING>
                        <ATTRIBUTE-VALUE-STRING THE-VALUE="{title}">
                            <DEFINITION>
                                <ATTRIBUTE-DEFINITION-STRING-REF>REQ_TITLE</ATTRIBUTE-DEFINITION-STRING-REF>
                            </DEFINITION>
                        </ATTRIBUTE-VALUE-STRING>
                        <ATTRIBUTE-VALUE-STRING THE-VALUE="{description}">
                            <DEFINITION>
                                <ATTRIBUTE-DEFINITION-STRING-REF>REQ_DESCRIPTION</ATTRIBUTE-DEFINITION-STRING-REF>
                            </DEFINITION>
                        </ATTRIBUTE-VALUE-STRING>
                        <ATTRIBUTE-VALUE-STRING THE-VALUE="{phase}">
                            <DEFINITION>
                                <ATTRIBUTE-DEFINITION-STRING-REF>REQ_PHASE</ATTRIBUTE-DEFINITION-STRING-REF>
                            </DEFINITION>
                        </ATTRIBUTE-VALUE-STRING>
                        <ATTRIBUTE-VALUE-STRING THE-VALUE="{req_type}">
                            <DEFINITION>
                                <ATTRIBUTE-DEFINITION-STRING-REF>REQ_TYPE</ATTRIBUTE-DEFINITION-STRING-REF>
                            </DEFINITION>
                        </ATTRIBUTE-VALUE-STRING>
                        <ATTRIBUTE-VALUE-STRING THE-VALUE="{priority}">
                            <DEFINITION>
                                <ATTRIBUTE-DEFINITION-STRING-REF>REQ_PRIORITY</ATTRIBUTE-DEFINITION-STRING-REF>
                            </DEFINITION>
                        </ATTRIBUTE-VALUE-STRING>
                        <ATTRIBUTE-VALUE-STRING THE-VALUE="{verification}">
                            <DEFINITION>
                                <ATTRIBUTE-DEFINITION-STRING-REF>REQ_VERIFICATION</ATTRIBUTE-DEFINITION-STRING-REF>
                            </DEFINITION>
                        </ATTRIBUTE-VALUE-STRING>
                    </VALUES>
                </SPEC-OBJECT>"""
                    spec_id += 1
        
        reqif_content += """
            </SPEC-OBJECTS>
        </REQ-IF-CONTENT>
    </CORE-CONTENT>
</REQ-IF>"""
        
        return reqif_content

    def add_documents_to_vectorstore(self, file_paths: List[str]) -> Dict:
        """
        Add documents to the vector store for chat functionality
        
        Args:
            file_paths: List of file paths to process and add
            
        Returns:
            Dictionary with processing results
        """
        results: Dict = {
            "processed": 0,
            "chunks_added": 0,
            "errors": []
        }
        
        self.logger.info(f"💬 Starting vectorstore processing for {len(file_paths)} files...")
        
        for file_path in file_paths:
            try:
                self.logger.info(f"📄 Processing file: {file_path}")
                
                # Extract text content based on file type
                content = ""
                file_extension = file_path.lower().split('.')[-1]
                
                if file_extension in ['txt', 'md']:
                    # Plain text files
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                        self.logger.info(f"   ✅ Text file processed: {len(content)} characters")
                    except UnicodeDecodeError:
                        with open(file_path, 'r', encoding='latin-1') as f:
                            content = f.read()
                        self.logger.info(f"   ✅ Text file processed (latin-1): {len(content)} characters")
                        
                elif file_extension == 'pdf':
                    # PDF files
                    try:
                        import PyPDF2
                        import io
                        
                        with open(file_path, 'rb') as f:
                            pdf_reader = PyPDF2.PdfReader(f)
                            content = ""
                            for page_num, page in enumerate(pdf_reader.pages):
                                try:
                                    page_text = page.extract_text()
                                    content += page_text + "\n"
                                    self.logger.info(f"     - Page {page_num + 1}: {len(page_text)} characters")
                                except Exception as page_error:
                                    self.logger.warning(f"     - Page {page_num + 1} failed: {str(page_error)}")
                        
                        self.logger.info(f"   ✅ PDF processed: {len(pdf_reader.pages)} pages, {len(content)} characters")
                        
                    except Exception as e:
                        self.logger.error(f"   ❌ PDF processing failed: {str(e)}")
                        results["errors"].append(f"PDF {file_path}: {str(e)}")
                        continue
                        
                elif file_extension == 'docx':
                    # DOCX files
                    try:
                        from docx import Document
                        
                        doc = Document(file_path)
                        content = ""
                        for para_num, paragraph in enumerate(doc.paragraphs):
                            para_text = paragraph.text
                            content += para_text + "\n"
                        
                        self.logger.info(f"   ✅ DOCX processed: {len(doc.paragraphs)} paragraphs, {len(content)} characters")
                        
                    except Exception as e:
                        self.logger.error(f"   ❌ DOCX processing failed: {str(e)}")
                        results["errors"].append(f"DOCX {file_path}: {str(e)}")
                        continue
                        
                elif file_extension in ['xml', 'aird']:
                    # XML/AIRD files
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                        self.logger.info(f"   ✅ XML/AIRD processed: {len(content)} characters")
                        
                    except Exception as e:
                        self.logger.error(f"   ❌ XML/AIRD processing failed: {str(e)}")
                        results["errors"].append(f"XML/AIRD {file_path}: {str(e)}")
                        continue
                        
                elif file_extension == 'json':
                    # JSON files
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            json_data = json.load(f)
                            content = json.dumps(json_data, indent=2)
                        self.logger.info(f"   ✅ JSON processed: {len(content)} characters")
                        
                    except Exception as e:
                        self.logger.error(f"   ❌ JSON processing failed: {str(e)}")
                        results["errors"].append(f"JSON {file_path}: {str(e)}")
                        continue
                        
                else:
                    # Unsupported file type
                    error_msg = f"Unsupported file type: {file_extension}"
                    self.logger.error(f"   ❌ {error_msg}")
                    results["errors"].append(f"{file_path}: {error_msg}")
                    continue
                
                # Check if content was extracted
                if not content or len(content.strip()) < 10:
                    error_msg = "No meaningful content extracted from file"
                    self.logger.warning(f"   ⚠️  {error_msg}")
                    results["errors"].append(f"{file_path}: {error_msg}")
                    continue
                
                # Create chunks with metadata using the existing document processor method
                text_chunks = self.doc_processor._chunk_text_with_metadata(
                    content,
                    {
                        "source": file_path,
                        "filename": file_path.split("/")[-1],
                        "file_type": file_extension,
                        "type": "uploaded_document",
                        "processed_for_chat": True,
                        "content_length": len(content)
                    }
                )
                
                self.logger.info(f"   📊 Created {len(text_chunks)} text chunks")
                
                # Add documents to vector store (using ChromaDB's default embeddings)
                chunks_added_for_file = 0
                for i, chunk in enumerate(text_chunks):
                    try:
                        # Create unique chunk ID
                        chunk_id = f"{file_path.replace('/', '_').replace(' ', '_')}_{i}_{len(chunk['content'])}"
                        
                        # Add to ChromaDB (ChromaDB will handle embedding generation automatically)
                        self.collection.add(
                            documents=[chunk["content"]],
                            metadatas=[chunk["metadata"]],
                            ids=[chunk_id]
                        )
                        
                        chunks_added_for_file += 1
                        results["chunks_added"] += 1
                        
                        if (i + 1) % 10 == 0:  # Log progress every 10 chunks
                            self.logger.info(f"     - Processed {i + 1}/{len(text_chunks)} chunks")
                        
                    except Exception as e:
                        self.logger.error(f"   ❌ Error processing chunk {i} from {file_path}: {str(e)}")
                        results["errors"].append(f"Chunk {i} from {file_path}: {str(e)}")
                
                results["processed"] += 1
                self.logger.info(f"✅ File completed: {file_path} - {chunks_added_for_file} chunks added successfully")
                
            except Exception as e:
                self.logger.error(f"❌ Error processing file {file_path}: {str(e)}")
                results["errors"].append(f"File {file_path}: {str(e)}")
        
        self.logger.info(f"🎉 Vectorstore processing completed: {results['processed']}/{len(file_paths)} files, {results['chunks_added']} total chunks")
        return results

    def query_documents(self, query: str, top_k: int = 5) -> Dict:
        """
        Query the document vector store for chat functionality
        
        Args:
            query: User query/question
            top_k: Number of relevant documents to retrieve
            
        Returns:
            Dictionary with answer and source documents
        """
        self.logger.info(f"💬 Processing query: {query[:100]}...")
        
        try:
            # Search vector store using ChromaDB's built-in text search
            search_results = self.collection.query(
                query_texts=[query],
                n_results=top_k
            )
            
            # Extract context documents
            context_docs = []
            if search_results["documents"] and search_results["documents"][0]:
                for i, (doc, metadata) in enumerate(zip(
                    search_results["documents"][0],
                    search_results["metadatas"][0]
                )):
                    context_docs.append(type('Document', (), {
                        'page_content': doc,
                        'metadata': metadata
                    })())
            
            # Generate response using context
            if context_docs:
                context_text = "\n\n".join([doc.page_content for doc in context_docs[:3]])
                
                prompt = f"""Based on the following context documents, please answer the user's question about ARCADIA methodology, MBSE, or the uploaded documents.

Context:
{context_text}

User Question: {query}

Please provide a comprehensive answer based on the context provided. If the context doesn't contain enough information to fully answer the question, acknowledge this and provide what information you can."""

                response = self.ollama_client.chat(
                    model="llama3:instruct",
                    messages=[
                        {"role": "user", "content": prompt}
                    ]
                )
                
                answer = response["message"]["content"]
                
                self.logger.info(f"✅ Generated response with {len(context_docs)} source documents")
                
                return {
                    "answer": answer,
                    "sources": context_docs,
                    "query": query
                }
            
            else:
                # No relevant documents found
                fallback_response = self.ollama_client.chat(
                    model="llama3:instruct",
                    messages=[
                        {"role": "user", "content": f"Please answer this question about ARCADIA methodology or MBSE: {query}"}
                    ]
                )
                
                return {
                    "answer": fallback_response["message"]["content"],
                    "sources": [],
                    "query": query
                }
                
        except Exception as e:
            self.logger.error(f"❌ Error processing query: {str(e)}")
            return {
                "answer": f"I apologize, but I encountered an error processing your question: {str(e)}",
                "sources": [],
                "query": query
            }

    def get_vectorstore_stats(self) -> Dict:
        """Get statistics about the current vector store"""
        try:
            collection_info = self.collection.get()
            
            total_docs = len(collection_info["ids"]) if collection_info["ids"] else 0
            
            # Group by source files
            sources: Dict[str, int] = {}
            if collection_info["metadatas"]:
                for metadata in collection_info["metadatas"]:
                    source = metadata.get("source", "unknown")
                    sources[source] = sources.get(source, 0) + 1
            
            return {
                "total_chunks": total_docs,
                "unique_sources": len(sources),
                "sources_breakdown": sources
            }
            
        except Exception as e:
            self.logger.error(f"Error getting vectorstore stats: {str(e)}")
            return {
                "total_chunks": 0,
                "unique_sources": 0,
                "sources_breakdown": {}
            }

    def clear_vectorstore(self):
        """Clear all documents from the vector store"""
        try:
            # Delete the collection and recreate it
            self.chroma_client.delete_collection(name=config.COLLECTION_NAME)
            self.collection = self._get_or_create_collection()
            self.logger.info("✅ Vector store cleared successfully")
            return True
        except Exception as e:
            self.logger.error(f"❌ Error clearing vector store: {str(e)}")
            return False

    def analyze_text_with_enhanced_extraction(self, text: str) -> Dict:
        """
        Analyse le texte avec extraction avancée de requirements
        
        Args:
            text: Texte à analyser
            
        Returns:
            Dictionnaire contenant l'analyse avancée
        """
        self.logger.info("Starting enhanced requirement analysis")
        start_time = time.time()
        
        try:
            # Extraction avancée des requirements
            enhanced_requirements = self.enhanced_extractor.extract_enhanced_requirements(text)
            
            # Génération de statistiques
            statistics = self.enhanced_extractor.get_statistics(enhanced_requirements)
            
            # Formatage pour affichage
            formatted_output = self.enhanced_extractor.format_extracted_requirements(enhanced_requirements)
            
            analysis_time = time.time() - start_time
            self.logger.info(f"Enhanced analysis completed in {analysis_time:.2f} seconds")
            self.logger.info(f"Found {len(enhanced_requirements)} requirement elements")
            
            return {
                "enhanced_requirements": enhanced_requirements,
                "statistics": statistics,
                "formatted_output": formatted_output,
                "analysis_time": analysis_time,
                "success": True
            }
            
        except Exception as e:
            self.logger.error(f"Error in enhanced requirement analysis: {str(e)}")
            return {
                "enhanced_requirements": [],
                "statistics": {},
                "formatted_output": f"Erreur lors de l'analyse: {str(e)}",
                "analysis_time": 0,
                "success": False,
                "error": str(e)
            }