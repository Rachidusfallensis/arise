import streamlit as st
import os
import json
import sys
from pathlib import Path
from datetime import datetime
import time

# Add the project root to Python path for proper imports
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from src.core.rag_system import SAFEMBSERAGSystem
from src.core.enhanced_structured_rag_system import EnhancedStructuredRAGSystem
from src.services.evaluation_service import EvaluationService
from config import config, arcadia_config
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import logging

# Configure logging for terminal output (Streamlit-compatible)
import os
from pathlib import Path

# Ensure logs directory exists
Path("logs").mkdir(exist_ok=True)

# Clear any existing handlers
logging.getLogger().handlers.clear()

# Configure logging with force and explicit stream
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stderr),  # Use stderr instead of stdout for Streamlit
        logging.FileHandler('logs/requirements_generation.log', mode='a')
    ],
    force=True  # Force reconfiguration
)

# Create logger with explicit configuration
logger = logging.getLogger("ARCADIA_RAG_System")
logger.setLevel(logging.INFO)

# Add console handler if not already present
if not any(isinstance(h, logging.StreamHandler) for h in logger.handlers):
    console_handler = logging.StreamHandler(sys.stderr)
    console_handler.setLevel(logging.INFO)
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    console_handler.setFormatter(formatter)
    logger.addHandler(console_handler)

# Test logging on startup
logger.info("Logger initialized successfully - terminal logging active")

# Page configuration
st.set_page_config(
    page_title="ARCADIA Requirements Generator",
    page_icon="🏗️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Modern CSS styling
st.markdown("""
<style>
    /* Main theme */
    .stApp {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
    }
    
    /* Header styling */
    .main-header {
        background: linear-gradient(90deg, #667eea 0%, #2c5aa0 100%);
        padding: 2rem;
        border-radius: 10px;
        margin-bottom: 2rem;
        color: white;
        text-align: center;
    }
    
    .main-header h1 {
        margin: 0;
        font-size: 2.5rem;
        font-weight: 300;
    }
    
    .main-header p {
        margin: 0.5rem 0 0 0;
        font-size: 1.1rem;
        opacity: 0.9;
    }
    
    /* Card styling */
    .info-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        border: 1px solid #e1e5e9;
        margin: 1rem 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #2c5aa0 100%);
        color: white;
        padding: 1.5rem;
        border-radius: 10px;
        text-align: center;
        margin: 0.5rem 0;
    }
    
    .metric-card h3 {
        margin: 0;
        font-size: 2rem;
        font-weight: 300;
    }
    
    .metric-card p {
        margin: 0.5rem 0 0 0;
        opacity: 0.9;
    }
    
    /* Requirements display */
    .requirement-item {
        background: white;
        padding: 1.5rem;
        border-radius: 8px;
        border-left: 4px solid #667eea;
        margin: 1rem 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    
    .requirement-header {
        font-weight: 600;
        font-size: 1.1rem;
        color: #333;
        margin-bottom: 0.5rem;
    }
    
    .requirement-description {
        color: #666;
        line-height: 1.6;
        margin-bottom: 1rem;
    }
    
    .requirement-meta {
        display: flex;
        gap: 1rem;
        font-size: 0.9rem;
    }
    
    .priority-high { color: #e74c3c; font-weight: 600; }
    .priority-medium { color: #f39c12; font-weight: 600; }
    .priority-low { color: #27ae60; font-weight: 600; }
    
    /* Chat interface */
    .chat-container {
        background: white;
        border-radius: 10px;
        padding: 1rem;
        margin: 1rem 0;
        border: 1px solid #e1e5e9;
    }
    
    .chat-message {
        margin: 1rem 0;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #667eea;
        background: #f8f9fa;
    }
    
    .user-message {
        border-left-color: #667eea;
        background: linear-gradient(135deg, #667eea 0%, #2c5aa0 100%);
        color: white;
        margin-left: 2rem;
    }
    
    .assistant-message {
        border-left-color: #27ae60;
        background: #f8f9fa;
        margin-right: 2rem;
    }
    
    .source-citation {
        background: #e8f4f8;
        border: 1px solid #d1ecf1;
        border-radius: 6px;
        padding: 1rem;
        margin: 0.5rem 0;
        font-size: 0.9rem;
    }
    
    /* Phase headers */
    .phase-section {
        background: linear-gradient(90deg, #667eea 0%, #2c5aa0 100%);
        color: white;
        padding: 1rem 1.5rem;
        border-radius: 8px;
        margin: 1.5rem 0 1rem 0;
    }
    
    .phase-section h3 {
        margin: 0;
        font-weight: 400;
    }
    
    /* Structured analysis styling */
    .analysis-overview {
        background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
        padding: 2rem;
        border-radius: 12px;
        margin: 1rem 0;
        border: 1px solid #dee2e6;
    }
    
    .actor-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        border-left: 4px solid #28a745;
        margin: 1rem 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    
    .capability-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        border-left: 4px solid #007bff;
        margin: 1rem 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    
    .function-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        border-left: 4px solid #ffc107;
        margin: 1rem 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    
    .traceability-link {
        background: #f8f9fa;
        border: 1px solid #dee2e6;
        border-radius: 8px;
        padding: 1rem;
        margin: 0.5rem 0;
    }
    
    .gap-analysis {
        background: #fff3cd;
        border: 1px solid #ffeaa7;
        border-radius: 8px;
        padding: 1rem;
        margin: 0.5rem 0;
    }
    
    .gap-critical {
        background: #f8d7da;
        border-color: #f5c6cb;
    }
    
    .gap-major {
        background: #fff3cd;
        border-color: #ffeaa7;
    }
    
    .gap-minor {
        background: #d1ecf1;
        border-color: #bee5eb;
    }
    
    /* Enhanced metrics */
    .enhanced-metric {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 1.5rem;
        border-radius: 12px;
        text-align: center;
        margin: 0.5rem 0;
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
    }
    
    .enhanced-metric h3 {
        margin: 0;
        font-size: 2.5rem;
        font-weight: 300;
    }
    
    .enhanced-metric p {
        margin: 0.5rem 0 0 0;
        opacity: 0.9;
        font-size: 1.1rem;
    }
    
    /* Component alignment improvements */
    .stColumn {
        padding: 0.5rem;
    }
    
    .stColumn > div {
        height: 100%;
        display: flex;
        flex-direction: column;
    }
    
    /* Enhanced ARCADIA analysis layout */
    .arcadia-overview {
        background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
        padding: 2rem;
        border-radius: 12px;
        margin: 1.5rem 0;
        border: 1px solid #dee2e6;
        box-shadow: 0 2px 8px rgba(0,0,0,0.05);
    }
    
    .arcadia-metrics-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
        gap: 1rem;
        margin: 1rem 0;
    }
    
    .arcadia-metric-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        border: 1px solid #e1e5e9;
        text-align: center;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }
    
    .arcadia-metric-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0,0,0,0.1);
    }
    
    /* Phase tabs styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 4px;
        margin: 1rem 0;
    }
    
    .stTabs [data-baseweb="tab"] {
        background: #f8f9fa;
        border-radius: 8px 8px 0 0;
        padding: 1rem 1.5rem;
        font-weight: 500;
        border: 1px solid #dee2e6;
        border-bottom: none;
        transition: all 0.2s ease;
    }
    
    .stTabs [data-baseweb="tab"]:hover {
        background: #e9ecef;
        transform: translateY(-1px);
    }
    
    .stTabs [data-baseweb="tab"][aria-selected="true"] {
        background: linear-gradient(90deg, #667eea 0%, #2c5aa0 100%);
        color: white;
        border-color: #667eea;
    }
    
    .stTabs [data-baseweb="tab-panel"] {
        border: 1px solid #dee2e6;
        border-radius: 0 0 8px 8px;
        padding: 1.5rem;
        background: white;
        min-height: 400px;
    }
    
    /* Buttons */
    .stButton > button {
        background: linear-gradient(90deg, #667eea 0%, #2c5aa0 100%);
        color: white;
        border: none;
        border-radius: 6px;
        padding: 0.5rem 1rem;
        font-weight: 500;
        transition: all 0.3s ease;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 8px rgba(0,0,0,0.2);
    }
    
    /* Sidebar */
    .css-1d391kg {
        background: #f8f9fa;
    }
    
    /* File uploader */
    .stFileUploader {
        border: 2px dashed #667eea;
        border-radius: 8px;
        padding: 1rem;
        background: #f8f9fa;
    }
    
    /* Progress indicators */
    .stProgress > div > div {
        background: linear-gradient(90deg, #667eea 0%, #2c5aa0 100%);
    }
    
    /* Tabs */
    .stTabs [data-baseweb="tab-list"] {
        gap: 2px;
    }
    
    .stTabs [data-baseweb="tab"] {
        background: #f8f9fa;
        border-radius: 8px 8px 0 0;
        padding: 0.75rem 1.5rem;
        font-weight: 500;
    }
    
    .stTabs [data-baseweb="tab"][aria-selected="true"] {
        background: linear-gradient(90deg, #667eea 0%, #2c5aa0 100%);
        color: white;
    }
    
    /* Professional improvements */
    .stAlert {
        border-radius: 8px;
        border: none;
        font-weight: 500;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }
    
    .stMetric {
        background: white;
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #e1e5e9;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
    }
    
    .stExpander {
        border: 1px solid #e1e5e9;
        border-radius: 8px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
    }
    
    .stSelectbox, .stTextInput, .stTextArea {
        border-radius: 6px;
    }
    
    .stColumn {
        padding: 0.25rem;
    }
    
    /* Professional typography */
    h1, h2, h3, h4 {
        font-weight: 600 !important;
        color: #2c3e50 !important;
        margin-bottom: 1rem !important;
    }
    
    /* Clean caption text */
    .stCaption {
        color: #6c757d !important;
        font-size: 0.875rem !important;
    }

    /* Hide streamlit elements */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .stDeployButton {display:none;}
</style>
""", unsafe_allow_html=True)

# ARCADIA methodology reference links
ARCADIA_REFERENCES = {
    "Getting Started": "https://gettingdesignright.com/GDR-Educate/Capella_Tutorial_v6_0/GettingStarted.html",
    "Installation": "https://gettingdesignright.com/GDR-Educate/Capella_Tutorial_v6_0/GettingStarted.html#Installation",
    "Operational Analysis": "https://gettingdesignright.com/GDR-Educate/Capella_Tutorial_v6_0/OperationalAnalysis.html",
    "System Analysis": "https://gettingdesignright.com/GDR-Educate/Capella_Tutorial_v6_0/SystemAnalysis.html",
    "Logical Architecture": "https://gettingdesignright.com/GDR-Educate/Capella_Tutorial_v6_0/LogicalArchitecture.html",
    "Physical Architecture": "https://gettingdesignright.com/GDR-Educate/Capella_Tutorial_v6_0/PhysicalArchitecture.html"
}

# MBSE context prompts for better responses
MBSE_CONTEXT_PROMPTS = {
    "operational": "Focus on operational activities, actors, and capabilities from an Operational Analysis perspective.",
    "system": "Consider system functions, interfaces, and requirements from a System Analysis viewpoint.",
    "logical": "Analyze logical components, their interactions, and behavioral aspects from a Logical Architecture perspective.",
    "physical": "Examine physical components, deployment, and implementation from a Physical Architecture standpoint.",
    "verification": "Address verification and validation approaches for MBSE artifacts.",
    "traceability": "Consider traceability links between different architectural levels."
}

# Initialize services
@st.cache_resource
def init_services(use_enhanced=True):
    """Initialize and cache the RAG system and evaluation service"""
    try:
        if use_enhanced:
            # Try to use enhanced persistent system first
            try:
                from src.core.enhanced_persistent_rag_system import EnhancedPersistentRAGSystem
                rag_system = EnhancedPersistentRAGSystem()
                logger.info("Enhanced Persistent RAG System initialized successfully")
                return rag_system, EvaluationService(), True
            except Exception as persistent_error:
                logger.warning(f"Enhanced Persistent system failed: {persistent_error}")
                logger.info("Falling back to Enhanced Structured system")
                rag_system = EnhancedStructuredRAGSystem()
                logger.info("Enhanced Structured RAG System initialized successfully")
                return rag_system, EvaluationService(), True
        else:
            # Try simple persistent system
            try:
                from src.core.simple_persistent_rag_system import SimplePersistentRAGSystem
                rag_system = SimplePersistentRAGSystem()
                logger.info("Simple Persistent RAG System initialized successfully")
                return rag_system, EvaluationService(), True
            except Exception as persistent_error:
                logger.warning(f"Simple Persistent system failed: {persistent_error}")
                rag_system = SAFEMBSERAGSystem()
                logger.info("Traditional RAG System initialized successfully")
                return rag_system, EvaluationService(), False
        
    except Exception as e:
        logger.error(f"Error initializing enhanced services: {str(e)}")
        logger.info("Falling back to traditional RAG system")
        # Final fallback to traditional system if all fails
        return SAFEMBSERAGSystem(), EvaluationService(), False

# Initialize project manager
@st.cache_resource
def init_project_manager(_rag_system):
    """Initialize and cache the project manager"""
    try:
        from ui.components.project_manager import ProjectManager
        project_manager = ProjectManager(_rag_system)
        logger.info("Project Manager initialized successfully")
        return project_manager
    except Exception as e:
        logger.error(f"Error initializing Project Manager: {str(e)}")
        return None

def load_project_chats(project_id=None):
    """Load saved chats for a specific project"""
    chats_file = Path("data/project_chats.json")
    chats_file.parent.mkdir(parents=True, exist_ok=True)
    
    if chats_file.exists():
        try:
            with open(chats_file, "r", encoding="utf-8") as f:
                all_chats = json.load(f)
                
            # Return chats for specific project or all if no project specified
            if project_id:
                return all_chats.get(str(project_id), {})
            else:
                return all_chats
        except Exception as e:
            logger.error(f"Error loading project chats: {str(e)}")
            return {}
    return {}

def save_project_chats(project_chats, project_id):
    """Save chats for a specific project"""
    try:
        chats_file = Path("data/project_chats.json")
        chats_file.parent.mkdir(parents=True, exist_ok=True)
        
        # Load existing chats
        all_chats = {}
        if chats_file.exists():
            try:
                with open(chats_file, "r", encoding="utf-8") as f:
                    all_chats = json.load(f)
            except:
                all_chats = {}
        
        # Update chats for this project
        all_chats[str(project_id)] = project_chats
        
        # Save all chats back
        with open(chats_file, "w", encoding="utf-8") as f:
            json.dump(all_chats, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Saved {len(project_chats)} chat sessions for project {project_id}")
    except Exception as e:
        logger.error(f"Error saving project chats: {str(e)}")

def load_chats():
    """Legacy function for backward compatibility"""
    return load_project_chats()

def save_chats(chats):
    """Legacy function for backward compatibility"""
    # For backward compatibility, save as global chats
    try:
        chats_file = Path("data/saved_chats.json")
        chats_file.parent.mkdir(parents=True, exist_ok=True)
        with open(chats_file, "w", encoding="utf-8") as f:
            json.dump(chats, f, indent=2, ensure_ascii=False)
        logger.info(f"Saved {len(chats)} global chat sessions")
    except Exception as e:
        logger.error(f"Error saving global chats: {str(e)}")

def get_enhanced_prompt(user_prompt, context_type=None):
    """Enhance user prompt with MBSE context"""
    if context_type and context_type != "None" and context_type in MBSE_CONTEXT_PROMPTS:
        enhanced_prompt = f"{MBSE_CONTEXT_PROMPTS[context_type]}\n\nUser Question: {user_prompt}"
        return enhanced_prompt
    return user_prompt

def display_arcadia_references():
    """Display Arcadia methodology reference links"""
    st.markdown("### ARCADIA Methodology References")
    st.markdown("Educational resources for each ARCADIA phase:")
    for phase, url in ARCADIA_REFERENCES.items():
        st.markdown(f"• [{phase}]({url})")

def get_fallback_example(example_type: str) -> str:
    """Provide fallback examples when files are not available."""
    
    if example_type == "safe":
        return """
# Systems Engineering Project Proposal

## Project Overview
Development of an intelligent transportation management system for smart city infrastructure using Model-Based Systems Engineering (MBSE) approach.

## Objectives
1. Design integrated traffic management system following ARCADIA methodology
2. Implement real-time traffic optimization algorithms  
3. Create stakeholder-centric interface for city operators
4. Establish quality assessment framework for system requirements

## Stakeholders
- Traffic Control Operators: Monitor and manage traffic flow
- City Planners: Strategic planning and infrastructure development
- Emergency Services: Priority routing and incident response
- Citizens: End users expecting efficient transportation

## Work Packages
WP1: Operational Analysis - Define stakeholder needs and use cases
WP2: System Analysis - Specify system-level functions and requirements
WP3: Logical Architecture - Design logical components and interfaces
WP4: Physical Architecture - Define implementation and deployment

## Technical Requirements
- The system shall process traffic data from 500+ sensors in real-time
- Response time for traffic signal adjustments must be under 2 seconds  
- System availability shall be 99.9% or higher
- Integration with emergency services dispatch systems is mandatory
"""
    
    elif example_type == "cyderco":
        return """
# Cybersecurity Infrastructure Project

## Project Description  
CYDERCO (Cyber Defense and Response Coordination) serves as a reference implementation for critical infrastructure protection systems.

## Mission Statement
Develop comprehensive cyber defense capabilities that can detect, analyze, and respond to sophisticated threats targeting critical infrastructure.

## Key Objectives
1. Implement advanced threat detection using machine learning
2. Develop automated incident response protocols
3. Create real-time monitoring and alerting systems
4. Establish coordination mechanisms for multi-agency response

## Stakeholders
- Security Operations Center (SOC) Analysts: Threat monitoring and analysis
- Incident Response Teams: Handle and contain security incidents  
- Infrastructure Operators: Maintain critical systems security
- Regulatory Bodies: Ensure compliance with security standards

## Functional Requirements
- Threat detection algorithms shall identify anomalies within 5 seconds
- Automated response must activate within 30 seconds of confirmed threat
- All security events must be logged for forensic analysis
- System shall integrate with existing SIEM platforms

## Non-Functional Requirements
- System availability must exceed 99.99%
- Threat detection accuracy shall be above 95%
- False positive rate must be below 2%
- Data encryption required for all communications
"""
    
    else:  # simple example
        return """
# Industrial Automation System

## Project Goal
Design and implement an intelligent manufacturing automation system that optimizes production efficiency while ensuring worker safety.

## Objectives
1. Automate production line operations and quality control
2. Implement predictive maintenance capabilities
3. Develop safety monitoring and emergency response systems
4. Create operator interface for system management

## Stakeholders
- Production Managers: Oversee manufacturing operations and efficiency
- Machine Operators: Interface with automated systems daily
- Quality Assurance Teams: Monitor and validate product quality
- Safety Officers: Ensure compliance with safety regulations

## System Requirements
- The system shall monitor production line status in real-time
- Predictive maintenance algorithms must forecast equipment failures 48 hours in advance
- Emergency stop functionality shall halt all operations within 0.5 seconds
- Quality control systems must inspect 100% of manufactured products
- Operator interfaces should provide intuitive control and monitoring capabilities
"""

def main():
    # Header with improved ARISE branding
    st.markdown("""
    <div class="main-header">
        <h1 style="margin-bottom: 0.5rem;">ARISE</h1>
        <p style="font-size: 1.3rem; margin: 0.5rem 0; font-weight: 300; opacity: 0.95;">
            ARCADIA Intelligent System Engine
        </p>
        <p style="font-size: 1.1rem; margin: 0.5rem 0 0 0; font-style: italic; opacity: 0.9;">
            Professional Model-Based Systems Engineering Platform
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    logger.info("=== ARCADIA Requirements Generator Started ===")
    
    # Initialize services
    rag_system, eval_service, is_enhanced = init_services(use_enhanced=True)
    logger.info("Core services initialized successfully")
    
    # Initialize project manager
    project_manager = init_project_manager(rag_system)
    has_project_management = project_manager is not None and project_manager.has_persistence
    
    # Initialize project-specific chat system in session state
    if "project_chats" not in st.session_state:
        st.session_state.project_chats = {}
    if "current_chat_id" not in st.session_state:
        st.session_state.current_chat_id = None
    if "current_project_chat_context" not in st.session_state:
        st.session_state.current_project_chat_context = None
    
    # Maintain backward compatibility with legacy chats
    if "chats" not in st.session_state:
        st.session_state.chats = load_chats()
    
    # Initialize project information for header (call render_project_sidebar once)
    current_project_id = None
    current_project = None
    
    if has_project_management:
        # Call project manager only once to avoid duplicate forms
        current_project_id = project_manager.render_project_sidebar()
        current_project = rag_system.get_current_project() if hasattr(rag_system, 'get_current_project') else None
        
        # Persistent Header with Project Status (always visible)
        header_col1, header_col2 = st.columns([1, 1])
        
        with header_col1:
            # Project status display
            if current_project:
                st.success(f"**Active Project:** {current_project.name}")
            else:
                st.info("No project selected")
        
        with header_col2:
            # Current project metrics

                
                # Log project session on startup
                if hasattr(rag_system, 'persistence_service'):
                    rag_system.persistence_service.log_project_session(
                        current_project_id, 
                        "app_access", 
                        "User accessed application interface"
                    )
        
        st.markdown("---")
    
    # Default configuration values for requirements analysis
    target_phase = "all"
    req_types = ["functional", "non_functional", "stakeholder"]
    export_format = "JSON"
    enable_structured_analysis = True
    enable_cross_phase_analysis = True
    
    # Simplified Sidebar
    with st.sidebar:

        
        # ARCADIA references (kept from original)
        display_arcadia_references()
        
        st.markdown("---")
        
        # System tools
        st.markdown("### System Tools")
        if st.button("Clear Cache", help="Clear system cache"):
            st.cache_resource.clear()
            st.success("Cache cleared")
        
        if st.button("Test System", help="Test RAG system"):
            try:
                test_rag = SAFEMBSERAGSystem()
                st.success("System operational")
            except Exception as e:
                st.error(f"System error: {str(e)}")
    
    # Main interface tabs - Phase 2 Reorganization: Combined Requirements & Analysis
    if has_project_management:
        # Professional tab layout
        tab1, tab2, tab3 = st.tabs([
            "Document Management",
            "Requirements & Analysis", 
            "Project Insights"
        ])
        
        with tab1:
            project_documents_tab(rag_system, current_project, has_project_management)
        
        with tab2:
            requirements_analysis_tab(rag_system, eval_service, target_phase, req_types, export_format, 
                                    enable_structured_analysis, enable_cross_phase_analysis, is_enhanced)
        
        with tab3:
            project_insights_tab(rag_system, current_project, has_project_management)
    
    elif is_enhanced:
        # Enhanced without project management - Phase 2 Reorganization
        tab1, tab2, tab3 = st.tabs([
            "Document Management",
            "Requirements & Analysis",
            "Insights"
        ])
        
        with tab1:
            project_documents_tab(rag_system, None, has_project_management)
        
        with tab2:
            requirements_analysis_tab(rag_system, eval_service, target_phase, req_types, export_format, 
                                    enable_structured_analysis, enable_cross_phase_analysis, is_enhanced)
        
        with tab3:
            project_insights_tab(rag_system, None, has_project_management)
    else:
        # Basic mode - Phase 2 Reorganization
        tab1, tab2, tab3 = st.tabs([
            "Document Management",
            "Requirements & Analysis",
            "Basic Insights"
        ])
        
        with tab1:
            project_documents_tab(rag_system, None, has_project_management)
        
        with tab2:
            requirements_analysis_tab(rag_system, eval_service, target_phase, req_types, export_format, 
                                    False, False, is_enhanced)
        
        with tab3:
            project_insights_tab(rag_system, None, has_project_management)

def generate_requirements_tab(rag_system, target_phase, req_types, export_format, 
                             enable_structured_analysis=False, enable_cross_phase_analysis=False, is_enhanced=False):
    st.markdown("### Requirements Generation")
    
    # Clear any previous file upload errors
    if 'upload_error' in st.session_state:
        del st.session_state['upload_error']
    
    # Input methods
    input_method = st.radio(
        "Input Method",
        ["Upload Document", "Paste Text", "Load Example"],
        key="input_method_radio"
    )
    
    proposal_text = ""
    
    if input_method == "Upload Document":
        st.markdown("#### Document Upload")
        
        # Add file size warning
        st.info("Maximum file size: 50MB • Supported formats: TXT, MD, PDF, DOCX")
        
        # Add error handling wrapper for file upload
        try:
            uploaded_file = st.file_uploader(
                "Upload project proposal",
                type=['txt', 'md', 'pdf', 'docx'],
                help="Upload your project proposal document",
                accept_multiple_files=False,
                key="document_uploader"
            )
            logger.info("File uploader component rendered successfully")
            
        except Exception as upload_error:
            logger.error(f"File uploader component error: {str(upload_error)}")
            st.error("File upload component error. Please refresh the page and try again.")
            uploaded_file = None
        
        if uploaded_file is not None:
            logger.info(f"File upload initiated: {uploaded_file.name}")
            logger.info(f"   • File type: {uploaded_file.type}")
            logger.info(f"   • File size: {uploaded_file.size} bytes ({uploaded_file.size / 1024 / 1024:.2f} MB)")
            
            # Check file size limit (50MB)
            max_size_mb = 50
            if uploaded_file.size > max_size_mb * 1024 * 1024:
                error_msg = f"File too large: {uploaded_file.size / 1024 / 1024:.2f} MB (max: {max_size_mb} MB)"
                logger.error(f"{error_msg}")
                st.error(error_msg)
                proposal_text = ""
            else:
                try:
                    logger.info("Starting document processing...")
                    
                    # Process uploaded file based on type
                    if uploaded_file.type == "text/plain" or uploaded_file.name.endswith('.txt'):
                        logger.info("Processing TXT file...")
                        file_content = uploaded_file.read()
                        logger.info(f"   • Raw content size: {len(file_content)} bytes")
                        proposal_text = str(file_content, "utf-8")
                        logger.info(f"TXT file processed successfully: {len(proposal_text)} characters")
                        
                    elif uploaded_file.name.endswith('.md'):
                        logger.info("Processing Markdown file...")
                        file_content = uploaded_file.read()
                        logger.info(f"   • Raw content size: {len(file_content)} bytes")
                        proposal_text = str(file_content, "utf-8")
                        logger.info(f"Markdown file processed successfully: {len(proposal_text)} characters")
                        
                    elif uploaded_file.type == "application/pdf" or uploaded_file.name.endswith('.pdf'):
                        logger.info("Processing PDF file...")
                        try:
                            import PyPDF2
                            import io
                            
                            # Read file content
                            file_content = uploaded_file.read()
                            logger.info(f"   • PDF raw size: {len(file_content)} bytes")
                            
                            # Create PDF reader
                            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                            logger.info(f"   • PDF pages detected: {len(pdf_reader.pages)}")
                            
                            proposal_text = ""
                            for page_num, page in enumerate(pdf_reader.pages):
                                logger.info(f"   • Processing page {page_num + 1}...")
                                try:
                                    page_text = page.extract_text()
                                    proposal_text += page_text + "\n"
                                    logger.info(f"     - Page {page_num + 1}: {len(page_text)} characters extracted")
                                except Exception as page_error:
                                    logger.warning(f"     - Page {page_num + 1} extraction failed: {str(page_error)}")
                            
                            if not proposal_text.strip():
                                logger.warning("PDF text extraction yielded empty result")
                                st.error("Could not extract text from PDF. The PDF might be image-based or encrypted. Please try copy-paste instead.")
                                proposal_text = ""
                            else:
                                logger.info(f"PDF processed successfully: {len(pdf_reader.pages)} pages, {len(proposal_text)} characters")
                                
                        except Exception as e:
                            logger.error(f"PDF processing error: {str(e)}")
                            logger.error(f"   • Error type: {type(e).__name__}")
                            st.error(f"Error processing PDF file: {str(e)}. Try using a different PDF or copy-paste the text.")
                            proposal_text = ""
                            
                    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or uploaded_file.name.endswith('.docx'):
                        logger.info("Processing DOCX file...")
                        try:
                            from docx import Document
                            import io
                            
                            # Read file content
                            file_content = uploaded_file.read()
                            logger.info(f"   • DOCX raw size: {len(file_content)} bytes")
                            
                            # Create document reader
                            doc = Document(io.BytesIO(file_content))
                            logger.info(f"   • DOCX paragraphs detected: {len(doc.paragraphs)}")
                            
                            proposal_text = ""
                            for para_num, paragraph in enumerate(doc.paragraphs):
                                para_text = paragraph.text
                                proposal_text += para_text + "\n"
                                if para_text.strip():  # Only log non-empty paragraphs
                                    logger.info(f"     - Paragraph {para_num + 1}: {len(para_text)} characters")
                                
                            if not proposal_text.strip():
                                logger.warning("DOCX text extraction yielded empty result")
                                st.error("Could not extract text from DOCX. The document might be empty or contain only images.")
                                proposal_text = ""
                            else:
                                logger.info(f"DOCX processed successfully: {len(doc.paragraphs)} paragraphs, {len(proposal_text)} characters")
                                
                        except Exception as e:
                            logger.error(f"DOCX processing error: {str(e)}")
                            logger.error(f"   • Error type: {type(e).__name__}")
                            st.error(f"Error processing DOCX file: {str(e)}. Try using a different file or copy-paste the text.")
                            proposal_text = ""
                            
                    else:
                        logger.warning(f"Unsupported file type: {uploaded_file.type}")
                        st.error(f"Unsupported file type: {uploaded_file.type}. Please upload TXT, MD, PDF, or DOCX files.")
                        proposal_text = ""
                        
                    # Show file info if successfully processed
                    if proposal_text:
                        st.success(f"File processed: {uploaded_file.name} ({len(proposal_text)} characters)")
                        logger.info(f"Document processing completed successfully!")
                        
                        # Show preview of extracted text
                        with st.expander("Preview extracted text (first 500 characters)"):
                            preview_text = proposal_text[:500] + "..." if len(proposal_text) > 500 else proposal_text
                            st.text(preview_text)
                        
                except Exception as e:
                    logger.error(f"File processing error: {str(e)}")
                    logger.error(f"   • Error type: {type(e).__name__}")
                    logger.error(f"   • Full traceback: ", exc_info=True)
                    st.error(f"Error reading file: {str(e)}")
                    proposal_text = ""
    
    elif input_method == "Paste Text":
        proposal_text = st.text_area(
            "Paste your project proposal text",
            height=300,
            help="Paste the text of your project proposal here"
        )
    
    elif input_method == "Load Example":
        example_choice = st.selectbox(
            "Choose Example",
            ["Transportation System", "Cybersecurity Infrastructure", "Industrial Automation"]
        )
        
        if st.button("Load Example"):
            try:
                if example_choice == "Transportation System":
                    # Load transportation system example
                    example_path = Path("data/examples/transportation_system.md")
                    if example_path.exists():
                        with open(example_path, "r", encoding="utf-8") as f:
                            proposal_text = f.read()
                    else:
                        st.info("Using built-in transportation system example.")
                        proposal_text = get_fallback_example("safe")
                        
                elif example_choice == "Cybersecurity Infrastructure":
                    # Load cybersecurity example (CYDERCO reference)
                    example_path = Path("data/examples/cyberco_reference.md")
                    if example_path.exists():
                        with open(example_path, "r", encoding="utf-8") as f:
                            proposal_text = f.read()
                    else:
                        st.info("Using built-in cybersecurity infrastructure example (CYDERCO reference).")
                        proposal_text = get_fallback_example("cyderco")
                        
                else:  # Industrial Automation
                    proposal_text = get_fallback_example("simple")
                    st.info("Using built-in industrial automation example.")
                    
                if proposal_text:
                    st.success(f"Example loaded: {example_choice}")
                    logger.info(f"Loaded example: {example_choice} ({len(proposal_text)} characters)")
                    
            except Exception as e:
                st.error(f"Error loading example: {str(e)}")
                logger.error(f"Error loading example {example_choice}: {str(e)}")
                proposal_text = get_fallback_example("simple")
    
    # Generation controls
    col1, col2 = st.columns([3, 1])
    
    with col1:
        st.markdown("#### Generation Parameters")
        
        # Advanced options
        with st.expander("Advanced Options"):
            max_requirements = st.slider("Max Requirements per Type", 5, 50, 20)
            include_rationale = st.checkbox("Include Rationale", True)
            include_verification = st.checkbox("Include Verification Methods", True)
            quality_threshold = st.slider("Quality Threshold", 0.5, 1.0, 0.7)
    
    with col2:
        st.markdown("#### Actions")
        generate_btn = st.button("Generate Requirements", type="primary")
        
        if proposal_text:
            st.success(f"Input ready ({len(proposal_text)} characters)")
    
    # Generation process
    if generate_btn and proposal_text:
        start_time = time.time()
        
        logger.info("Starting requirements generation process...")
        logger.info(f"Input: {len(proposal_text)} characters, Target Phase: {target_phase}, Types: {req_types}")
        logger.info(f"Generation started at: {datetime.now().strftime('%H:%M:%S')}")
        
        with st.spinner("Generating requirements..."):
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            # Step 1: Analyze proposal
            step1_start = time.time()
            status_text.text("Analyzing proposal...")
            logger.info("Step 1: Analyzing project proposal and extracting context...")
            progress_bar.progress(20)
            
            # Step 2: Generate requirements
            step2_start = time.time()
            step1_duration = step2_start - step1_start
            logger.info(f"Step 1 completed in {step1_duration:.1f} seconds")
            
            status_text.text("Generating requirements...")
            logger.info("Step 2: Initiating AI-driven requirements generation...")
            logger.info(f"Targeting ARCADIA phases: {target_phase}")
            logger.info(f"Requirement types: {', '.join(req_types)}")
            logger.info("Note: This may take several minutes due to LLM processing time...")
            
            # Estimate time based on phases and types
            phases_to_process = [target_phase] if target_phase != "all" else list(arcadia_config.ARCADIA_PHASES.keys())
            estimated_calls = len(phases_to_process) * len(req_types) + 1  # +1 for stakeholders
            logger.info(f"Estimated LLM calls: {estimated_calls} (each may take 30-120 seconds)")
            logger.info(f"Estimated total time: {estimated_calls * 60 / 60:.1f}-{estimated_calls * 120 / 60:.1f} minutes")
            
            progress_bar.progress(60)
            
            try:
                if is_enhanced and hasattr(rag_system, 'generate_enhanced_requirements_from_proposal'):
                    # Use enhanced generation with structured analysis
                    results = rag_system.generate_enhanced_requirements_from_proposal(
                        proposal_text=proposal_text,
                        target_phase=target_phase,
                        requirement_types=req_types,
                        enable_structured_analysis=enable_structured_analysis,
                        enable_cross_phase_analysis=enable_cross_phase_analysis
                    )
                    # Store enhanced results for the new tab
                    st.session_state['enhanced_results'] = results
                    logger.info(f"Enhanced generation completed - Traditional requirements: {len(results.get('requirements', {}))}")
                else:
                    # Use traditional generation
                    results = rag_system.generate_requirements_from_proposal(
                        proposal_text, target_phase, req_types
                    )
                    logger.info(f"Traditional generation completed - Requirements: {len(results.get('requirements', {}))}")
                
                # Auto-save to current project if available
                if hasattr(rag_system, 'get_current_project') and hasattr(rag_system, 'persistence_service'):
                    current_project = rag_system.get_current_project()
                    if current_project:
                        try:
                            # Save traditional requirements
                            traditional_requirements = results.get('traditional_requirements', results) if 'traditional_requirements' in results else results
                            success = rag_system.persistence_service.save_project_requirements(
                                current_project.id, 
                                traditional_requirements
                            )
                            
                            if success:
                                logger.info(f"Requirements auto-saved to project: {current_project.name}")
                                st.session_state['requirements_saved'] = True
                                
                                # Save stakeholders if available
                                if traditional_requirements.get('stakeholders'):
                                    stakeholders_list = []
                                    for stakeholder in traditional_requirements['stakeholders'].values():
                                        if isinstance(stakeholder, dict):
                                            stakeholders_list.append(stakeholder)
                                        elif isinstance(stakeholder, list):
                                            stakeholders_list.extend(stakeholder)
                                    
                                    if stakeholders_list:
                                        rag_system.persistence_service.save_stakeholders(current_project.id, stakeholders_list)
                                        logger.info(f"Stakeholders auto-saved: {len(stakeholders_list)}")
                                
                                # Save ARCADIA analyses if enhanced results available
                                if 'structured_analysis' in results:
                                    structured_analysis = results['structured_analysis']
                                    
                                    # Save each phase analysis
                                    if hasattr(structured_analysis, 'operational_analysis') and structured_analysis.operational_analysis:
                                        rag_system.persistence_service.save_arcadia_analysis(
                                            current_project.id, 
                                            'operational', 
                                            structured_analysis.operational_analysis.__dict__
                                        )
                                    
                                    if hasattr(structured_analysis, 'system_analysis') and structured_analysis.system_analysis:
                                        rag_system.persistence_service.save_arcadia_analysis(
                                            current_project.id, 
                                            'system', 
                                            structured_analysis.system_analysis.__dict__
                                        )
                                    
                                    if hasattr(structured_analysis, 'logical_architecture') and structured_analysis.logical_architecture:
                                        rag_system.persistence_service.save_arcadia_analysis(
                                            current_project.id, 
                                            'logical', 
                                            structured_analysis.logical_architecture.__dict__
                                        )
                                    
                                    if hasattr(structured_analysis, 'physical_architecture') and structured_analysis.physical_architecture:
                                        rag_system.persistence_service.save_arcadia_analysis(
                                            current_project.id, 
                                            'physical', 
                                            structured_analysis.physical_architecture.__dict__
                                        )
                                    
                                    if hasattr(structured_analysis, 'cross_phase_analysis') and structured_analysis.cross_phase_analysis:
                                        rag_system.persistence_service.save_arcadia_analysis(
                                            current_project.id, 
                                            'cross_phase', 
                                            structured_analysis.cross_phase_analysis.__dict__
                                        )
                                    
                                    logger.info("ARCADIA analyses auto-saved")
                                
                                # Log the session
                                rag_system.persistence_service.log_project_session(
                                    current_project.id,
                                    "requirements_generation",
                                    f"Generated requirements for phases: {target_phase}, types: {', '.join(req_types)}",
                                    {
                                        "target_phase": target_phase,
                                        "requirement_types": req_types,
                                        "total_requirements": sum(len(reqs) for phase_reqs in traditional_requirements.get('requirements', {}).values() for reqs in phase_reqs.values() if isinstance(reqs, list)),
                                        "generation_time": time.time() - start_time,
                                        "enhanced_analysis": enable_structured_analysis
                                    }
                                )
                            
                        except Exception as save_error:
                            logger.error(f"Error auto-saving requirements: {str(save_error)}")
                            st.session_state['save_error'] = str(save_error)
                
                step3_start = time.time()
                step2_duration = step3_start - step2_start
                logger.info(f"Step 2 completed in {step2_duration:.1f} seconds ({step2_duration/60:.1f} minutes)")
                logger.info("Requirements generation completed successfully")
                
                # Log generation statistics
                stats = results.get('statistics', {})
                total_reqs = stats.get('total_requirements', 0)
                logger.info(f"Generation Statistics:")
                logger.info(f"   • Total Requirements Generated: {total_reqs}")
                logger.info(f"   • Phases Covered: {len(results.get('requirements', {}))}")
                logger.info(f"   • Stakeholders Identified: {len(results.get('stakeholders', {}))}")
                
                # Log priority distribution
                priority_dist = stats.get('by_priority', {})
                logger.info(f"   • Priority Distribution:")
                logger.info(f"     - MUST (Critical): {priority_dist.get('MUST', 0)}")
                logger.info(f"     - SHOULD (Important): {priority_dist.get('SHOULD', 0)}")
                logger.info(f"     - COULD (Nice-to-have): {priority_dist.get('COULD', 0)}")
                
                # Log requirements by phase and type
                for phase, phase_reqs in results.get('requirements', {}).items():
                    phase_total = sum(len(reqs) if isinstance(reqs, list) else 0 for reqs in phase_reqs.values())
                    logger.info(f"   • {phase.upper()} Phase: {phase_total} requirements")
                    for req_type, reqs in phase_reqs.items():
                        count = len(reqs) if isinstance(reqs, list) else 0
                        logger.info(f"     - {req_type}: {count} requirements")
                
            except Exception as e:
                logger.error(f"Error during requirements generation: {str(e)}")
                st.error(f"Error during generation: {str(e)}")
                return
            
            # Step 3: Post-processing
            status_text.text("Post-processing...")
            logger.info("Step 3: Post-processing and quality validation...")
            progress_bar.progress(80)
            
            progress_bar.progress(100)
            status_text.text("Generation completed!")
            
            # Calculate total time
            total_time = time.time() - start_time
            logger.info(f"Requirements generation process completed successfully!")
            logger.info(f"Total generation time: {total_time:.1f} seconds ({total_time/60:.1f} minutes)")
            logger.info(f"Completed at: {datetime.now().strftime('%H:%M:%S')}")
            
            # Store results in session state
            st.session_state['generation_results'] = results
            st.session_state['proposal_text'] = proposal_text
            st.session_state['generation_time'] = total_time
    
    # Display results
    if 'generation_results' in st.session_state:
        results = st.session_state['generation_results']
        display_generation_results(results, export_format, rag_system)

def export_requirements_to_json(results):
    """Export requirements to JSON format with enhanced ARCADIA analysis"""
    # Check if enhanced results are available
    enhanced_results = st.session_state.get('enhanced_results')
    
    export_data = {
        "metadata": {
            "export_timestamp": datetime.now().isoformat(),
            "total_requirements": results.get('statistics', {}).get('total_requirements', 0),
            "phases_covered": len(results.get('requirements', {})),
            "stakeholders_count": len(results.get('stakeholders', {})),
            "enhanced_analysis_included": enhanced_results is not None
        },
        "requirements": results.get('requirements', {}),
        "stakeholders": results.get('stakeholders', {}),
        "statistics": results.get('statistics', {})
    }
    
    # Add enhanced ARCADIA analysis if available
    if enhanced_results and enhanced_results.get('structured_analysis'):
        structured_analysis = enhanced_results['structured_analysis']
        
        export_data["arcadia_analysis"] = {}
        
        # Operational Analysis
        if hasattr(structured_analysis, 'operational_analysis') and structured_analysis.operational_analysis:
            op_analysis = structured_analysis.operational_analysis
            export_data["arcadia_analysis"]["operational_analysis"] = {
                "actors": [
                    {
                        "id": actor.id,
                        "name": actor.name,
                        "role_definition": actor.role_definition,
                        "description": actor.description,
                        "responsibilities": actor.responsibilities,
                        "capabilities": actor.capabilities
                    } for actor in (op_analysis.actors or [])
                ],
                "capabilities": [
                    {
                        "id": cap.id,
                        "name": cap.name,
                        "mission_statement": cap.mission_statement,
                        "description": cap.description,
                        "involved_actors": cap.involved_actors,
                        "performance_constraints": cap.performance_constraints
                    } for cap in (op_analysis.capabilities or [])
                ],
                "scenarios": [
                    {
                        "id": scenario.id,
                        "name": scenario.name,
                        "scenario_type": scenario.scenario_type,
                        "description": scenario.description
                    } for scenario in (op_analysis.scenarios or [])
                ]
            }
        
        # System Analysis
        if hasattr(structured_analysis, 'system_analysis') and structured_analysis.system_analysis:
            sys_analysis = structured_analysis.system_analysis
            export_data["arcadia_analysis"]["system_analysis"] = {
                "system_boundary": {
                    "scope_definition": sys_analysis.system_boundary.scope_definition,
                    "included_elements": sys_analysis.system_boundary.included_elements,
                    "excluded_elements": sys_analysis.system_boundary.excluded_elements
                } if sys_analysis.system_boundary else {},
                "functions": [
                    {
                        "id": func.id,
                        "name": func.name,
                        "function_type": func.function_type,
                        "description": func.description,
                        "allocated_actors": func.allocated_actors,
                        "performance_requirements": func.performance_requirements
                    } for func in (sys_analysis.functions or [])
                ],
                "capabilities": [
                    {
                        "id": cap.id,
                        "name": cap.name,
                        "description": cap.description,
                        "implementing_functions": cap.implementing_functions
                    } for cap in (sys_analysis.capabilities or [])
                ]
            }
        
        # Logical Architecture
        if hasattr(structured_analysis, 'logical_architecture') and structured_analysis.logical_architecture:
            logical_arch = structured_analysis.logical_architecture
            export_data["arcadia_analysis"]["logical_architecture"] = {
                "components": [
                    {
                        "id": comp.id,
                        "name": comp.name,
                        "component_type": comp.component_type,
                        "description": comp.description,
                        "responsibilities": comp.responsibilities,
                        "sub_components": comp.sub_components
                    } for comp in (logical_arch.components or [])
                ],
                "functions": [
                    {
                        "id": func.id,
                        "name": func.name,
                        "behavioral_specification": "; ".join([model.get('spec', str(model)) for model in func.behavioral_models]) if func.behavioral_models else "",
                        "description": func.description,
                        "allocated_components": func.allocated_components,
                        "input_flows": [str(iface) for iface in func.input_interfaces] if func.input_interfaces else [],
                        "output_flows": [str(iface) for iface in func.output_interfaces] if func.output_interfaces else []
                    } for func in (logical_arch.functions or [])
                ],
                "interfaces": [
                    {
                        "id": iface.id,
                        "name": iface.name,
                        "protocol": iface.protocol,
                        "provider": iface.provider,
                        "consumer": iface.consumer,
                        "data_elements": iface.data_elements
                    } for iface in (logical_arch.interfaces or [])
                ]
            }
        
        # Physical Architecture
        if hasattr(structured_analysis, 'physical_architecture') and structured_analysis.physical_architecture:
            physical_arch = structured_analysis.physical_architecture
            export_data["arcadia_analysis"]["physical_architecture"] = {
                "components": [
                    {
                        "id": comp.id,
                        "name": comp.name,
                        "component_type": comp.component_type,
                        "technology_platform": comp.technology_platform,
                        "description": comp.description,
                        "resource_requirements": comp.resource_requirements,
                        "deployment_location": comp.deployment_location
                    } for comp in (physical_arch.components or [])
                ],
                "constraints": [
                    {
                        "constraint_type": const.constraint_type,
                        "description": const.description,
                        "rationale": getattr(const, 'rationale', '')
                    } for const in (physical_arch.constraints or [])
                ],
                "functions": [
                    {
                        "id": func.id,
                        "name": func.name,
                        "description": func.description,
                        "resource_requirements": func.resource_requirements,
                        "timing_constraints": func.timing_constraints
                    } for func in (physical_arch.functions or [])
                ]
            }
        
        # Cross-Phase Analysis
        if hasattr(structured_analysis, 'cross_phase_analysis') and structured_analysis.cross_phase_analysis:
            cross_phase = structured_analysis.cross_phase_analysis
            export_data["arcadia_analysis"]["cross_phase_analysis"] = {
                "traceability_links": [
                    {
                        "source_element": link.source_element,
                        "target_element": link.target_element,
                        "relationship_type": link.relationship_type,
                        "confidence_score": link.confidence_score,
                        "source_phase": link.source_phase.value if hasattr(link.source_phase, 'value') else str(link.source_phase),
                        "target_phase": link.target_phase.value if hasattr(link.target_phase, 'value') else str(link.target_phase)
                    } for link in (cross_phase.traceability_links or [])
                ],
                "gap_analysis": [
                    {
                        "gap_type": gap.gap_type,
                        "description": gap.description,
                        "severity": gap.severity,
                        "recommendations": gap.recommendations
                    } for gap in (cross_phase.gap_analysis or [])
                ],
                "coverage_matrix": cross_phase.coverage_matrix or {}
            }
        
        # Add enhancement summary
        if enhanced_results.get('enhancement_summary'):
            export_data["arcadia_analysis"]["enhancement_summary"] = enhanced_results['enhancement_summary']
    
    return json.dumps(export_data, indent=2, ensure_ascii=False)

def export_requirements_to_csv(results):
    """Export requirements to CSV format with enhanced ARCADIA analysis"""
    import io
    
    output = io.StringIO()
    enhanced_results = st.session_state.get('enhanced_results')
    
    # Requirements section
    output.write("=== REQUIREMENTS ===\n")
    output.write("Phase,Type,ID,Title,Description,Priority,Verification Method,Rationale\n")
    
    # Process requirements by phase
    for phase, phase_reqs in results.get('requirements', {}).items():
        for req_type, reqs in phase_reqs.items():
            if isinstance(reqs, list):
                for req in reqs:
                    row = [
                        csv_escape(phase),
                        csv_escape(req_type),
                        csv_escape(req.get('id', '')),
                        csv_escape(req.get('title', '')),
                        csv_escape(req.get('description', '')),
                        csv_escape(req.get('priority', '')),
                        csv_escape(req.get('verification_method', '')),
                        csv_escape(req.get('rationale', ''))
                    ]
                    output.write(','.join(row) + '\n')
    
    # Add ARCADIA analysis sections if available
    if enhanced_results and enhanced_results.get('structured_analysis'):
        structured_analysis = enhanced_results['structured_analysis']
        
        # Operational Actors
        if hasattr(structured_analysis, 'operational_analysis') and structured_analysis.operational_analysis:
            op_analysis = structured_analysis.operational_analysis
            if op_analysis.actors:
                output.write("\n=== OPERATIONAL ACTORS ===\n")
                output.write("ID,Name,Role,Description,Responsibilities,Capabilities\n")
                for actor in op_analysis.actors:
                    row = [
                        csv_escape(actor.id),
                        csv_escape(actor.name),
                        csv_escape(actor.role_definition),
                        csv_escape(actor.description),
                        csv_escape('; '.join(actor.responsibilities) if actor.responsibilities else ''),
                        csv_escape('; '.join(actor.capabilities) if actor.capabilities else '')
                    ]
                    output.write(','.join(row) + '\n')
            
            # Operational Capabilities
            if op_analysis.capabilities:
                output.write("\n=== OPERATIONAL CAPABILITIES ===\n")
                output.write("ID,Name,Mission Statement,Description,Involved Actors,Performance Constraints\n")
                for cap in op_analysis.capabilities:
                    row = [
                        csv_escape(cap.id),
                        csv_escape(cap.name),
                        csv_escape(cap.mission_statement),
                        csv_escape(cap.description),
                        csv_escape('; '.join(cap.involved_actors) if cap.involved_actors else ''),
                        csv_escape('; '.join(cap.performance_constraints) if cap.performance_constraints else '')
                    ]
                    output.write(','.join(row) + '\n')
        
        # System Functions
        if hasattr(structured_analysis, 'system_analysis') and structured_analysis.system_analysis:
            sys_analysis = structured_analysis.system_analysis
            if sys_analysis.functions:
                output.write("\n=== SYSTEM FUNCTIONS ===\n")
                output.write("ID,Name,Type,Description,Allocated Actors,Performance Requirements\n")
                for func in sys_analysis.functions:
                    row = [
                        csv_escape(func.id),
                        csv_escape(func.name),
                        csv_escape(func.function_type),
                        csv_escape(func.description),
                        csv_escape('; '.join(func.allocated_actors) if func.allocated_actors else ''),
                        csv_escape('; '.join(func.performance_requirements) if func.performance_requirements else '')
                    ]
                    output.write(','.join(row) + '\n')
        
        # Logical Components
        if hasattr(structured_analysis, 'logical_architecture') and structured_analysis.logical_architecture:
            logical_arch = structured_analysis.logical_architecture
            if logical_arch.components:
                output.write("\n=== LOGICAL COMPONENTS ===\n")
                output.write("ID,Name,Type,Description,Responsibilities,Sub Components\n")
                for comp in logical_arch.components:
                    row = [
                        csv_escape(comp.id),
                        csv_escape(comp.name),
                        csv_escape(comp.component_type),
                        csv_escape(comp.description),
                        csv_escape('; '.join(comp.responsibilities) if comp.responsibilities else ''),
                        csv_escape('; '.join(comp.sub_components) if comp.sub_components else '')
                    ]
                    output.write(','.join(row) + '\n')
        
        # Physical Components
        if hasattr(structured_analysis, 'physical_architecture') and structured_analysis.physical_architecture:
            physical_arch = structured_analysis.physical_architecture
            if physical_arch.components:
                output.write("\n=== PHYSICAL COMPONENTS ===\n")
                output.write("ID,Name,Type,Technology Platform,Description,Resource Requirements,Deployment Location\n")
                for comp in physical_arch.components:
                    row = [
                        csv_escape(comp.id),
                        csv_escape(comp.name),
                        csv_escape(comp.component_type),
                        csv_escape(comp.technology_platform),
                        csv_escape(comp.description),
                        csv_escape('; '.join(comp.resource_requirements) if comp.resource_requirements else ''),
                        csv_escape(comp.deployment_location)
                    ]
                    output.write(','.join(row) + '\n')
        
        # Traceability Links
        if hasattr(structured_analysis, 'cross_phase_analysis') and structured_analysis.cross_phase_analysis:
            cross_phase = structured_analysis.cross_phase_analysis
            if cross_phase.traceability_links:
                output.write("\n=== TRACEABILITY LINKS ===\n")
                output.write("Source Element,Target Element,Relationship Type,Confidence Score,Source Phase,Target Phase\n")
                for link in cross_phase.traceability_links:
                    row = [
                        csv_escape(link.source_element),
                        csv_escape(link.target_element),
                        csv_escape(link.relationship_type),
                        csv_escape(str(link.confidence_score)),
                        csv_escape(link.source_phase.value if hasattr(link.source_phase, 'value') else str(link.source_phase)),
                        csv_escape(link.target_phase.value if hasattr(link.target_phase, 'value') else str(link.target_phase))
                    ]
                    output.write(','.join(row) + '\n')
    
    # Stakeholders section
    output.write("\n=== STAKEHOLDERS ===\n")
    output.write("Stakeholder,Role,Description\n")
    
    for stakeholder_name, stakeholder_data in results.get('stakeholders', {}).items():
        if isinstance(stakeholder_data, dict):
            role = stakeholder_data.get('role', '')
            desc = stakeholder_data.get('description', '')
        else:
            role = ''
            desc = str(stakeholder_data) if stakeholder_data else ''
        
        row = [
            csv_escape(stakeholder_name),
            csv_escape(role),
            csv_escape(desc)
        ]
        output.write(','.join(row) + '\n')
    
    return output.getvalue()

def csv_escape(value):
    """Helper function to escape CSV values"""
    if value is None:
        return ""
    value = str(value).replace('"', '""')  # Escape quotes
    if ',' in value or '"' in value or '\n' in value:
        return f'"{value}"'
    return value

def export_requirements_to_excel_csv(results):
    """Export requirements to Excel-compatible CSV format with comprehensive ARCADIA analysis"""
    import io
    
    output = io.StringIO()
    enhanced_results = st.session_state.get('enhanced_results')
    
    # Main requirements sheet
    output.write("=== REQUIREMENTS ===\n")
    output.write("Phase,Type,ID,Title,Description,Priority,Verification Method,Priority Confidence,Created Date\n")
    
    for phase, phase_reqs in results.get('requirements', {}).items():
        for req_type, reqs in phase_reqs.items():
            if isinstance(reqs, list):
                for req in reqs:
                    row = [
                        csv_escape(phase),
                        csv_escape(req_type),
                        csv_escape(req.get('id', '')),
                        csv_escape(req.get('title', '')),
                        csv_escape(req.get('description', '')),
                        csv_escape(req.get('priority', '')),
                        csv_escape(req.get('verification_method', '')),
                        csv_escape(req.get('priority_confidence', '')),
                        csv_escape(datetime.now().strftime('%Y-%m-%d'))
                    ]
                    output.write(','.join(row) + '\n')
    
    # Add comprehensive ARCADIA analysis sections if available
    if enhanced_results and enhanced_results.get('structured_analysis'):
        structured_analysis = enhanced_results['structured_analysis']
        
        # Operational Analysis - Actors
        if hasattr(structured_analysis, 'operational_analysis') and structured_analysis.operational_analysis:
            op_analysis = structured_analysis.operational_analysis
            if op_analysis.actors:
                output.write("\n=== OPERATIONAL ACTORS ===\n")
                output.write("ID,Name,Role Definition,Description,Responsibilities,Capabilities,Phase\n")
                for actor in op_analysis.actors:
                    row = [
                        csv_escape(actor.id),
                        csv_escape(actor.name),
                        csv_escape(actor.role_definition),
                        csv_escape(actor.description),
                        csv_escape('; '.join(actor.responsibilities) if actor.responsibilities else ''),
                        csv_escape('; '.join(actor.capabilities) if actor.capabilities else ''),
                        csv_escape('Operational')
                    ]
                    output.write(','.join(row) + '\n')
            
            # Operational Capabilities
            if op_analysis.capabilities:
                output.write("\n=== OPERATIONAL CAPABILITIES ===\n")
                output.write("ID,Name,Mission Statement,Description,Involved Actors,Performance Constraints,Phase\n")
                for cap in op_analysis.capabilities:
                    row = [
                        csv_escape(cap.id),
                        csv_escape(cap.name),
                        csv_escape(cap.mission_statement),
                        csv_escape(cap.description),
                        csv_escape('; '.join(cap.involved_actors) if cap.involved_actors else ''),
                        csv_escape('; '.join(cap.performance_constraints) if cap.performance_constraints else ''),
                        csv_escape('Operational')
                    ]
                    output.write(','.join(row) + '\n')
            
            # Operational Scenarios
            if op_analysis.scenarios:
                output.write("\n=== OPERATIONAL SCENARIOS ===\n")
                output.write("ID,Name,Scenario Type,Description,Phase\n")
                for scenario in op_analysis.scenarios:
                    row = [
                        csv_escape(scenario.id),
                        csv_escape(scenario.name),
                        csv_escape(scenario.scenario_type),
                        csv_escape(scenario.description),
                        csv_escape('Operational')
                    ]
                    output.write(','.join(row) + '\n')
        
        # System Analysis
        if hasattr(structured_analysis, 'system_analysis') and structured_analysis.system_analysis:
            sys_analysis = structured_analysis.system_analysis
            
            # System Boundary
            if sys_analysis.system_boundary:
                output.write("\n=== SYSTEM BOUNDARY ===\n")
                output.write("Scope Definition,Included Elements,Excluded Elements\n")
                row = [
                    csv_escape(sys_analysis.system_boundary.scope_definition),
                    csv_escape('; '.join(sys_analysis.system_boundary.included_elements) if sys_analysis.system_boundary.included_elements else ''),
                    csv_escape('; '.join(sys_analysis.system_boundary.excluded_elements) if sys_analysis.system_boundary.excluded_elements else '')
                ]
                output.write(','.join(row) + '\n')
            
            # System Functions
            if sys_analysis.functions:
                output.write("\n=== SYSTEM FUNCTIONS ===\n")
                output.write("ID,Name,Function Type,Description,Allocated Actors,Performance Requirements,Phase\n")
                for func in sys_analysis.functions:
                    row = [
                        csv_escape(func.id),
                        csv_escape(func.name),
                        csv_escape(func.function_type),
                        csv_escape(func.description),
                        csv_escape('; '.join(func.allocated_actors) if func.allocated_actors else ''),
                        csv_escape('; '.join(func.performance_requirements) if func.performance_requirements else ''),
                        csv_escape('System')
                    ]
                    output.write(','.join(row) + '\n')
        
        # Logical Architecture
        if hasattr(structured_analysis, 'logical_architecture') and structured_analysis.logical_architecture:
            logical_arch = structured_analysis.logical_architecture
            
            # Logical Components
            if logical_arch.components:
                output.write("\n=== LOGICAL COMPONENTS ===\n")
                output.write("ID,Name,Component Type,Description,Responsibilities,Sub Components,Phase\n")
                for comp in logical_arch.components:
                    row = [
                        csv_escape(comp.id),
                        csv_escape(comp.name),
                        csv_escape(comp.component_type),
                        csv_escape(comp.description),
                        csv_escape('; '.join(comp.responsibilities) if comp.responsibilities else ''),
                        csv_escape('; '.join(comp.sub_components) if comp.sub_components else ''),
                        csv_escape('Logical')
                    ]
                    output.write(','.join(row) + '\n')
            
            # Logical Functions
            if logical_arch.functions:
                output.write("\n=== LOGICAL FUNCTIONS ===\n")
                output.write("ID,Name,Behavioral Specification,Description,Allocated Components,Input Flows,Output Flows,Phase\n")
                for func in logical_arch.functions:
                    # Format behavioral models for CSV
                    behavioral_spec = "; ".join([model.get('spec', str(model)) for model in func.behavioral_models]) if func.behavioral_models else ""
                    row = [
                        csv_escape(func.id),
                        csv_escape(func.name),
                        csv_escape(behavioral_spec),
                        csv_escape(func.description),
                        csv_escape('; '.join(func.allocated_components) if func.allocated_components else ''),
                        csv_escape('; '.join([str(iface) for iface in func.input_interfaces]) if func.input_interfaces else ''),
                        csv_escape('; '.join([str(iface) for iface in func.output_interfaces]) if func.output_interfaces else ''),
                        csv_escape('Logical')
                    ]
                    output.write(','.join(row) + '\n')
            
            # Logical Interfaces
            if logical_arch.interfaces:
                output.write("\n=== LOGICAL INTERFACES ===\n")
                output.write("ID,Name,Protocol,Provider,Consumer,Data Elements,Phase\n")
                for iface in logical_arch.interfaces:
                    row = [
                        csv_escape(iface.id),
                        csv_escape(iface.name),
                        csv_escape(iface.protocol),
                        csv_escape(iface.provider),
                        csv_escape(iface.consumer),
                        csv_escape('; '.join(iface.data_elements) if iface.data_elements else ''),
                        csv_escape('Logical')
                    ]
                    output.write(','.join(row) + '\n')
        
        # Physical Architecture
        if hasattr(structured_analysis, 'physical_architecture') and structured_analysis.physical_architecture:
            physical_arch = structured_analysis.physical_architecture
            
            # Physical Components
            if physical_arch.components:
                output.write("\n=== PHYSICAL COMPONENTS ===\n")
                output.write("ID,Name,Component Type,Technology Platform,Description,Resource Requirements,Deployment Location,Phase\n")
                for comp in physical_arch.components:
                    row = [
                        csv_escape(comp.id),
                        csv_escape(comp.name),
                        csv_escape(comp.component_type),
                        csv_escape(comp.technology_platform),
                        csv_escape(comp.description),
                        csv_escape('; '.join(comp.resource_requirements) if comp.resource_requirements else ''),
                        csv_escape(comp.deployment_location),
                        csv_escape('Physical')
                    ]
                    output.write(','.join(row) + '\n')
            
            # Implementation Constraints
            if physical_arch.constraints:
                output.write("\n=== IMPLEMENTATION CONSTRAINTS ===\n")
                output.write("Constraint Type,Description,Rationale,Phase\n")
                for const in physical_arch.constraints:
                    row = [
                        csv_escape(const.constraint_type),
                        csv_escape(const.description),
                        csv_escape(getattr(const, 'rationale', '')),
                        csv_escape('Physical')
                    ]
                    output.write(','.join(row) + '\n')
        
        # Cross-Phase Analysis
        if hasattr(structured_analysis, 'cross_phase_analysis') and structured_analysis.cross_phase_analysis:
            cross_phase = structured_analysis.cross_phase_analysis
            
            # Traceability Links
            if cross_phase.traceability_links:
                output.write("\n=== TRACEABILITY LINKS ===\n")
                output.write("Source Element,Target Element,Relationship Type,Confidence Score,Source Phase,Target Phase\n")
                for link in cross_phase.traceability_links:
                    row = [
                        csv_escape(link.source_element),
                        csv_escape(link.target_element),
                        csv_escape(link.relationship_type),
                        csv_escape(str(link.confidence_score)),
                        csv_escape(link.source_phase.value if hasattr(link.source_phase, 'value') else str(link.source_phase)),
                        csv_escape(link.target_phase.value if hasattr(link.target_phase, 'value') else str(link.target_phase))
                    ]
                    output.write(','.join(row) + '\n')
            
            # Gap Analysis
            if cross_phase.gap_analysis:
                output.write("\n=== GAP ANALYSIS ===\n")
                output.write("Gap Type,Description,Severity,Recommendations\n")
                for gap in cross_phase.gap_analysis:
                    row = [
                        csv_escape(gap.gap_type),
                        csv_escape(gap.description),
                        csv_escape(gap.severity),
                        csv_escape('; '.join(gap.recommendations) if gap.recommendations else '')
                    ]
                    output.write(','.join(row) + '\n')
        
        # Enhancement Summary
        if enhanced_results.get('enhancement_summary'):
            summary = enhanced_results['enhancement_summary']
            output.write("\n=== ENHANCEMENT SUMMARY ===\n")
            output.write("Metric,Value\n")
            for key, value in summary.items():
                if isinstance(value, list):
                    output.write(f'"{key}","{"; ".join(value)}"\n')
                else:
                    output.write(f'"{key}","{value}"\n')
    
    # Add stakeholders section
    output.write("\n=== STAKEHOLDERS ===\n")
    output.write("Stakeholder,Role,Description,Type,Influence,Phase,Mentions\n")
    
    for stakeholder_name, stakeholder_data in results.get('stakeholders', {}).items():
        if isinstance(stakeholder_data, dict):
            role = stakeholder_data.get('role', '')
            desc = stakeholder_data.get('description', '')
            stakeholder_type = stakeholder_data.get('type', '')
            influence = stakeholder_data.get('influence', '')
            phase = stakeholder_data.get('phase', '')
            mentions = stakeholder_data.get('mentions_count', '')
        else:
            role = ''
            desc = str(stakeholder_data) if stakeholder_data else ''
            stakeholder_type = ''
            influence = ''
            phase = ''
            mentions = ''
        
        row = [
            csv_escape(stakeholder_name),
            csv_escape(role),
            csv_escape(desc),
            csv_escape(stakeholder_type),
            csv_escape(influence),
            csv_escape(phase),
            csv_escape(str(mentions))
        ]
        output.write(','.join(row) + '\n')
    
    # Add statistics section
    stats = results.get('statistics', {})
    if stats:
        output.write("\n=== STATISTICS ===\n")
        output.write("Metric,Value\n")
        
        for metric, value in stats.items():
            if metric == 'by_priority' and isinstance(value, dict):
                for priority, count in value.items():
                    output.write(f'"{metric}_{priority}","{count}"\n')
            elif metric == 'by_phase' and isinstance(value, dict):
                for phase, phase_data in value.items():
                    if isinstance(phase_data, dict):
                        for sub_metric, sub_value in phase_data.items():
                            output.write(f'"{metric}_{phase}_{sub_metric}","{sub_value}"\n')
                    else:
                        output.write(f'"{metric}_{phase}","{phase_data}"\n')
            else:
                output.write(f'"{metric}","{value}"\n')
    
    return output.getvalue()

def display_generation_results(results, export_format, rag_system):
    st.markdown("### Generated Requirements")
    
    # Project persistence information
    if hasattr(rag_system, 'get_current_project'):
        current_project = rag_system.get_current_project()
        if current_project:
            if st.session_state.get('requirements_saved'):
                st.success(f"Requirements automatically saved to project: {current_project.name}")
            elif st.session_state.get('save_error'):
                st.error(f"Auto-save failed: {st.session_state['save_error']}")
                if st.button("Retry Save"):
                    try:
                        # Retry saving
                        traditional_requirements = results.get('traditional_requirements', results) if 'traditional_requirements' in results else results
                        success = rag_system.persistence_service.save_project_requirements(
                            current_project.id, 
                            traditional_requirements
                        )
                        if success:
                            st.success("Requirements saved successfully!")
                            st.session_state['requirements_saved'] = True
                            del st.session_state['save_error']
                            st.rerun()
                    except Exception as e:
                        st.error(f"Retry failed: {str(e)}")
        else:
            st.info("Tip: Create or select a project to automatically save requirements for future access!")
    
    # Performance information
    if 'generation_time' in st.session_state:
        generation_time = st.session_state['generation_time']
        with st.expander("Performance Information"):
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Generation Time", f"{generation_time:.1f}s")
            with col2:
                st.metric("Time (minutes)", f"{generation_time/60:.1f}min")
            with col3:
                avg_per_req = generation_time / max(1, results.get('statistics', {}).get('total_requirements', 1))
                st.metric("Avg per Requirement", f"{avg_per_req:.1f}s")
            
            st.info("""
            **Why does generation take time?**
            - Each ARCADIA phase requires separate LLM processing
            - Complex reasoning for requirements quality and traceability  
            - Multiple requirement types (functional, non-functional, stakeholder)
            - Network latency to Ollama server
            """)
    
    # Statistics overview
    stats = results.get('statistics', {})
    priority_dist = stats.get('by_priority', {})
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total Requirements", stats.get('total_requirements', 0))
    with col2:
        st.metric("Stakeholders", len(results.get('stakeholders', {})))
    with col3:
        st.metric("Phases Covered", len(results.get('requirements', {})))
    with col4:
        st.metric("Quality Score", f"{stats.get('average_quality', 0):.1%}")
    
    # Priority distribution overview
    if priority_dist:
        st.markdown("#### Priority Distribution Analysis")
        
        priority_col1, priority_col2, priority_col3 = st.columns(3)
        total_reqs = max(stats.get('total_requirements', 1), 1)
        
        with priority_col1:
            must_count = priority_dist.get('MUST', 0)
            must_percentage = (must_count / total_reqs) * 100
            st.metric("MUST (Critical)", must_count, delta=f"{must_percentage:.1f}%")
            
        with priority_col2:
            should_count = priority_dist.get('SHOULD', 0)
            should_percentage = (should_count / total_reqs) * 100
            st.metric("SHOULD (Important)", should_count, delta=f"{should_percentage:.1f}%")
            
        with priority_col3:
            could_count = priority_dist.get('COULD', 0)
            could_percentage = (could_count / total_reqs) * 100
            st.metric("COULD (Nice-to-have)", could_count, delta=f"{could_percentage:.1f}%")
        
        # Priority distribution chart
        import plotly.express as px
        priority_data = pd.DataFrame([
            {"Priority": "MUST (Critical)", "Count": must_count, "Percentage": must_percentage},
            {"Priority": "SHOULD (Important)", "Count": should_count, "Percentage": should_percentage},
            {"Priority": "COULD (Nice-to-have)", "Count": could_count, "Percentage": could_percentage}
        ])
        
        if priority_data['Count'].sum() > 0:
            fig = px.pie(priority_data, values='Count', names='Priority', 
                       title="Requirements Priority Distribution (ARCADIA-Compliant)",
                       color_discrete_map={
                           "MUST (Critical)": "#FF4B4B",
                           "SHOULD (Important)": "#FFA500", 
                           "COULD (Nice-to-have)": "#00D4AA"
                       })
            fig.update_layout(height=400)
            st.plotly_chart(fig, use_container_width=True)
            
            # Document-driven distribution analysis
            st.info("""
            **ARCADIA Priority Guidelines:**
            - MUST (SHALL): Core security, regulatory compliance, safety-critical operations
            - SHOULD: Important operational features, performance requirements  
            - COULD: Enhancement features, convenience functions
            
            **Note:** Priority distribution reflects the actual content and nature of your technical documents.
            Different project types naturally have different priority profiles.
            """)
            
            # Document-content based priority assessment
            if must_percentage > 70:
                st.info("Safety/Security-Critical System: High proportion of MUST requirements reflects critical system nature")
            elif must_percentage > 50:
                st.info("Mission-Critical System: Moderate-high MUST requirements indicate operational importance")
            elif must_percentage < 10:
                st.info("Research/Enhancement Project: Low MUST requirements suggest experimental or enhancement nature")
            else:
                st.success("Balanced System: Priority distribution reflects mixed criticality levels")
    
    # Requirements by phase
    for phase, phase_reqs in results.get('requirements', {}).items():
        phase_info = arcadia_config.ARCADIA_PHASES.get(phase, {})
        
        st.markdown(f"""
        <div class="phase-section">
            <h3>{phase_info.get('name', phase.title())} Phase</h3>
        </div>
        """, unsafe_allow_html=True)
        
        # Phase description
        st.write(phase_info.get('description', ''))
        
        # Requirements tabs for this phase
        if phase_reqs:
            phase_tabs = st.tabs([req_type.title() for req_type in phase_reqs.keys()])
            
            for i, (req_type, reqs) in enumerate(phase_reqs.items()):
                with phase_tabs[i]:
                    display_requirements_list(reqs, req_type)
    
    # Export section
    st.markdown("### Export Requirements")
    
    # Get the appropriate results for export
    enhanced_results = st.session_state.get('enhanced_results')
    export_results = enhanced_results.get('traditional_requirements', results) if enhanced_results else results
    
    col1, col2 = st.columns([2, 1])
    with col1:
        # Generate export content based on format
        try:
            if export_format == "JSON":
                exported_content = export_requirements_to_json(export_results)
                filename = "requirements.json"
                mime_type = "application/json"
                
            elif export_format == "CSV":
                exported_content = export_requirements_to_csv(export_results)
                filename = "requirements.csv"
                mime_type = "text/csv"
                
            elif export_format == "Excel":
                exported_content = export_requirements_to_excel_csv(export_results)
                filename = "requirements_excel.csv"
                mime_type = "text/csv"
                
            elif export_format == "ARCADIA_JSON" and enhanced_results:
                # Export enhanced ARCADIA analysis
                arcadia_export = {
                    "metadata": {
                        "export_timestamp": datetime.now().isoformat(),
                        "export_type": "ARCADIA_Analysis"
                    },
                    "traditional_requirements": enhanced_results.get('traditional_requirements', {}),
                    "structured_analysis": enhanced_results.get('structured_analysis', {}).__dict__ if hasattr(enhanced_results.get('structured_analysis', {}), '__dict__') else {},
                    "enhancement_summary": enhanced_results.get('enhancement_summary', {})
                }
                exported_content = json.dumps(arcadia_export, indent=2, ensure_ascii=False, default=str)
                filename = "arcadia_analysis.json"
                mime_type = "application/json"
                
            else:
                # Fallback to JSON for unsupported formats
                exported_content = export_requirements_to_json(export_results)
                filename = "requirements.json"
                mime_type = "application/json"
                if export_format not in ["JSON", "CSV", "Excel", "ARCADIA_JSON"]:
                    st.warning(f"Format {export_format} not fully supported, exporting as JSON")
            
            # Create download button - always available
            st.download_button(
                f"Download {export_format}",
                exported_content,
                filename,
                mime_type,
                key=f"download_requirements_{export_format}",
                help=f"Download requirements in {export_format} format"
            )
            
            # Show format info
            total_reqs = sum(len(reqs) for phase_reqs in export_results.get('requirements', {}).values() 
                           for reqs in phase_reqs.values() if isinstance(reqs, list))
            st.caption(f"Ready to export {total_reqs} requirements in {export_format} format")
            
        except Exception as e:
            st.error(f"Export preparation error: {str(e)}")
            logger.error(f"Requirements export preparation error: {str(e)}")
    
    with col2:
        st.info(f"Export format: {export_format}")
        
        # Show export stats
        if export_results:
            stats = export_results.get('statistics', {})
            if stats.get('total_requirements'):
                st.metric("Total Requirements", stats['total_requirements'])
            st.metric("Phases", len(export_results.get('requirements', {})))
            st.metric("Stakeholders", len(export_results.get('stakeholders', {})))

def display_requirements_list(requirements, req_type):
    if not requirements:
        st.info(f"No {req_type} requirements generated")
        return
    
    for req in requirements:
        # Determine priority class
        priority = req.get('priority', 'N/A').lower()
        if priority == 'must':
            priority_class = 'priority-high'
        elif priority == 'should':
            priority_class = 'priority-medium'
        else:
            priority_class = 'priority-low'
        
        # Get priority confidence and rationale
        priority_confidence = req.get('priority_confidence', 0.0)
        priority_rationale = req.get('rationale', 'No rationale provided')
        
        # Create expandable details for priority analysis
        priority_analysis = req.get('priority_analysis', {})
        analysis_summary = ""
        if priority_analysis:
            indicators = priority_analysis.get('indicators_found', [])
            if indicators:
                categories = set(ind['category'] for ind in indicators)
                analysis_summary = f"Based on: {', '.join(categories)}"
        
        st.markdown(f"""
        <div class="requirement-item">
            <div class="requirement-header">{req.get('id', 'N/A')}: {req.get('title', 'Untitled')}</div>
            <div class="requirement-description">{req.get('description', 'No description')}</div>
            <div class="requirement-meta">
                <span><strong>Priority:</strong> <span class="{priority_class}">{req.get('priority', 'N/A')}</span> 
                      (Confidence: {priority_confidence:.1f})</span>
                <span><strong>Verification:</strong> {req.get('verification_method', 'N/A')}</span>
            </div>
            <div class="requirement-analysis" style="font-size: 0.9em; color: #666; margin-top: 8px;">
                <strong>Priority Analysis:</strong> {analysis_summary}
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        # Add expandable details for priority rationale
        with st.expander(f"Priority Rationale for {req.get('id', 'N/A')}", expanded=False):
            st.write(priority_rationale)
            
            if priority_analysis:
                st.subheader("Detailed Analysis")
                col1, col2 = st.columns(2)
                
                with col1:
                    st.metric("Stakeholder Alignment", f"{priority_analysis.get('stakeholder_alignment', 0):.2f}")
                    st.metric("Regulatory Compliance", f"{priority_analysis.get('regulatory_compliance', 0):.2f}")
                    st.metric("Safety Criticality", f"{priority_analysis.get('safety_criticality', 0):.2f}")
                
                with col2:
                    st.metric("Component Specificity", f"{priority_analysis.get('component_specificity', 0):.2f}")
                    st.metric("Phase Relevance", f"{priority_analysis.get('phase_relevance', 0):.2f}")
                
                # Show indicators found
                indicators_found = priority_analysis.get('indicators_found', [])
                if indicators_found:
                    st.subheader("Criticality Indicators Found")
                    for indicator in indicators_found:
                        st.write(f"• **{indicator['keyword']}** ({indicator['category']}) - Weight: {indicator['weight']:.2f}")
        



def evaluation_tab(rag_system, eval_service):
    st.markdown("### System Evaluation")
    
    if 'generation_results' in st.session_state:
        results = st.session_state['generation_results']
        
        # Evaluation metrics
        st.markdown("#### Evaluation Metrics")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("Evaluate against CYDERCO"):
                with st.spinner("Evaluating..."):
                    evaluation = rag_system.evaluate_against_cyderco(results)
                    st.session_state['evaluation_results'] = evaluation
        
        with col2:
            if st.button("Quality Assessment"):
                with st.spinner("Assessing quality..."):
                    quality_assessment = eval_service.assess_requirement_quality(results)
                    st.session_state['quality_assessment'] = quality_assessment
        
        # Display evaluation results
        if 'evaluation_results' in st.session_state:
            evaluation = st.session_state['evaluation_results']
            
            st.markdown("#### CYDERCO Compatibility")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Coverage Score", f"{evaluation['coverage_score']:.1f}%")
            with col2:
                st.metric("Missing Requirements", len(evaluation['missing_requirements']))
            with col3:
                st.metric("Additional Requirements", len(evaluation['additional_requirements']))
            
            # Quality metrics
            if 'quality_assessment' in st.session_state:
                quality = st.session_state['quality_assessment']
                
                st.markdown("#### Quality Metrics")
                
                metrics_df = pd.DataFrame([
                    {"Metric": "Completeness", "Score": quality['completeness']},
                    {"Metric": "Consistency", "Score": quality['consistency']},
                    {"Metric": "Traceability", "Score": quality['traceability']},
                    {"Metric": "Testability", "Score": quality['testability']},
                    {"Metric": "Clarity", "Score": quality['clarity']}
                ])
                
                # Display metrics as a chart
                fig = px.bar(
                    metrics_df,
                    x="Metric",
                    y="Score",
                    title="Requirements Quality Assessment",
                    color="Score",
                    color_continuous_scale="RdYlGn"
                )
                fig.update_layout(showlegend=False)
                st.plotly_chart(fig)
                
                # Display detailed metrics
                for _, row in metrics_df.iterrows():
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.write(f"**{row['Metric']}**")
                    with col2:
                        score_color = "green" if row['Score'] > 0.8 else "orange" if row['Score'] > 0.6 else "red"
                        st.markdown(f"<span style='color: {score_color}'>{row['Score']:.1%}</span>", unsafe_allow_html=True)
    
    else:
        st.info("Generate requirements first to see evaluation results")



def chat_tab(rag_system):
    """Enhanced chat interface with document interaction"""
    st.markdown("### Chat with Documents")
    
    # Create layout
    col1, col2, col3 = st.columns([2, 6, 2])
    
    with col1:
        st.markdown("#### Chat Management")
        
        # New chat button - Legacy support for traditional mode
        if st.button("New Chat", type="primary", use_container_width=True):
            new_chat_id = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
            st.session_state.current_chat_id = new_chat_id
            
            # Initialize legacy chats if not exists
            if "chats" not in st.session_state:
                st.session_state.chats = {}
                
            st.session_state.chats[new_chat_id] = {
                "title": "New Chat",
                "messages": [],
                "created_at": datetime.now().isoformat(),
                "document_count": 0
            }
            save_chats(st.session_state.chats)
            st.rerun()
        
        # Chat history
        st.markdown("#### Chat History")
        
        # Initialize legacy chats if needed
        if "chats" not in st.session_state:
            st.session_state.chats = {}
            
        if st.session_state.chats:
            with st.container():
                for chat_id, chat_data in sorted(st.session_state.chats.items(), 
                                               key=lambda x: x[1].get('created_at', ''), reverse=True):
                    title = chat_data.get('title', 'Untitled Chat')
                    msg_count = len(chat_data.get('messages', []))
                    
                    # Create a more informative button label
                    button_label = f"{title[:25]}{'...' if len(title) > 25 else ''}"
                    if msg_count > 0:
                        button_label += f" ({msg_count} msgs)"
                    
                    if st.button(button_label, key=f"chat_{chat_id}", use_container_width=True):
                        st.session_state.current_chat_id = chat_id
                        st.rerun()
                        
                    # Delete chat option
                    if st.button("Delete", key=f"delete_{chat_id}", help="Delete this chat"):
                        del st.session_state.chats[chat_id]
                        if st.session_state.current_chat_id == chat_id:
                            st.session_state.current_chat_id = None
                        save_chats(st.session_state.chats)
                        st.rerun()
        else:
            st.info("No previous chats. Create a new chat to get started!")
        
        # Clear all chats
        if st.session_state.chats and st.button("Clear All Chats", type="secondary"):
            if st.checkbox("Confirm deletion of all chats"):
                st.session_state.chats = {}
                st.session_state.current_chat_id = None
                save_chats(st.session_state.chats)
                st.success("All chats cleared!")
                st.rerun()
    
    with col2:
        st.markdown("#### Document Management")
        
        # Document Upload Section
        with st.expander("Upload Documents", expanded=False):
            st.markdown("""
            **Supported formats:** PDF, DOCX, TXT, MD, XML, JSON, AIRD, Capella  
            **Purpose:** Add documents to the knowledge base for chat interaction
            """)
            
            uploaded_files = st.file_uploader(
                "Add documents to the knowledge base",
                accept_multiple_files=True,
                type=['pdf', 'docx', 'txt', 'md', 'xml', 'json', 'aird', 'capella'],
                help="Upload multiple documents to enhance the chat knowledge base"
            )
            
            if uploaded_files:
                st.info(f"{len(uploaded_files)} file(s) selected")
                
                if st.button("Process Documents", type="primary"):
                    with st.spinner("Processing documents..."):
                        logger.info(f"Processing {len(uploaded_files)} documents for chat...")
                        
                        temp_paths = []
                        processed_files = []
                        
                        try:
                            # Save uploaded files temporarily
                            for uploaded_file in uploaded_files:
                                temp_path = Path(f"temp_{uploaded_file.name}")
                                with open(temp_path, "wb") as f:
                                    f.write(uploaded_file.getbuffer())
                                temp_paths.append(str(temp_path))
                                processed_files.append(uploaded_file.name)
                                logger.info(f"   Saved: {uploaded_file.name} ({uploaded_file.size} bytes)")
                            
                            # Process documents through RAG system
                            results = rag_system.add_documents_to_vectorstore(temp_paths)
                            
                            # Clean up temporary files
                            for temp_path in temp_paths:
                                try:
                                    os.remove(temp_path)
                                except:
                                    pass
                            
                            # Display results
                            st.success(f"Processed {results.get('processed', 0)} files successfully!")
                            st.info(f"Added {results.get('chunks_added', 0)} text chunks to knowledge base")
                            
                            # Show errors outside the expander to avoid nesting
                            if results.get('errors'):
                                st.error("Processing Errors:")
                                for error in results['errors']:
                                    st.error(f"{error}")
                            
                            # Update current chat document count
                            if st.session_state.current_chat_id and st.session_state.current_chat_id in st.session_state.chats:
                                current_chat = st.session_state.chats[st.session_state.current_chat_id]
                                current_chat['document_count'] = current_chat.get('document_count', 0) + results.get('processed', 0)
                                save_chats(st.session_state.chats)
                            
                            logger.info(f"Document processing completed for chat interface")
                            
                        except Exception as e:
                            logger.error(f"Error processing documents for chat: {str(e)}")
                            st.error(f"Error processing documents: {str(e)}")
                            
                            # Clean up on error
                            for temp_path in temp_paths:
                                try:
                                    os.remove(temp_path)
                                except:
                                    pass
        
        # Chat Interface
        st.markdown("#### Conversation")
        
        if st.session_state.current_chat_id:
            current_chat = st.session_state.chats[st.session_state.current_chat_id]
            
            # Display chat title
            chat_title = st.text_input(
                "Chat Title", 
                value=current_chat.get('title', 'New Chat'),
                key="chat_title_input"
            )
            
            # Update title if changed
            if chat_title != current_chat.get('title', 'New Chat'):
                current_chat['title'] = chat_title
                save_chats(st.session_state.chats)
            
            # Chat messages container
            chat_container = st.container()
            
            with chat_container:
                # Display existing messages
                for i, message in enumerate(current_chat["messages"]):
                    if message["role"] == "user":
                        st.markdown(f"""
                        <div class="chat-message user-message">
                            <strong>You:</strong><br>
                            {message["content"]}
                        </div>
                        """, unsafe_allow_html=True)
                    else:  # assistant
                        st.markdown(f"""
                        <div class="chat-message assistant-message">
                            <strong>Assistant:</strong><br>
                            {message["content"]}
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # Show context sources if available
                        if "context" in message and message["context"]:
                            with st.container():
                                st.markdown(f"**Sources ({len(message['context'])} documents):**")
                                for j, doc in enumerate(message["context"]):
                                    st.markdown(f"""
                                    <div class="source-citation">
                                        <strong>Source {j+1}:</strong> {doc.get('metadata', {}).get('source', 'Unknown')}<br>
                                        <strong>Content Preview:</strong> {doc.get('content', '')[:300]}...
                                    </div>
                                    """, unsafe_allow_html=True)
            
            # Chat input
            user_prompt = st.chat_input("Ask about your documents, ARCADIA methodology, or MBSE concepts...")
            
            if user_prompt:
                logger.info(f"New chat message in session {st.session_state.current_chat_id}: {user_prompt[:100]}...")
                
                # Update chat title if it's the first message
                if not current_chat["messages"]:
                    current_chat["title"] = user_prompt[:50] + ("..." if len(user_prompt) > 50 else "")
                
                # Add user message
                current_chat["messages"].append({
                    "role": "user", 
                    "content": user_prompt,
                    "timestamp": datetime.now().isoformat()
                })
                
                # Generate response
                with st.spinner("Analyzing documents and generating response..."):
                    try:
                        # Get MBSE context from settings
                        mbse_context = st.session_state.get('mbse_context', 'None')
                        enhanced_prompt = get_enhanced_prompt(user_prompt, mbse_context)
                        
                        # Use RAG system to generate response
                        response_data = rag_system.query_documents(enhanced_prompt)
                        
                        if isinstance(response_data, dict):
                            response = response_data.get('answer', 'I apologize, but I could not generate a response.')
                            context_docs = response_data.get('sources', [])
                        else:
                            response = str(response_data)
                            context_docs = []
                        
                        logger.info(f"Generated response with {len(context_docs)} context sources")
                        
                        # Add assistant response
                        current_chat["messages"].append({
                            "role": "assistant",
                            "content": response,
                            "context": [{"content": doc.page_content, "metadata": doc.metadata} for doc in context_docs] if context_docs else [],
                            "timestamp": datetime.now().isoformat()
                        })
                        
                        # Save chats
                        save_chats(st.session_state.chats)
                        
                        # Rerun to show new messages
                        st.rerun()
                        
                    except Exception as e:
                        logger.error(f"Error generating chat response: {str(e)}")
                        error_message = f"I apologize, but I encountered an error: {str(e)}"
                        
                        current_chat["messages"].append({
                            "role": "assistant",
                            "content": error_message,
                            "timestamp": datetime.now().isoformat()
                        })
                        
                        save_chats(st.session_state.chats)
                        st.rerun()
        
        else:
            st.info("Create a new chat or select an existing one to start chatting with your documents!")
            
            # Quick start suggestions
            st.markdown("""
            **Quick Start Tips:**
            1. Click **"New Chat"** to create a conversation
            2. Upload documents to build your knowledge base
            3. Ask questions about your documents, ARCADIA methodology, or MBSE concepts
            4. Use the settings panel to configure MBSE context for specialized responses
            """)
    
    with col3:
        st.markdown("#### Chat Settings")
        
        # Model selection
        available_models = ["llama3:instruct", "gemma3:27b", "gemma3:12b", "gemma3:4b", "mistral:latest", "deepseek-r1:7b"]
        selected_model = st.selectbox(
            "Chat Model",
            available_models,
            index=0,
            key="chat_model",
            help="Select the language model for chat responses"
        )
        
        # Context settings
        context_length = st.slider(
            "Context Documents",
            min_value=1,
            max_value=10,
            value=5,
            key="context_length",
            help="Number of document chunks to use as context"
        )
        
        # MBSE context selection
        mbse_context = st.selectbox(
            "MBSE Context",
            ["None"] + list(MBSE_CONTEXT_PROMPTS.keys()),
            key="mbse_context",
            help="Apply specialized MBSE perspective to responses"
        )
        
        if mbse_context != "None":
            st.info(f"**Context Applied:** {MBSE_CONTEXT_PROMPTS[mbse_context]}")
        
        # Chat statistics
        if st.session_state.chats:
            st.markdown("#### Chat Statistics")
            
            total_chats = len(st.session_state.chats)
            total_messages = sum(len(chat.get('messages', [])) for chat in st.session_state.chats.values())
            
            st.metric("Total Chats", total_chats)
            st.metric("Total Messages", total_messages)
            
            if st.session_state.current_chat_id and st.session_state.current_chat_id in st.session_state.chats:
                current_chat = st.session_state.chats[st.session_state.current_chat_id]
                current_messages = len(current_chat.get('messages', []))
                st.metric("Current Chat Messages", current_messages)

def structured_arcadia_tab(rag_system):
    """Enhanced tab for structured ARCADIA analysis with visualizations"""
    st.markdown("### 🏗️ Structured ARCADIA Analysis")
    
    # Check if we have enhanced results
    enhanced_results = st.session_state.get('enhanced_results')
    
    if not enhanced_results:
        st.info("""
        **No structured analysis available yet.**
        
        To generate structured ARCADIA analysis:
        1. Go to the **Generate Requirements** tab
        2. Enable **"Structured ARCADIA Analysis"** in the sidebar
        3. Upload a document or paste text and generate requirements
        4. Return here to view detailed structured insights
        """)
        
        # Add example of what will be available
        st.markdown("#### What you'll see here:")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("""
            **🎭 Operational Analysis**
            - Stakeholder identification
            - Operational capabilities
            - Mission scenarios
            - Activity processes
            
            **🏗️ System Analysis**  
            - System boundary definition
            - System functions & interfaces
            - Capability realization
            - Functional chains
            """)
        with col2:
            st.markdown("""
            **🧩 Logical Architecture**
            - Logical components & hierarchies
            - Logical functions & behaviors
            - Logical interfaces & data flows
            - Logical scenarios & interactions
            
            **🔧 Physical Architecture**
            - Physical components & platforms
            - Implementation constraints
            - Deployment configurations
            - Performance characteristics
            """)
        
        st.markdown("""
        **🔗 Cross-Phase Analysis**
        - Bidirectional traceability between all phases
        - Gap identification across architecture levels
        - Architecture consistency validation
        - Quality metrics and recommendations
        """)
        return
    
    # Get structured analysis summary
    if hasattr(rag_system, 'get_structured_analysis_summary'):
        try:
            analysis_summary = rag_system.get_structured_analysis_summary(enhanced_results)
        except Exception as e:
            st.error(f"Error getting analysis summary: {str(e)}")
            analysis_summary = {}
    else:
        analysis_summary = {}
    
    # Main content tabs for structured analysis (All ARCADIA phases)
    struct_tab1, struct_tab2, struct_tab3, struct_tab4, struct_tab5, struct_tab6 = st.tabs([
        "📊 Analysis Overview",
        "🎭 Operational Analysis", 
        "🏗️ System Analysis",
        "🧩 Logical Architecture",
        "🔧 Physical Architecture",
        "🔗 Cross-Phase Insights"
    ])
    
    with struct_tab1:
        display_analysis_overview(enhanced_results, analysis_summary)
    
    with struct_tab2:
        display_operational_analysis(enhanced_results)
    
    with struct_tab3:
        display_system_analysis(enhanced_results)
    
    with struct_tab4:
        display_logical_analysis(enhanced_results)
    
    with struct_tab5:
        display_physical_analysis(enhanced_results)
    
    with struct_tab6:
        display_cross_phase_analysis(enhanced_results, analysis_summary)

def display_analysis_overview(enhanced_results, analysis_summary):
    """Display high-level overview of the structured analysis"""
    st.markdown("#### Analysis Overview")
    
    # Enhancement summary
    enhancement_summary = enhanced_results.get('enhancement_summary', {})
    
    # Key metrics - All ARCADIA phases
    col1, col2, col3, col4, col5, col6 = st.columns(6)
    
    with col1:
        st.metric(
            "Phases Analyzed", 
            len(enhancement_summary.get('phases_analyzed', [])),
            help="Number of ARCADIA phases successfully analyzed"
        )
    
    with col2:
        st.metric(
            "Total Actors", 
            enhancement_summary.get('total_actors_identified', 0),
            help="Actors identified in operational and system phases"
        )
    
    with col3:
        st.metric(
            "Total Capabilities", 
            enhancement_summary.get('total_capabilities_identified', 0),
            help="Capabilities identified in operational and system phases"
        )
    
    with col4:
        st.metric(
            "Total Components", 
            enhancement_summary.get('total_components_identified', 0),
            help="Components identified in logical and physical phases"
        )
    
    with col5:
        st.metric(
            "Total Functions", 
            enhancement_summary.get('total_functions_identified', 0),
            help="Functions identified across system, logical, and physical phases"
        )
    
    with col6:
        st.metric(
            "Cross-Phase Links", 
            enhancement_summary.get('cross_phase_links', 0),
            help="Traceability links between phases"
        )
    
    # Phase coverage visualization
    if 'extraction_statistics' in analysis_summary:
        st.markdown("#### Phase Coverage Analysis")
        
        extraction_stats = analysis_summary['extraction_statistics']
        
        # Create coverage data for visualization (all ARCADIA phases)
        coverage_data = []
        for phase, stats in extraction_stats.items():
            total_elements = sum(stats.values())
            phase_data = {
                'Phase': phase.title(),
                'Elements': total_elements,
                'Actors': stats.get('actors', 0),
                'Capabilities': stats.get('capabilities', 0),
                'Functions': stats.get('functions', 0),
                'Scenarios': stats.get('scenarios', 0)
            }
            
            # Add phase-specific metrics
            if phase == 'logical':
                phase_data.update({
                    'Components': stats.get('components', 0),
                    'Interfaces': stats.get('interfaces', 0)
                })
            elif phase == 'physical':
                phase_data.update({
                    'Components': stats.get('components', 0),
                    'Constraints': stats.get('implementation_constraints', 0)
                })
            
            coverage_data.append(phase_data)
        
        if coverage_data:
            coverage_df = pd.DataFrame(coverage_data)
            
            # Create stacked bar chart
            fig = px.bar(
                coverage_df, 
                x='Phase', 
                y=['Actors', 'Capabilities', 'Functions', 'Scenarios'],
                title="Elements Extracted by Phase",
                color_discrete_sequence=['#667eea', '#2c5aa0', '#5a67d8', '#4c51bf']
            )
            fig.update_layout(height=400)
            st.plotly_chart(fig, use_container_width=True)
    
    # Quality scores
    if 'quality_scores' in analysis_summary:
        st.markdown("#### Quality Assessment")
        
        quality_scores = analysis_summary['quality_scores']
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Quality metrics as gauge charts
            for metric_name, score_info in quality_scores.items():
                percentage = score_info['percentage']
                
                # Color based on score
                if percentage >= 80:
                    color = "green"
                elif percentage >= 60:
                    color = "orange"
                else:
                    color = "red"
                
                st.metric(
                    metric_name, 
                    f"{percentage:.1f}%",
                    help=f"Score: {score_info['score']:.2f}/{score_info['max_score']}"
                )
        
        with col2:
            # Quality score visualization
            quality_data = []
            for metric_name, score_info in quality_scores.items():
                quality_data.append({
                    'Metric': metric_name,
                    'Score': score_info['percentage']
                })
            
            if quality_data:
                quality_df = pd.DataFrame(quality_data)
                fig = px.bar(
                    quality_df,
                    x='Score',
                    y='Metric',
                    orientation='h',
                    title="Quality Metrics",
                    color='Score',
                    color_continuous_scale='RdYlGn',
                    range_color=[0, 100]
                )
                fig.update_layout(height=300)
                st.plotly_chart(fig, use_container_width=True)
    
    # Recommendations
    if 'recommendations' in analysis_summary:
        st.markdown("#### Recommendations")
        recommendations = analysis_summary['recommendations']
        
        if recommendations:
            for i, rec in enumerate(recommendations, 1):
                st.info(f"**{i}.** {rec}")
        else:
            st.success("✅ Analysis completed successfully with good quality scores!")

def display_operational_analysis(enhanced_results):
    """Display detailed operational analysis results"""
    st.markdown("#### 🎭 Operational Analysis Details")
    
    structured_analysis = enhanced_results.get('structured_analysis')
    if not structured_analysis or not structured_analysis.operational_analysis:
        st.warning("No operational analysis data available")
        return
    
    op_analysis = structured_analysis.operational_analysis
    
    # Operational Actors Section
    with st.expander("👥 Operational Actors", expanded=True):
        if op_analysis.actors:
            for i, actor in enumerate(op_analysis.actors):
                with st.container():
                    st.markdown(f"""
                    <div class="requirement-item">
                        <div class="requirement-header">{actor.name} ({actor.id})</div>
                        <div class="requirement-description">
                            <strong>Role:</strong> {actor.role_definition}<br>
                            <strong>Description:</strong> {actor.description}
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if actor.responsibilities:
                        st.markdown("**Responsibilities:**")
                        for resp in actor.responsibilities:
                            st.markdown(f"• {resp}")
                    
                    if actor.capabilities:
                        st.markdown("**Capabilities:**")
                        for cap in actor.capabilities:
                            st.markdown(f"• {cap}")
                    
                    if i < len(op_analysis.actors) - 1:
                        st.markdown("---")
        else:
            st.info("No operational actors identified")
    
    # Operational Capabilities Section
    with st.expander("🎯 Operational Capabilities", expanded=True):
        if op_analysis.capabilities:
            for i, capability in enumerate(op_analysis.capabilities):
                with st.container():
                    st.markdown(f"""
                    <div class="requirement-item">
                        <div class="requirement-header">{capability.name} ({capability.id})</div>
                        <div class="requirement-description">
                            <strong>Mission:</strong> {capability.mission_statement}<br>
                            <strong>Description:</strong> {capability.description}
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if capability.involved_actors:
                        st.markdown(f"**Involved Actors:** {', '.join(capability.involved_actors)}")
                    
                    if capability.performance_constraints:
                        st.markdown("**Performance Constraints:**")
                        for constraint in capability.performance_constraints:
                            st.markdown(f"• {constraint}")
                    
                    if i < len(op_analysis.capabilities) - 1:
                        st.markdown("---")
        else:
            st.info("No operational capabilities identified")
    
    # Operational Scenarios Section
    with st.expander("📋 Operational Scenarios", expanded=False):
        if op_analysis.scenarios:
            for scenario in op_analysis.scenarios:
                st.markdown(f"**{scenario.name} ({scenario.id})**")
                st.write(f"Type: {scenario.scenario_type}")
                st.write(f"Description: {scenario.description}")
                st.markdown("---")
        else:
            st.info("No operational scenarios identified")

def display_system_analysis(enhanced_results):
    """Display detailed system analysis results"""
    st.markdown("#### 🏗️ System Analysis Details")
    
    structured_analysis = enhanced_results.get('structured_analysis')
    if not structured_analysis or not structured_analysis.system_analysis:
        st.warning("No system analysis data available")
        return
    
    sys_analysis = structured_analysis.system_analysis
    
    # System Boundary Section
    with st.expander("🔲 System Boundary", expanded=True):
        boundary = sys_analysis.system_boundary
        st.markdown(f"""
        <div class="info-card">
            <h4>Scope Definition</h4>
            <p>{boundary.scope_definition}</p>
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        with col1:
            if boundary.included_elements:
                st.markdown("**Included Elements:**")
                for element in boundary.included_elements:
                    st.markdown(f"✅ {element}")
        
        with col2:
            if boundary.excluded_elements:
                st.markdown("**Excluded Elements:**")
                for element in boundary.excluded_elements:
                    st.markdown(f"❌ {element}")
    
    # System Functions Section
    with st.expander("⚙️ System Functions", expanded=True):
        if sys_analysis.functions:
            for i, function in enumerate(sys_analysis.functions):
                with st.container():
                    st.markdown(f"""
                    <div class="requirement-item">
                        <div class="requirement-header">{function.name} ({function.id})</div>
                        <div class="requirement-description">
                            <strong>Type:</strong> {function.function_type}<br>
                            <strong>Description:</strong> {function.description}
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if function.allocated_actors:
                        st.markdown(f"**Allocated Actors:** {', '.join(function.allocated_actors)}")
                    
                    if function.performance_requirements:
                        st.markdown("**Performance Requirements:**")
                        for req in function.performance_requirements:
                            st.markdown(f"• {req}")
                    
                    if i < len(sys_analysis.functions) - 1:
                        st.markdown("---")
        else:
            st.info("No system functions identified")
    
    # System Capabilities Section
    with st.expander("🎯 System Capabilities", expanded=False):
        if sys_analysis.capabilities:
            for capability in sys_analysis.capabilities:
                st.markdown(f"**{capability.name} ({capability.id})**")
                st.write(f"Description: {capability.description}")
                if capability.implementing_functions:
                    st.write(f"Implementing Functions: {', '.join(capability.implementing_functions)}")
                st.markdown("---")
        else:
            st.info("No system capabilities identified")

def display_logical_analysis(enhanced_results):
    """Display detailed logical architecture results"""
    st.markdown("#### 🧩 Logical Architecture Details")
    
    structured_analysis = enhanced_results.get('structured_analysis')
    if not structured_analysis or not structured_analysis.logical_architecture:
        st.warning("No logical architecture data available. Please ensure logical architecture extraction is enabled.")
        return
    
    logical_arch = structured_analysis.logical_architecture
    
    # Logical Components Section
    with st.expander("🧩 Logical Components", expanded=True):
        if logical_arch.components:
            for i, component in enumerate(logical_arch.components):
                with st.container():
                    st.markdown(f"""
                    <div class="requirement-item">
                        <div class="requirement-header">{component.name} ({component.id})</div>
                        <div class="requirement-description">
                            <strong>Type:</strong> {component.component_type}<br>
                            <strong>Description:</strong> {component.description}
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if component.responsibilities:
                        st.markdown("**Responsibilities:**")
                        for resp in component.responsibilities:
                            st.markdown(f"• {resp}")
                    
                    if component.sub_components:
                        st.markdown(f"**Sub-components:** {', '.join(component.sub_components)}")
                    
                    if i < len(logical_arch.components) - 1:
                        st.markdown("---")
        else:
            st.info("No logical components identified")
    
    # Logical Functions Section
    with st.expander("⚙️ Logical Functions", expanded=True):
        if logical_arch.functions:
            for i, function in enumerate(logical_arch.functions):
                with st.container():
                    # Format behavioral models for display
                    behavioral_display = "; ".join([model.get('spec', str(model)) for model in function.behavioral_models]) if function.behavioral_models else "Not specified"
                    
                    st.markdown(f"""
                    <div class="requirement-item">
                        <div class="requirement-header">{function.name} ({function.id})</div>
                        <div class="requirement-description">
                            <strong>Behavior:</strong> {behavioral_display}<br>
                            <strong>Description:</strong> {function.description}
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if function.allocated_components:
                        st.markdown(f"**Allocated Components:** {', '.join(function.allocated_components)}")
                    
                    if function.input_interfaces:
                        st.markdown("**Input Interfaces:**")
                        for interface in function.input_interfaces:
                            st.markdown(f"• {interface}")
                    
                    if function.output_interfaces:
                        st.markdown("**Output Interfaces:**")
                        for interface in function.output_interfaces:
                            st.markdown(f"• {interface}")
                    
                    if i < len(logical_arch.functions) - 1:
                        st.markdown("---")
        else:
            st.info("No logical functions identified")
    
    # Logical Interfaces Section
    with st.expander("🔌 Logical Interfaces", expanded=False):
        if logical_arch.interfaces:
            for interface in logical_arch.interfaces:
                st.markdown(f"**{interface.name} ({interface.id})**")
                st.write(f"Protocol: {interface.protocol}")
                st.write(f"Provider: {interface.provider} → Consumer: {interface.consumer}")
                if interface.data_elements:
                    st.write(f"Data Elements: {', '.join(interface.data_elements)}")
                st.markdown("---")
        else:
            st.info("No logical interfaces identified")
    
    # Logical Scenarios Section
    with st.expander("📋 Logical Scenarios", expanded=False):
        if logical_arch.scenarios:
            for scenario in logical_arch.scenarios:
                st.markdown(f"**{scenario.name} ({scenario.id})**")
                st.write(f"Context: {scenario.scenario_context}")
                st.write(f"Description: {scenario.description}")
                if scenario.component_interactions:
                    st.markdown("**Component Interactions:**")
                    for interaction in scenario.component_interactions:
                        st.markdown(f"• {interaction}")
                st.markdown("---")
        else:
            st.info("No logical scenarios identified")

def display_physical_analysis(enhanced_results):
    """Display detailed physical architecture results"""
    st.markdown("#### 🔧 Physical Architecture Details")
    
    structured_analysis = enhanced_results.get('structured_analysis')
    if not structured_analysis or not structured_analysis.physical_architecture:
        st.warning("No physical architecture data available. Please ensure physical architecture extraction is enabled.")
        return
    
    physical_arch = structured_analysis.physical_architecture
    
    # Physical Components Section
    with st.expander("🔧 Physical Components", expanded=True):
        if physical_arch.components:
            for i, component in enumerate(physical_arch.components):
                with st.container():
                    st.markdown(f"""
                    <div class="requirement-item">
                        <div class="requirement-header">{component.name} ({component.id})</div>
                        <div class="requirement-description">
                            <strong>Type:</strong> {component.component_type}<br>
                            <strong>Technology Platform:</strong> {component.technology_platform}<br>
                            <strong>Description:</strong> {component.description}
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if component.resource_requirements:
                        st.markdown("**Resource Requirements:**")
                        for req in component.resource_requirements:
                            st.markdown(f"• {req}")
                    
                    if component.deployment_location:
                        st.markdown(f"**Deployment Location:** {component.deployment_location}")
                    
                    if i < len(physical_arch.components) - 1:
                        st.markdown("---")
        else:
            st.info("No physical components identified")
    
    # Implementation Constraints Section
    with st.expander("⚠️ Implementation Constraints", expanded=True):
        if physical_arch.constraints:
            # Group constraints by type
            constraint_types = {}
            for constraint in physical_arch.constraints:
                constraint_type = constraint.constraint_type
                if constraint_type not in constraint_types:
                    constraint_types[constraint_type] = []
                constraint_types[constraint_type].append(constraint)
            
            for constraint_type, constraints in constraint_types.items():
                st.markdown(f"**{constraint_type.title()} Constraints:**")
                for constraint in constraints:
                    st.markdown(f"• **{constraint.description}**")
                    if hasattr(constraint, 'rationale') and constraint.rationale:
                        st.markdown(f"  💡 *{constraint.rationale}*")
                st.markdown("---")
        else:
            st.info("No implementation constraints identified")
    
    # Physical Functions Section
    with st.expander("⚙️ Physical Functions", expanded=False):
        if physical_arch.functions:
            for function in physical_arch.functions:
                st.markdown(f"**{function.name} ({function.id})**")
                st.write(f"Description: {function.description}")
                if function.resource_requirements:
                    st.write(f"Resource Requirements: {', '.join(function.resource_requirements)}")
                if function.timing_constraints:
                    st.write(f"Timing Constraints: {', '.join(function.timing_constraints)}")
                st.markdown("---")
        else:
            st.info("No physical functions identified")
    
    # Physical Scenarios Section
    with st.expander("📋 Physical Scenarios", expanded=False):
        if physical_arch.scenarios:
            for scenario in physical_arch.scenarios:
                st.markdown(f"**{scenario.name} ({scenario.id})**")
                st.write(f"Type: {scenario.scenario_type}")
                st.write(f"Description: {scenario.description}")
                if scenario.involved_components:
                    st.write(f"Involved Components: {', '.join(scenario.involved_components)}")
                st.markdown("---")
        else:
            st.info("No physical scenarios identified")

def display_cross_phase_analysis(enhanced_results, analysis_summary):
    """Display cross-phase analysis insights"""
    st.markdown("#### 🔗 Cross-Phase Analysis Insights")
    
    structured_analysis = enhanced_results.get('structured_analysis')
    if not structured_analysis or not structured_analysis.cross_phase_analysis:
        st.warning("No cross-phase analysis data available")
        return
    
    cross_phase = structured_analysis.cross_phase_analysis
    
    # Traceability Links Section
    with st.expander("🔗 Traceability Links", expanded=True):
        if cross_phase.traceability_links:
            st.markdown(f"**Total Links:** {len(cross_phase.traceability_links)}")
            
            # Create traceability visualization
            trace_data = []
            for link in cross_phase.traceability_links:
                trace_data.append({
                    'Source': link.source_element,
                    'Target': link.target_element,
                    'Relationship': link.relationship_type,
                    'Confidence': link.confidence_score,
                    'Source Phase': link.source_phase.value,
                    'Target Phase': link.target_phase.value
                })
            
            if trace_data:
                trace_df = pd.DataFrame(trace_data)
                
                # Display as table
                st.dataframe(trace_df, use_container_width=True)
                
                # Create sankey diagram would be nice here, but simpler for now
                st.markdown("**Traceability Summary:**")
                for link in cross_phase.traceability_links[:5]:  # Show first 5
                    confidence_color = "🟢" if link.confidence_score > 0.7 else "🟡" if link.confidence_score > 0.5 else "🔴"
                    st.markdown(f"{confidence_color} {link.source_element} → {link.target_element} ({link.relationship_type})")
        else:
            st.info("No traceability links identified")
    
    # Gap Analysis Section
    with st.expander("🔍 Gap Analysis", expanded=True):
        if cross_phase.gap_analysis:
            st.markdown(f"**Gaps Identified:** {len(cross_phase.gap_analysis)}")
            
            # Group by severity
            gaps_by_severity = {}
            for gap in cross_phase.gap_analysis:
                severity = gap.severity
                if severity not in gaps_by_severity:
                    gaps_by_severity[severity] = []
                gaps_by_severity[severity].append(gap)
            
            # Display by severity
            severity_colors = {
                'critical': '🔴',
                'major': '🟠', 
                'medium': '🟡',
                'minor': '🟢'
            }
            
            for severity in ['critical', 'major', 'medium', 'minor']:
                if severity in gaps_by_severity:
                    st.markdown(f"**{severity_colors[severity]} {severity.title()} ({len(gaps_by_severity[severity])})**")
                    for gap in gaps_by_severity[severity]:
                        st.markdown(f"• **{gap.gap_type.upper()}**: {gap.description}")
                        if gap.recommendations:
                            for rec in gap.recommendations:
                                st.markdown(f"  💡 {rec}")
        else:
            st.success("✅ No gaps identified - analysis appears complete!")
    
    # Coverage Matrix
    if cross_phase.coverage_matrix:
        with st.expander("📊 Coverage Matrix", expanded=False):
            st.markdown("**Phase-to-Phase Coverage:**")
            
            coverage_data = []
            for source_target, coverage_info in cross_phase.coverage_matrix.items():
                for coverage_type, score in coverage_info.items():
                    coverage_data.append({
                        'Relationship': source_target.replace('_', ' → ').title(),
                        'Coverage Type': coverage_type.replace('_', ' ').title(),
                        'Coverage %': score * 100
                    })
            
            if coverage_data:
                coverage_df = pd.DataFrame(coverage_data)
                fig = px.bar(
                    coverage_df,
                    x='Relationship',
                    y='Coverage %',
                    color='Coverage Type',
                    title="Cross-Phase Coverage Analysis",
                    color_discrete_sequence=['#667eea', '#2c5aa0']
                )
                fig.update_layout(height=400)
                st.plotly_chart(fig, use_container_width=True)

def project_management_tab(rag_system, project_manager):
    """Comprehensive project management interface"""
    st.markdown("### 🗂️ Project Management")
    
    if not hasattr(rag_system, 'get_current_project'):
        st.error("❌ Project management not available with this RAG system")
        return
    
    current_project = rag_system.get_current_project()
    
    if not current_project:
        st.warning("⚠️ No project selected. Please create or select a project in the sidebar.")
        
        # Show available projects
        projects = rag_system.get_all_projects()
        if projects:
            st.markdown("#### Available Projects:")
            for project in projects[:5]:  # Show first 5
                col1, col2, col3 = st.columns([3, 1, 1])
                with col1:
                    st.write(f"**{project.name}**")
                    if project.description:
                        st.caption(project.description)
                with col2:
                    st.metric("Docs", project.documents_count)
                with col3:
                    st.metric("Reqs", project.requirements_count)
                
                st.markdown("---")
        
        # Quick project creation
        with st.expander("➕ Quick Create Project"):
            with st.form("quick_project"):
                name = st.text_input("Project Name", placeholder="My MBSE Project")
                description = st.text_area("Description", placeholder="Project description...")
                
                if st.form_submit_button("Create Project", type="primary"):
                    if name.strip():
                        try:
                            project_id = rag_system.create_project(name.strip(), description.strip())
                            st.success(f"✅ Project created: {name}")
                            st.rerun()
                        except Exception as e:
                            st.error(f"❌ Error creating project: {str(e)}")
                    else:
                        st.error("Project name is required")
        return
    
    # Display current project dashboard
    st.markdown(f"## 📋 {current_project.name}")
    if current_project.description:
        st.info(current_project.description)
    
    # Project metadata
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("📄 Documents", current_project.documents_count)
    with col2:
        st.metric("📝 Requirements", current_project.requirements_count)
    with col3:
        st.metric("📅 Created", current_project.created_at.strftime("%d/%m/%Y"))
    with col4:
        st.metric("🔄 Updated", current_project.updated_at.strftime("%d/%m/%Y"))
    
    # Project tabs
    project_tab1, project_tab2, project_tab3, project_tab4, project_tab5 = st.tabs([
        "📄 Documents", 
        "📝 Requirements", 
        "🔍 Search & Analysis",
        "📊 Project Statistics",
        "⚙️ Project Settings"
    ])
    
    with project_tab1:
        documents_project_tab(rag_system, current_project)
    
    with project_tab2:
        requirements_project_tab(rag_system, current_project)
    
    with project_tab3:
        search_analysis_project_tab(rag_system, current_project)
    
    with project_tab4:
        statistics_project_tab(rag_system, current_project)
    
    with project_tab5:
        settings_project_tab(rag_system, current_project)

def documents_project_tab(rag_system, current_project):
    """Document management for the current project"""
    st.markdown("#### 📄 Project Documents")
    
    # Document upload section with duplicate detection
    with st.expander("📤 Upload Documents", expanded=False):
        st.markdown("**Upload new documents to this project**")
        
        uploaded_files = st.file_uploader(
            "Select files",
            accept_multiple_files=True,
            type=['pdf', 'docx', 'txt', 'md', 'xml', 'json', 'aird'],
            help="Upload multiple documents for processing."
        )
        
        if uploaded_files:
            st.info(f"📄 {len(uploaded_files)} file(s) selected for upload")
            
            if st.button("🚀 Process Documents", type="primary"):
                process_documents_with_duplicate_detection(rag_system, current_project, uploaded_files)
    
    # List existing documents
    st.markdown("#### 📚 Existing Documents")
    
    try:
        documents = rag_system.persistence_service.get_project_documents(current_project.id)
        
        if documents:
            for doc in documents:
                with st.container():
                    col1, col2, col3, col4 = st.columns([3, 2, 2, 1])
                    
                    with col1:
                        st.write(f"**{doc.filename}**")
                        st.caption(f"ID: `{doc.id}`")
                    
                    with col2:
                        st.write(f"Size: {doc.file_size / 1024:.1f} KB")
                        st.write(f"Chunks: {doc.chunks_count}")
                    
                    with col3:
                        st.write(f"Model: {doc.embedding_model}")
                        st.write(f"Status: {doc.processing_status}")
                    
                    with col4:
                        status_color = "🟢" if doc.processing_status == "completed" else "🟡"
                        st.write(status_color)
                    
                    st.markdown("---")
        else:
            st.info("📭 No documents uploaded yet. Upload some documents to get started!")
            
    except Exception as e:
        st.error(f"❌ Error loading documents: {str(e)}")

def requirements_project_tab(rag_system, current_project):
    """Requirements management for the current project"""
    st.markdown("#### 📝 Project Requirements")
    
    try:
        requirements_data = rag_system.persistence_service.get_project_requirements(current_project.id)
        requirements = requirements_data.get("requirements", {})
        
        if requirements:
            # Requirements summary
            total_reqs = sum(len(reqs) for phase_reqs in requirements.values() for reqs in phase_reqs.values())
            st.metric("Total Requirements", total_reqs)
            
            # Display by phase
            for phase, phase_reqs in requirements.items():
                phase_info = arcadia_config.ARCADIA_PHASES.get(phase, {})
                st.markdown(f"### {phase_info.get('name', phase.title())} Phase")
                
                for req_type, reqs in phase_reqs.items():
                    if reqs:
                        st.markdown(f"#### {req_type.title()} Requirements ({len(reqs)})")
                        
                        for req in reqs:
                            with st.expander(f"{req.get('id', 'N/A')}: {req.get('title', 'Untitled')}"):
                                st.write(f"**Description:** {req.get('description', 'No description')}")
                                st.write(f"**Priority:** {req.get('priority', 'N/A')}")
                                st.write(f"**Verification:** {req.get('verification_method', 'N/A')}")
                                if req.get('rationale'):
                                    st.write(f"**Rationale:** {req.get('rationale')}")
                                st.write(f"**Created:** {req.get('created_at', 'N/A')}")
        else:
            st.info("📋 No requirements generated yet. Go to 'Generate Requirements' tab to create some!")
            
    except Exception as e:
        st.error(f"❌ Error loading requirements: {str(e)}")

def search_analysis_project_tab(rag_system, current_project):
    """Search and analysis within the project"""
    st.markdown("#### 🔍 Search & Analysis")
    
    # Search in project documents
    search_query = st.text_input("🔍 Search in project documents", placeholder="Enter your search query...")
    
    if search_query:
        with st.spinner("Searching..."):
            try:
                # Use the RAG system to search within the project
                results = rag_system.query_documents(search_query)
                
                if isinstance(results, dict) and results.get('sources'):
                    st.markdown(f"#### Search Results for: '{search_query}'")
                    st.write(f"**Answer:** {results.get('answer', 'No answer generated')}")
                    
                    st.markdown("**Sources:**")
                    for i, source in enumerate(results['sources'], 1):
                        with st.expander(f"Source {i}: {source.metadata.get('source', 'Unknown')}"):
                            st.write(source.page_content)
                            st.json(source.metadata)
                else:
                    st.info("No relevant documents found for your query.")
                    
            except Exception as e:
                st.error(f"❌ Search error: {str(e)}")
    
    # ARCADIA analysis for project
    st.markdown("#### 🏗️ ARCADIA Analysis")
    
    try:
        analyses = rag_system.persistence_service.get_project_arcadia_analyses(current_project.id)
        
        if analyses:
            st.success(f"✅ {len(analyses)} ARCADIA analysis(es) available")
            
            for analysis in analyses:
                with st.expander(f"{analysis['phase_type'].title()} Analysis - {analysis['created_at']}"):
                    st.json(analysis['analysis_data'])
        else:
            st.info("🏗️ No ARCADIA analyses yet. Generate structured analysis to see results here.")
            
    except Exception as e:
        st.error(f"❌ Error loading ARCADIA analyses: {str(e)}")

def statistics_project_tab(rag_system, current_project):
    """Project statistics and metrics"""
    st.markdown("#### 📊 Project Statistics")
    
    try:
        # Get project-specific statistics
        documents = rag_system.persistence_service.get_project_documents(current_project.id)
        requirements_data = rag_system.persistence_service.get_project_requirements(current_project.id)
        stakeholders = rag_system.persistence_service.get_project_stakeholders(current_project.id)
        sessions = rag_system.persistence_service.get_project_sessions(current_project.id, limit=10)
        
        # Statistics cards
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            total_size = sum(doc.file_size for doc in documents) / (1024 * 1024)  # MB
            st.metric("📄 Total Size", f"{total_size:.1f} MB")
        
        with col2:
            total_chunks = sum(doc.chunks_count for doc in documents)
            st.metric("🧩 Total Chunks", total_chunks)
        
        with col3:
            st.metric("👥 Stakeholders", len(stakeholders))
        
        with col4:
            st.metric("📝 Sessions", len(sessions))
        
        # Requirements by phase chart
        requirements = requirements_data.get("requirements", {})
        if requirements:
            st.markdown("#### Requirements by Phase")
            
            phase_data = []
            for phase, phase_reqs in requirements.items():
                total_phase_reqs = sum(len(reqs) for reqs in phase_reqs.values())
                phase_data.append({
                    "Phase": phase.title(),
                    "Requirements": total_phase_reqs
                })
            
            if phase_data:
                df = pd.DataFrame(phase_data)
                fig = px.bar(df, x="Phase", y="Requirements", title="Requirements Distribution by ARCADIA Phase")
                st.plotly_chart(fig, use_container_width=True)
        
        # Project activity timeline
        if sessions:
            st.markdown("#### Recent Project Activity")
            for session in sessions[:5]:  # Show last 5 activities
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.write(f"**{session['action_type']}**: {session['action_description']}")
                with col2:
                    st.caption(session['created_at'])
                
    except Exception as e:
        st.error(f"❌ Error loading statistics: {str(e)}")

def settings_project_tab(rag_system, current_project):
    """Project settings and management"""
    st.markdown("#### ⚙️ Project Settings")
    
    # Project editing
    with st.expander("✏️ Edit Project Information"):
        with st.form("edit_project"):
            new_name = st.text_input("Project Name", value=current_project.name)
            new_description = st.text_area("Description", value=current_project.description)
            new_proposal_text = st.text_area("Proposal Text", value=current_project.proposal_text, height=150)
            
            if st.form_submit_button("💾 Update Project"):
                try:
                    success = rag_system.persistence_service.update_project(
                        current_project.id, 
                        name=new_name if new_name != current_project.name else None,
                        description=new_description if new_description != current_project.description else None,
                        proposal_text=new_proposal_text if new_proposal_text != current_project.proposal_text else None
                    )
                    
                    if success:
                        st.success("✅ Project updated successfully!")
                        rag_system.persistence_service.log_project_session(
                            current_project.id, 
                            "project_update", 
                            "Project information updated"
                        )
                        st.rerun()
                    else:
                        st.error("❌ Failed to update project")
                        
                except Exception as e:
                    st.error(f"❌ Update error: {str(e)}")
    
    # Export project data
    with st.expander("📤 Export Project Data"):
        export_col1, export_col2 = st.columns(2)
        
        with export_col1:
            if st.button("📄 Export Requirements"):
                try:
                    requirements_data = rag_system.persistence_service.get_project_requirements(current_project.id)
                    st.download_button(
                        "Download Requirements JSON",
                        json.dumps(requirements_data, indent=2),
                        f"{current_project.name}_requirements.json",
                        "application/json"
                    )
                except Exception as e:
                    st.error(f"❌ Export error: {str(e)}")
        
        with export_col2:
            if st.button("🏗️ Export ARCADIA Analyses"):
                try:
                    analyses = rag_system.persistence_service.get_project_arcadia_analyses(current_project.id)
                    st.download_button(
                        "Download ARCADIA Analyses JSON",
                        json.dumps(analyses, indent=2),
                        f"{current_project.name}_arcadia_analyses.json",
                        "application/json"
                    )
                except Exception as e:
                    st.error(f"❌ Export error: {str(e)}")
    
    # Danger zone
    with st.expander("⚠️ Danger Zone", expanded=False):
        st.markdown("**Delete Project**")
        st.warning("⚠️ This action cannot be undone. All project data will be permanently deleted.")
        
        confirm_delete = st.checkbox("I understand that this action is irreversible")
        delete_confirmation = st.text_input("Type the project name to confirm deletion", placeholder=current_project.name)
        
        if confirm_delete and delete_confirmation == current_project.name:
            if st.button("🗑️ DELETE PROJECT", type="secondary"):
                try:
                    success = rag_system.persistence_service.delete_project(current_project.id)
                    if success:
                        st.success("✅ Project deleted successfully!")
                        st.info("🔄 Please refresh the page to continue.")
                    else:
                        st.error("❌ Failed to delete project")
                except Exception as e:
                    st.error(f"❌ Deletion error: {str(e)}")

def extract_file_content(file_path: str, filename: str) -> str:
    """Extract text content from various file types"""
    try:
        # Determine file type by extension
        file_ext = filename.lower().split('.')[-1]
        
        if file_ext in ['txt', 'md']:
            # Plain text files
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
                
        elif file_ext == 'pdf':
            # PDF files
            try:
                import PyPDF2
                import io
                
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in pdf_reader.pages:
                        try:
                            text += page.extract_text() + "\n"
                        except:
                            continue
                    return text
            except ImportError:
                logger.error("PyPDF2 not available for PDF processing")
                return ""
            except Exception as e:
                logger.error(f"Error extracting PDF content: {str(e)}")
                return ""
                
        elif file_ext == 'docx':
            # DOCX files
            try:
                from docx import Document
                
                doc = Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            except ImportError:
                logger.error("python-docx not available for DOCX processing")
                return ""
            except Exception as e:
                logger.error(f"Error extracting DOCX content: {str(e)}")
                return ""
                
        elif file_ext in ['json', 'xml']:
            # JSON/XML files
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
                
        else:
            # Try to read as plain text for other formats
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                # If UTF-8 fails, try with error handling
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
                    
    except Exception as e:
        logger.error(f"Error extracting content from {filename}: {str(e)}")
        return ""

def process_documents_with_duplicate_detection(rag_system, current_project, uploaded_files):
    """Process uploaded documents with intelligent duplicate detection"""
    
    # SIMPLIFIED: Process all files directly without duplicate detection
    st.info("Processing all files directly in project")
    
    # Simple processing approach
    temp_files = []
    all_extracted_content = []
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    try:
        status_text.text("Extracting content and processing files...")
        
        for i, uploaded_file in enumerate(uploaded_files):
            progress_bar.progress((i + 1) / len(uploaded_files))
            
            # Save temporarily
            temp_path = f"temp_{uploaded_file.name}"
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            temp_files.append(temp_path)
            
            # Extract content from file
            file_content = extract_file_content(temp_path, uploaded_file.name)
            if file_content:
                all_extracted_content.append({
                    'filename': uploaded_file.name,
                    'content': file_content,
                    'size': len(file_content)
                })
        
        # Store extracted content in session state for requirements generation
        if all_extracted_content:
            combined_content = "\n\n=== DOCUMENT SEPARATOR ===\n\n".join([
                f"=== {doc['filename']} ===\n{doc['content']}" 
                for doc in all_extracted_content
            ])
            st.session_state['extracted_document_content'] = combined_content
            st.session_state['extracted_files_info'] = all_extracted_content
        
        # Show processing summary
        st.success(f"{len(uploaded_files)} file(s) ready for processing")
        
        # Process all files directly
        status_text.text("Processing documents...")
        
        if hasattr(rag_system, 'add_documents_to_project'):
            results = rag_system.add_documents_to_project(temp_files, current_project.id)
        else:
            # Fallback for systems without project support
            results = rag_system.add_documents_to_vectorstore(temp_files)
        
        # Safely extract numeric results, handling cases where values might be lists
        processed_val = results.get('processed', results.get('processed_files', 0))
        chunks_val = results.get('new_chunks', results.get('chunks_added', 0))
        
        # Convert to integers if they're lists (take length) or other types
        if isinstance(processed_val, list):
            total_processed = len(processed_val)
        elif isinstance(processed_val, (int, float)):
            total_processed = int(processed_val)
        else:
            total_processed = 1 if processed_val else 0
            
        if isinstance(chunks_val, list):
            total_chunks = len(chunks_val)
        elif isinstance(chunks_val, (int, float)):
            total_chunks = int(chunks_val)
        else:
            total_chunks = 0
        
        # Show final results
        progress_bar.progress(1.0)
        
        if total_processed > 0:
            st.success("**Processing Complete**")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Files Processed", total_processed)
            with col2:
                st.metric("Chunks Created", total_chunks)
            with col3:
                st.metric("Project", current_project.name)
                
            # Show what's now available
            if st.session_state.get('extracted_document_content'):
                st.info("**Documents are now available for:**")
                st.markdown("""
                - **Chat & Analysis** - Ask questions about your documents
                - **Requirements Generation** - Generate structured requirements
                - **ARCADIA Analysis** - Perform systems engineering analysis
                """)
        else:
            st.error("No files were successfully processed")
            
    except Exception as e:
        st.error(f"Error during processing: {str(e)}")
        logger.error(f"Document processing error: {str(e)}")
    
    finally:
        # Cleanup temporary files
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except:
                pass
        
        progress_bar.empty()
        status_text.empty()
    
    return  # Exit here to avoid original duplicate detection logic
    
    # ORIGINAL DUPLICATE DETECTION CODE BELOW (NOT EXECUTED)
    # Save files temporarily and check for duplicates
    temp_files = []
    duplicate_info = []
    new_files = []
    all_extracted_content = []  # Store content from all files for requirements generation
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    try:
        status_text.text("Checking for duplicates and extracting content...")
        
        for i, uploaded_file in enumerate(uploaded_files):
            progress_bar.progress((i + 1) / (len(uploaded_files) * 2))  # First half for duplicate checking
            
            # Save temporarily
            temp_path = f"temp_{uploaded_file.name}"
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            temp_files.append(temp_path)
            
            # Extract content from file regardless of duplicate status
            file_content = extract_file_content(temp_path, uploaded_file.name)
            if file_content:
                all_extracted_content.append({
                    'filename': uploaded_file.name,
                    'content': file_content,
                    'size': len(file_content)
                })
            
            # Check for duplicates
            file_hash = rag_system.persistence_service.calculate_file_hash(temp_path)
            is_duplicate, doc_id, existing_project_id = rag_system.persistence_service.check_file_hash_globally(file_hash)
            
            if is_duplicate:
                duplicate_info.append({
                    'file': uploaded_file,
                    'temp_path': temp_path,
                    'existing_doc_id': doc_id,
                    'existing_project_id': existing_project_id,
                    'content': file_content  # Include content for requirements generation
                })
            else:
                new_files.append({
                    'file': uploaded_file,
                    'temp_path': temp_path,
                    'content': file_content  # Include content for requirements generation
                })
        
        # Store extracted content in session state for requirements generation
        if all_extracted_content:
            combined_content = "\n\n=== DOCUMENT SEPARATOR ===\n\n".join([
                f"=== {doc['filename']} ===\n{doc['content']}" 
                for doc in all_extracted_content
            ])
            st.session_state['extracted_document_content'] = combined_content
            st.session_state['extracted_files_info'] = all_extracted_content
        
        progress_bar.progress(0.5)
        status_text.text("Duplicate check complete!")
        
        # Show results with efficiency emphasis
        if duplicate_info:
            st.info(f" **Efficiency Notice:** Found {len(duplicate_info)} file(s) already processed in the system")
            
            # Create efficiency metrics
            total_duplicate_size = sum(dup['file'].size for dup in duplicate_info) / (1024 * 1024)  # MB
            
            # Use container instead of nested expander
            st.markdown(f"**📊Efficiency Savings ({len(duplicate_info)} files)**")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("💾 Storage Saved", f"{total_duplicate_size:.1f} MB")
            with col2:
                st.metric("⚡ Processing Saved", f"{len(duplicate_info)} files")
            with col3:
                st.metric("🧠 Embeddings Reused", f"{len(duplicate_info)} sets")
            
            st.info("""
            **🎯 Smart Deduplication Benefits:**
            • **Storage efficiency**
            
            • **Processing speed**
            
            • **Cost savings**
            
            • **Consistency**
            """)
            
            # Show duplicate details in a clean format
            st.markdown("**📄 Available Files:**")
            for dup in duplicate_info:
                st.write(f"• **{dup['file'].name}** - Available from project `{dup['existing_project_id']}`")
            
            # Show content availability for requirements generation
            available_content_files = [doc['filename'] for doc in st.session_state.get('extracted_files_info', [])]
            if available_content_files:
                st.success(f"✅ **Content ready for analysis:** {', '.join(available_content_files)}")
                st.info("🚀 **Instant access**: Pre-processed content and embeddings are immediately available!")
            
            st.markdown("**📋 Duplicate Handling Options:**")
            duplicate_action = st.radio(
                "How should we handle these efficient duplicates?",
                [
                    "🔗 Reuse existing (Recommended - Maximum efficiency)", 
                    "⏭️ Skip duplicates (Content still available for analysis)",
                    "🔄 Process anyway (Creates redundant copies - Not recommended)"
                ],
                key="duplicate_action",
                help="Recommended: Reuse existing for maximum storage and processing efficiency"
            )
            
            # Map back to original values for compatibility
            if duplicate_action.startswith("🔗"):
                duplicate_action = "Reuse existing (add to current project)"
            elif duplicate_action.startswith("⏭️"):
                duplicate_action = "Skip duplicates"
            else:
                duplicate_action = "Process anyway (create new entries)"
        else:
            duplicate_action = "Skip duplicates"  # Default when no duplicates
        
        # Determine files to process based on duplicate action
        files_to_process = new_files.copy()
        duplicate_files_to_process = []
        duplicate_files_to_reuse = []
        
        if duplicate_info and duplicate_action != "Skip duplicates":
            if duplicate_action == "Process anyway (create new entries)":
                # Add duplicate files to processing list
                duplicate_files_to_process = duplicate_info.copy()
                files_to_process.extend([{
                    'file': dup['file'],
                    'temp_path': dup['temp_path']
                } for dup in duplicate_info])
                st.info(f"🔄 Will process {len(duplicate_info)} duplicate file(s) as new entries")
                
            elif duplicate_action == "Reuse existing (add to current project)":
                # Prepare duplicate files for reuse
                duplicate_files_to_reuse = duplicate_info.copy()
                st.info(f"🔗 Will link {len(duplicate_info)} existing document(s) to current project")
        
        # Show processing summary
        total_files = len(files_to_process) + len(duplicate_files_to_reuse)
        if total_files > 0:
            if new_files:
                st.success(f"✅ {len(new_files)} new file(s) ready for processing")
            if duplicate_files_to_process:
                st.warning(f"🔄 {len(duplicate_files_to_process)} duplicate file(s) will be processed as new entries")
            if duplicate_files_to_reuse:
                st.info(f"🔗 {len(duplicate_files_to_reuse)} existing document(s) will be linked to project")
            
            # Single processing button for all actions
            button_text = f"🚀 Process {total_files} File(s)"
            if st.button(button_text, type="primary"):
                status_text.text("Processing documents...")
                
                try:
                    total_processed = 0
                    total_chunks = 0
                    all_errors = []
                    
                    # Process new files and duplicates marked for processing
                    if files_to_process:
                        status_text.text(f"Processing {len(files_to_process)} files...")
                        file_paths = [f['temp_path'] for f in files_to_process]
                        
                        if hasattr(rag_system, 'add_documents_to_project'):
                            results = rag_system.add_documents_to_project(file_paths, current_project.id)
                        else:
                            # Fallback for systems without project support
                            results = rag_system.add_documents_to_vectorstore(file_paths)
                        
                        total_processed += results.get('processed', 0)
                        total_chunks += results.get('new_chunks', 0) or results.get('chunks_added', 0)
                        if results.get('errors'):
                            all_errors.extend(results['errors'])
                    
                    # Handle reuse existing documents
                    if duplicate_files_to_reuse:
                        status_text.text(f"Linking {len(duplicate_files_to_reuse)} existing documents...")
                        reuse_success = 0
                        
                        for dup in duplicate_files_to_reuse:
                            try:
                                # Link existing document to current project
                                success = rag_system.persistence_service.link_document_to_project(
                                    dup['existing_doc_id'], 
                                    current_project.id
                                )
                                if success:
                                    reuse_success += 1
                                    logger.info(f"Linked existing document {dup['file'].name} to project {current_project.id}")
                                else:
                                    all_errors.append(f"Failed to link {dup['file'].name} to project")
                                    
                            except Exception as e:
                                error_msg = f"Error linking {dup['file'].name}: {str(e)}"
                                all_errors.append(error_msg)
                                logger.error(error_msg)
                        
                        total_processed += reuse_success
                        st.info(f"🔗 Successfully linked {reuse_success} existing document(s)")
                    
                    # Log the session with detailed information
                    action_description = []
                    if new_files:
                        action_description.append(f"processed {len(new_files)} new documents")
                    if duplicate_files_to_process:
                        action_description.append(f"processed {len(duplicate_files_to_process)} duplicates as new entries")
                    if duplicate_files_to_reuse:
                        action_description.append(f"linked {len(duplicate_files_to_reuse)} existing documents")
                    
                    rag_system.persistence_service.log_project_session(
                        current_project.id,
                        "document_upload",
                        f"Document upload: {', '.join(action_description)}",
                        {
                            "new_files": [f['file'].name for f in new_files],
                            "duplicate_action": duplicate_action,
                            "duplicates_processed": [dup['file'].name for dup in duplicate_files_to_process],
                            "duplicates_reused": [dup['file'].name for dup in duplicate_files_to_reuse],
                            "total_processed": total_processed,
                            "total_chunks": total_chunks
                        }
                    )
                    
                    # Show success message
                    st.success(f"✅ Successfully processed {total_processed} file(s)!")
                    if total_chunks > 0:
                        st.info(f"📊 Added {total_chunks} text chunks to the knowledge base")
                    
                    # Show errors if any
                    if all_errors:
                        st.error("⚠️ Some errors occurred:")
                        for error in all_errors:
                            st.error(f"• {error}")
                    
                    if total_processed > 0:
                        st.balloons()
                        st.success("🔄 Documents are now available! They should appear in the project.")
                        time.sleep(1)  # Give time for database to update
                        st.rerun()
                    
                except Exception as e:
                    st.error(f"❌ Processing error: {str(e)}")
                    rag_system.persistence_service.log_project_session(
                        current_project.id,
                        "document_upload_error",
                        f"Failed to process documents: {str(e)}"
                    )
        
        if not new_files and not duplicate_info:
            st.info("ℹ️ No files to process")
            
    finally:
        # Clean up temporary files
        for temp_path in temp_files:
            try:
                os.remove(temp_path)
            except:
                pass
        
        status_text.empty()
        progress_bar.empty()

def _query_project_documents_robust(rag_system, user_prompt, project_id, project_name, ready_docs):
    """
    Robust project-specific document query with similarity search on stored chunks
    
    Args:
        rag_system: The RAG system instance
        user_prompt: User's question
        project_id: Project ID
        project_name: Project name for context
        ready_docs: List of ready/processed documents
    
    Returns:
        Dict with 'answer' and 'sources' keys
    """
    try:
        # Method 1: Try dedicated project query method (preferred)
        if hasattr(rag_system, 'query_project_documents'):
            logger.info(f"Using dedicated query_project_documents method for project {project_id}")
            response_data = rag_system.query_project_documents(user_prompt, project_id)
            
            # Validate response format
            if isinstance(response_data, dict) and response_data.get('answer'):
                logger.info("✅ Project query successful")
                return response_data
            else:
                logger.warning("Project query returned invalid format, falling back")
        
        # Method 2: Enhanced ChromaDB similarity search with project filter
        if hasattr(rag_system, 'collection'):
            logger.info(f"🔍 Performing similarity search for project {project_id}")
            try:
                # Query with project filter for similarity search
                chroma_results = rag_system.collection.query(
                    query_texts=[user_prompt],
                    n_results=10,  # Get more results for better similarity filtering
                    where={"project_id": project_id}
                )
                
                # Build response from ChromaDB results with similarity scoring
                if chroma_results.get('documents') and chroma_results['documents'][0]:
                    context_docs = []
                    similarity_scores = chroma_results.get('distances', [[]])[0]
                    
                    for i, (doc, metadata, distance) in enumerate(zip(
                        chroma_results['documents'][0],
                        chroma_results.get('metadatas', [{}])[0],
                        similarity_scores
                    )):
                        # Add similarity score to metadata
                        enhanced_metadata = metadata.copy()
                        enhanced_metadata['similarity_score'] = 1 - distance  # Convert distance to similarity
                        enhanced_metadata['rank'] = i + 1
                        
                        context_docs.append(type('Document', (), {
                            'page_content': doc,
                            'metadata': enhanced_metadata
                        })())
                    
                    if context_docs:
                        # Filter for high-relevance chunks (similarity > 0.7)
                        relevant_docs = [doc for doc in context_docs if doc.metadata.get('similarity_score', 0) > 0.7]
                        
                        if not relevant_docs:
                            # If no high-relevance chunks, use top 3 results
                            relevant_docs = context_docs[:3]
                            logger.info(f"📊 No high-relevance chunks found, using top 3 results")
                        else:
                            logger.info(f"📊 Found {len(relevant_docs)} high-relevance chunks (similarity > 0.7)")
                        
                        # Create context from most relevant chunks
                        context_text = "\n\n".join([doc.page_content for doc in relevant_docs])
                        
                        # Enhanced prompt with similarity context
                        project_prompt = f"""Based on the following document excerpts from project '{project_name}', answer the question: {user_prompt}

Context from Project Documents (ranked by relevance):
{context_text}

IMPORTANT: 
- Base your answer ONLY on the provided document excerpts above
- If the excerpts don't contain enough information to answer the question, explicitly state this
- Include specific references to the document sources when possible
- Do not make assumptions beyond what's stated in the document excerpts

Question: {user_prompt}"""
                        
                        response = rag_system.ollama_client.chat(
                            model="llama3:instruct",
                            messages=[{"role": "user", "content": project_prompt}]
                        )
                        
                        logger.info(f"✅ ChromaDB similarity search successful - used {len(relevant_docs)} chunks")
                        return {
                            'answer': response["message"]["content"],
                            'sources': relevant_docs
                        }
            except Exception as chroma_error:
                logger.warning(f"ChromaDB similarity search failed: {str(chroma_error)}")
        
        # Method 3: Persistence service similarity search (for systems with stored chunks)
        if hasattr(rag_system, 'persistence_service') and ready_docs:
            logger.info(f"🔍 Performing manual similarity search on stored chunks for project {project_id}")
            try:
                # Get stored chunks from persistence service
                project_chunks = rag_system.persistence_service.get_project_chunks(project_id)
                
                if project_chunks:
                    # Filter chunks by ready documents
                    ready_filenames = [doc.filename for doc in ready_docs]
                    ready_chunks = [
                        chunk for chunk in project_chunks 
                        if chunk.get('metadata', {}).get('source_filename') in ready_filenames
                    ]
                    
                    if ready_chunks:
                        # Perform similarity search on chunks
                        similar_chunks = _calculate_chunk_similarity(user_prompt, ready_chunks)
                        
                        if similar_chunks:
                            # Use top similar chunks for context
                            top_chunks = similar_chunks[:3]
                            
                            context_text = "\n\n".join([
                                f"[Source: {chunk['metadata'].get('source_filename', 'Unknown')}]\n{chunk['content']}"
                                for chunk in top_chunks
                            ])
                            
                            project_prompt = f"""Based on the following document excerpts from project '{project_name}', answer the question: {user_prompt}

Context from Project Documents (ranked by similarity):
{context_text}

IMPORTANT: 
- Base your answer ONLY on the provided document excerpts above
- If the excerpts don't contain enough information to answer the question, explicitly state this
- Include specific references to the document sources when possible

Question: {user_prompt}"""
                            
                            response = rag_system.ollama_client.chat(
                                model="llama3:instruct",
                                messages=[{"role": "user", "content": project_prompt}]
                            )
                            
                            # Create mock sources for display
                            mock_sources = []
                            for chunk in top_chunks:
                                mock_sources.append(type('Document', (), {
                                    'page_content': chunk['content'],
                                    'metadata': chunk['metadata']
                                })())
                            
                            logger.info(f"✅ Manual similarity search successful - used {len(top_chunks)} chunks")
                            return {
                                'answer': response["message"]["content"],
                                'sources': mock_sources
                            }
            except Exception as manual_error:
                logger.warning(f"Manual similarity search failed: {str(manual_error)}")
        
        # Method 4: Fallback - inform user about document availability
        logger.warning("No similarity search method available")
        
        if ready_docs:
            doc_list = ", ".join([doc.filename for doc in ready_docs])
            return {
                'answer': f"I have access to {len(ready_docs)} documents in project '{project_name}': {doc_list}. However, I need a properly configured similarity search system to answer questions about their content. Please ensure your RAG system supports project-specific document queries.",
                'sources': []
            }
        else:
            return {
                'answer': f"No processed documents are available in project '{project_name}'. Please upload and process documents first before asking questions.",
                'sources': []
            }
    
    except Exception as e:
        logger.error(f"Error in robust project query: {str(e)}")
        return {
            'answer': f"I encountered an error while searching the project documents: {str(e)}. Please try again or contact support if the issue persists.",
            'sources': []
        }


def _calculate_chunk_similarity(query, chunks):
    """
    Calculate similarity between query and chunks using text-based scoring
    
    Args:
        query: User query string
        chunks: List of chunk dictionaries with 'content' and 'metadata'
    
    Returns:
        List of chunks sorted by similarity score (highest first)
    """
    import re
    
    # Normalize query
    query_lower = query.lower()
    query_terms = set(re.findall(r'\b\w+\b', query_lower))
    
    # Remove common stop words
    stop_words = {
        'a', 'an', 'and', 'are', 'as', 'at', 'be', 'by', 'for', 'from', 'has', 'he', 'in', 'is', 'it', 'its',
        'of', 'on', 'that', 'the', 'to', 'was', 'were', 'will', 'with', 'what', 'when', 'where', 'who', 'why', 'how'
    }
    query_terms = query_terms - stop_words
    
    if not query_terms:
        return chunks[:5]  # Return first 5 if no meaningful terms
    
    scored_chunks = []
    
    for chunk in chunks:
        content = chunk.get('content', '').lower()
        content_terms = set(re.findall(r'\b\w+\b', content))
        
        # Calculate similarity score
        score = 0
        
        # Exact term matches
        common_terms = query_terms.intersection(content_terms)
        score += len(common_terms) * 2
        
        # Phrase matching (bonus for consecutive terms)
        for term in query_terms:
            if term in content:
                score += content.count(term) * 1.5
        
        # Length normalization
        if len(content) > 0:
            score = score / (len(content) / 1000)  # Normalize by content length
        
        if score > 0:
            chunk_with_score = chunk.copy()
            chunk_with_score['similarity_score'] = round(score, 2)
            scored_chunks.append(chunk_with_score)
    
    # Sort by similarity score (highest first)
    scored_chunks.sort(key=lambda x: x['similarity_score'], reverse=True)
    
    return scored_chunks


def _load_project_documents_safely(rag_system, project_id):
    """
    Safely load project documents with error handling and recovery
    
    Args:
        rag_system: The RAG system instance
        project_id: Project ID to load documents for
    
    Returns:
        List of processed documents (empty list on error)
    """
    try:
        # Load documents from persistence service
        documents = rag_system.persistence_service.get_project_documents(project_id)
        
        # Filter for completed documents only
        completed_docs = [doc for doc in documents if doc.processing_status == "completed"]
        
        logger.info(f"Loaded {len(completed_docs)} completed documents for project {project_id}")
        return completed_docs
        
    except Exception as e:
        logger.error(f"Error loading project documents for {project_id}: {str(e)}")
        # Return empty list on error - don't crash the app
        return []


def _load_project_data_safely(rag_system, project_id):
    """
    Safely load all project data with error handling
    
    Args:
        rag_system: The RAG system instance
        project_id: Project ID to load data for
    
    Returns:
        Dict with safely loaded data or empty defaults
    """
    try:
        # Initialize with safe defaults
        project_data = {
            "documents": [],
            "chunks": [],
            "requirements": {},
            "stakeholders": [],
            "arcadia_analyses": [],
            "sessions": [],
            "errors": []
        }
        
        # Load documents safely
        try:
            documents = rag_system.persistence_service.get_project_documents(project_id)
            project_data["documents"] = [doc for doc in documents if doc.processing_status == "completed"]
        except Exception as e:
            logger.warning(f"Could not load documents for project {project_id}: {str(e)}")
            project_data["errors"].append(f"Documents: {str(e)}")
        
        # Load chunks safely
        try:
            chunks = rag_system.persistence_service.get_project_chunks(project_id)
            project_data["chunks"] = chunks
        except Exception as e:
            logger.warning(f"Could not load chunks for project {project_id}: {str(e)}")
            project_data["errors"].append(f"Chunks: {str(e)}")
        
        # Load requirements safely
        try:
            requirements = rag_system.persistence_service.get_project_requirements(project_id)
            project_data["requirements"] = requirements
        except Exception as e:
            logger.warning(f"Could not load requirements for project {project_id}: {str(e)}")
            project_data["errors"].append(f"Requirements: {str(e)}")
        
        # Load stakeholders safely
        try:
            stakeholders = rag_system.persistence_service.get_project_stakeholders(project_id)
            project_data["stakeholders"] = stakeholders
        except Exception as e:
            logger.warning(f"Could not load stakeholders for project {project_id}: {str(e)}")
            project_data["errors"].append(f"Stakeholders: {str(e)}")
        
        # Load ARCADIA analyses safely
        try:
            analyses = rag_system.persistence_service.get_project_arcadia_analyses(project_id)
            project_data["arcadia_analyses"] = analyses
        except Exception as e:
            logger.warning(f"Could not load ARCADIA analyses for project {project_id}: {str(e)}")
            project_data["errors"].append(f"ARCADIA Analyses: {str(e)}")
        
        # Load sessions safely
        try:
            sessions = rag_system.persistence_service.get_project_sessions(project_id, limit=10)
            project_data["sessions"] = sessions
        except Exception as e:
            logger.warning(f"Could not load sessions for project {project_id}: {str(e)}")
            project_data["errors"].append(f"Sessions: {str(e)}")
        
        return project_data
        
    except Exception as e:
        logger.error(f"Critical error loading project data for {project_id}: {str(e)}")
        # Return safe defaults even on critical error
        return {
            "documents": [],
            "chunks": [],
            "requirements": {},
            "stakeholders": [],
            "arcadia_analyses": [],
            "sessions": [],
            "errors": [f"Critical error: {str(e)}"]
        }


def _log_error_safely(rag_system, project_id, error_type, error_message):
    """
    Safely log errors without crashing the app
    
    Args:
        rag_system: The RAG system instance
        project_id: Project ID (can be None)
        error_type: Type of error
        error_message: Error message
    """
    try:
        if hasattr(rag_system, 'persistence_service') and project_id:
            rag_system.persistence_service.log_project_session(
                project_id,
                error_type,
                error_message
            )
        logger.error(f"Error logged for project {project_id}: {error_type} - {error_message}")
    except Exception as e:
        # Even error logging failed - just log to console
        logger.error(f"Failed to log error: {str(e)}")


def _handle_missing_data_gracefully(data_type, project_name=None):
    """
    Handle missing data gracefully with user-friendly messages
    
    Args:
        data_type: Type of missing data
        project_name: Optional project name for context
    
    Returns:
        Tuple of (should_continue, message)
    """
    project_context = f" for project '{project_name}'" if project_name else ""
    
    if data_type == "documents":
        st.info(f"📭 No documents found{project_context}. Upload some documents to get started!")
        return True, "No documents"
    
    elif data_type == "chunks":
        st.warning(f"⚠️ No text chunks available{project_context}. This may indicate documents are still processing or there was a processing error.")
        return True, "No chunks"
    
    elif data_type == "requirements":
        st.info(f"📝 No requirements generated{project_context}. Generate some requirements to view analysis.")
        return True, "No requirements"
    
    elif data_type == "project":
        st.error("❌ Project data is missing or corrupted. Please select a different project or contact support.")
        return False, "Missing project"
    
    else:
        st.warning(f"⚠️ Some data is missing{project_context}. The application will continue with available data.")
        return True, "Missing data"


def _handle_refresh_project_data(rag_system, current_project):
    """
    Handle refreshing project data with proper deletion and safety checks
    
    Args:
        rag_system: The RAG system instance
        current_project: Current project object
    """
    try:
        # Create a confirmation dialog
        st.warning("⚠️ **Data Refresh Warning**")
        st.markdown(f"""
        **This action will:**
        • Delete all processed documents and chunks for project '{current_project.name}'
        • Remove all requirements and ARCADIA analyses
        • Clear all stakeholder information
        • Reset project chat history
        • Keep project metadata (name, description, etc.)
        
        ⚠️ **This action cannot be undone!**
        """)
        
        # Confirmation checkbox
        confirm_refresh = st.checkbox("I understand that this will delete all project data except basic project information")
        
        if confirm_refresh:
            # Additional text confirmation
            confirmation_text = st.text_input(
                "Type 'DELETE DATA' to confirm the refresh action",
                placeholder="DELETE DATA"
            )
            
            if confirmation_text == "DELETE DATA":
                if st.button("🗑️ Confirm Data Refresh", type="secondary"):
                    with st.spinner("🔄 Refreshing project data..."):
                        # Call the new refresh function
                        success = _refresh_project_data_safely(rag_system, current_project.id)
                        
                        if success:
                            st.success("✅ Project data refreshed successfully!")
                            st.info("🔄 Reloading page with fresh data...")
                            time.sleep(2)
                            st.rerun()
                        else:
                            st.error("❌ Failed to refresh project data. Please try again.")
            else:
                if confirmation_text:
                    st.error("❌ Confirmation text doesn't match. Please type 'DELETE DATA' exactly.")
        
    except Exception as e:
        logger.error(f"Error in refresh handler: {str(e)}")
        st.error(f"❌ Refresh error: {str(e)}")


def _refresh_project_data_safely(rag_system, project_id):
    """
    Safely refresh project data by deleting all associated data except project metadata
    
    Args:
        rag_system: The RAG system instance
        project_id: Project ID to refresh
    
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        logger.info(f"Starting data refresh for project {project_id}")
        
        # Get persistence service
        persistence_service = rag_system.persistence_service
        
        # Log the refresh action
        persistence_service.log_project_session(
            project_id,
            "data_refresh_start",
            "Starting complete project data refresh"
        )
        
        # Delete data in safe order (respecting foreign key constraints)
        with persistence_service._get_db_connection() as conn:
            cursor = conn.cursor()
            
            # Step 1: Delete document chunks
            cursor.execute("DELETE FROM document_chunks WHERE project_id = ?", (project_id,))
            chunks_deleted = cursor.rowcount
            logger.info(f"Deleted {chunks_deleted} chunks")
            
            # Step 2: Delete processed documents
            cursor.execute("DELETE FROM processed_documents WHERE project_id = ?", (project_id,))
            docs_deleted = cursor.rowcount
            logger.info(f"Deleted {docs_deleted} documents")
            
            # Step 3: Delete requirements
            cursor.execute("DELETE FROM requirements WHERE project_id = ?", (project_id,))
            reqs_deleted = cursor.rowcount
            logger.info(f"Deleted {reqs_deleted} requirements")
            
            # Step 4: Delete ARCADIA analyses
            cursor.execute("DELETE FROM arcadia_analyses WHERE project_id = ?", (project_id,))
            analyses_deleted = cursor.rowcount
            logger.info(f"Deleted {analyses_deleted} ARCADIA analyses")
            
            # Step 5: Delete stakeholders
            cursor.execute("DELETE FROM stakeholders WHERE project_id = ?", (project_id,))
            stakeholders_deleted = cursor.rowcount
            logger.info(f"Deleted {stakeholders_deleted} stakeholders")
            
            # Step 6: Reset project counters
            cursor.execute("""
                UPDATE projects 
                SET documents_count = 0, 
                    requirements_count = 0, 
                    updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
            """, (project_id,))
            
            conn.commit()
            
            # Clear ChromaDB entries if available
            if hasattr(rag_system, 'collection'):
                try:
                    # Delete from ChromaDB
                    chroma_results = rag_system.collection.get(where={"project_id": project_id})
                    if chroma_results.get('ids'):
                        rag_system.collection.delete(ids=chroma_results['ids'])
                        logger.info(f"Deleted {len(chroma_results['ids'])} ChromaDB entries")
                except Exception as chroma_error:
                    logger.warning(f"ChromaDB deletion error: {str(chroma_error)}")
            
            # Clear caches
            if hasattr(st, 'cache_data'):
                st.cache_data.clear()
            if hasattr(st, 'cache_resource'):
                st.cache_resource.clear()
            
            # Log successful completion
            persistence_service.log_project_session(
                project_id,
                "data_refresh_complete",
                f"Successfully refreshed project data: {docs_deleted} docs, {chunks_deleted} chunks, {reqs_deleted} requirements, {analyses_deleted} analyses, {stakeholders_deleted} stakeholders deleted"
            )
            
            logger.info(f"Successfully refreshed project {project_id}")
            return True
            
    except Exception as e:
        logger.error(f"Error refreshing project data: {str(e)}")
        
        # Log the failure
        try:
            persistence_service.log_project_session(
                project_id,
                "data_refresh_failed",
                f"Failed to refresh project data: {str(e)}"
            )
        except:
            pass
        
        return False





def project_documents_tab(rag_system, current_project, has_project_management):
    """Document Management tab - Focused on documents only"""
    st.markdown("### Document Management")
    
    # Initialize project_documents to avoid undefined variable errors
    project_documents = []
    
    # Load project documents with enhanced safety
    if has_project_management and current_project:
        try:
            project_documents = _load_project_documents_safely(rag_system, current_project.id)
        except Exception as e:
            st.error(f"❌ Error loading project documents: {str(e)}")
            project_documents = []
            _log_error_safely(rag_system, current_project.id, "document_loading_error", str(e))
    
    # Simple status check without detailed project metrics
    if has_project_management and not current_project:
        st.warning("No project selected. Please create or select a project in the sidebar.")
        st.info("""
        **To create a new project:**
        1. Use the sidebar on the left to create a new project
        2. Or select an existing project from the dropdown
        3. Then return here to manage documents and chat
        """)
        return
    elif not has_project_management:
        st.info("**Traditional Mode** - Basic document management available")
    
    # Document upload section
    with st.expander("Upload Documents", expanded=False):
        if has_project_management and current_project:
            st.markdown(f"**Upload to project: {current_project.name}**")
            
            uploaded_files = st.file_uploader(
                "Select files",
                accept_multiple_files=True,
                type=['pdf', 'docx', 'txt', 'md', 'xml', 'json', 'aird'],
                help="Upload multiple documents. Smart deduplication saves storage and processing time.",
                key="project_document_uploader"
            )
            
            if uploaded_files:
                st.info(f"📄 {len(uploaded_files)} file(s) selected for upload")
                
                if st.button("🚀 Process Documents", type="primary", key="process_project_docs"):
                    process_documents_with_duplicate_detection(rag_system, current_project, uploaded_files)
        else:
            st.markdown("**Upload to knowledge base**")
            uploaded_files = st.file_uploader(
                "Add documents to the knowledge base",
                accept_multiple_files=True,
                type=['pdf', 'docx', 'txt', 'md', 'xml', 'json', 'aird', 'capella'],
                help="Upload multiple documents to enhance the chat knowledge base",
                key="general_document_uploader"
            )
            
            if uploaded_files:
                st.info(f"{len(uploaded_files)} file(s) selected")
                
                if st.button("Process Documents", type="primary", key="process_general_docs"):
                    with st.spinner("Processing documents..."):
                        logger.info(f"Processing {len(uploaded_files)} documents...")
                        
                        temp_paths = []
                        try:
                            # Save uploaded files temporarily
                            for uploaded_file in uploaded_files:
                                temp_path = Path(f"temp_{uploaded_file.name}")
                                with open(temp_path, "wb") as f:
                                    f.write(uploaded_file.getbuffer())
                                temp_paths.append(str(temp_path))
                            
                            # Process documents through RAG system
                            results = rag_system.add_documents_to_vectorstore(temp_paths)
                            
                            # Display results
                            st.success(f"Processed {results.get('processed', 0)} files successfully!")
                            st.info(f"Added {results.get('chunks_added', 0)} text chunks to knowledge base")
                            
                            if results.get('errors'):
                                st.error("Processing Errors:")
                                for error in results['errors']:
                                    st.error(f"{error}")
                            
                        except Exception as e:
                            logger.error(f"Error processing documents: {str(e)}")
                            st.error(f"Error processing documents: {str(e)}")
                        finally:
                            # Clean up temporary files
                            for temp_path in temp_paths:
                                try:
                                    os.remove(temp_path)
                                except:
                                    pass
    
    # Document list and preview
    if has_project_management and current_project:
        st.markdown("#### 📋 Project Documents")
        try:
            documents = rag_system.persistence_service.get_project_documents(current_project.id)
            
            if documents:
                for doc in documents:
                    with st.container():
                        # Document info
                        col1, col2 = st.columns([3, 1])
                        with col1:
                            st.write(f"**{doc.filename}**")
                            st.caption(f"Size: {doc.file_size / 1024:.1f} KB • Chunks: {doc.chunks_count}")
                        with col2:
                            status_color = "🟢" if doc.processing_status == "completed" else "🟡"
                            st.write(f"{status_color} {doc.processing_status}")
                        
                        # Preview option
                        if st.button(f"👁️ Preview", key=f"preview_{doc.id}"):
                            st.session_state[f"show_preview_{doc.id}"] = not st.session_state.get(f"show_preview_{doc.id}", False)
                        
                        if st.session_state.get(f"show_preview_{doc.id}", False):
                            try:
                                # Get document chunks for preview
                                chunks = rag_system.persistence_service.get_document_chunks(doc.id)
                                
                                if chunks:
                                    # Combine chunks into preview text
                                    preview_text = "\n\n".join([chunk["content"] for chunk in chunks])
                                    
                                    # Limit preview to first 2000 characters for readability
                                    if len(preview_text) > 2000:
                                        preview_text = preview_text[:2000] + "\n\n... (content truncated for preview)"
                                    
                                    st.text_area(f"Preview of {doc.filename} ({len(chunks)} chunks)", 
                                               value=preview_text, 
                                               height=200, disabled=True, key=f"preview_content_{doc.id}")
                                    
                                    # Show additional info
                                    st.caption(f"📊 Document info: {doc.file_size / 1024:.1f} KB • {len(chunks)} chunks • {len(preview_text)} characters shown")
                                else:
                                    st.info("No content available for preview. Document may still be processing.")
                                    
                            except Exception as e:
                                st.error(f"Error loading preview: {str(e)}")
                        
                        st.markdown("---")
            else:
                st.info("📭 No documents uploaded yet. Upload some documents to get started!")
                
        except Exception as e:
            st.error(f"❌ Error loading documents: {str(e)}")
    else:
        st.markdown("#### 📄 Available Documents")
        st.info("In traditional mode, document management is handled globally. Upload documents above to add them to the knowledge base.")
    
    # Chat with Documents section (moved under document management)
    st.markdown("#### 💬 Chat with Documents")
    
    # Project-specific chat management
    if has_project_management and current_project:
        # Ensure we have project chats loaded for current project
        project_id = current_project.id
        if project_id not in st.session_state.project_chats:
            st.session_state.project_chats[project_id] = load_project_chats(project_id)
        
        # Check if we need to switch chat context when project changes
        if st.session_state.current_project_chat_context != project_id:
            st.session_state.current_project_chat_context = project_id
            st.session_state.current_chat_id = None  # Reset chat selection when switching projects
        
        current_project_chats = st.session_state.project_chats[project_id]
        
        # Chat management
        col1, col2 = st.columns([3, 1])
        with col1:
            st.markdown(f"**Project: {current_project.name}** - Chat with project documents only")
            
            if st.button("📝 New Project Chat", type="primary", use_container_width=True, key="new_project_chat"):
                new_chat_id = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
                st.session_state.current_chat_id = new_chat_id
                current_project_chats[new_chat_id] = {
                    "title": "New Project Chat",
                    "messages": [],
                    "created_at": datetime.now().isoformat(),
                    "project_id": project_id,
                    "project_name": current_project.name,
                    "document_count": len(project_documents) if 'project_documents' in locals() else 0
                }
                # Update session state and save
                st.session_state.project_chats[project_id] = current_project_chats
                save_project_chats(current_project_chats, project_id)
                st.rerun()
        
        with col2:
            # Project-specific chat history dropdown
            if current_project_chats:
                chat_options = {chat_id: chat_data.get('title', 'Untitled Chat') 
                              for chat_id, chat_data in sorted(current_project_chats.items(), 
                                                             key=lambda x: x[1].get('created_at', ''), reverse=True)}
                
                # Only show dropdown if there are chats
                if chat_options:
                    # Add placeholder option
                    chat_options_with_placeholder = {"": "Select a chat..."} | chat_options
                    
                    selected_chat = st.selectbox(
                        "Project Chat History", 
                        options=list(chat_options_with_placeholder.keys()), 
                        format_func=lambda x: chat_options_with_placeholder[x],
                        index=0,  # Start with placeholder
                        key="project_chat_history_selector"
                    )
                    
                    if selected_chat and selected_chat != st.session_state.current_chat_id:
                        st.session_state.current_chat_id = selected_chat
                        st.rerun()
                else:
                    st.caption("No chats yet")
            else:
                st.caption("No project chats")
    
    # Chat interface
    if has_project_management and current_project and st.session_state.current_chat_id:
        # Get current project chats
        project_id = current_project.id
        current_project_chats = st.session_state.project_chats.get(project_id, {})
        
        if st.session_state.current_chat_id in current_project_chats:
            current_chat = current_project_chats[st.session_state.current_chat_id]
            
            # Chat title
            chat_title = st.text_input(
                "Chat Title", 
                value=current_chat.get('title', 'New Project Chat'),
                key="integrated_project_chat_title"
            )
            
            # Update title if changed
            if chat_title != current_chat.get('title', 'New Project Chat'):
                current_chat['title'] = chat_title
                # Update session state and save
                st.session_state.project_chats[project_id] = current_project_chats
                save_project_chats(current_project_chats, project_id)
        
            # Chat messages container with scroll
            chat_container = st.container()
            
            with chat_container:
                # Display existing messages
                for i, message in enumerate(current_chat["messages"]):
                    if message["role"] == "user":
                        st.markdown(f"""
                        <div class="chat-message user-message">
                            <strong>You:</strong><br>
                            {message["content"]}
                        </div>
                        """, unsafe_allow_html=True)
                    else:  # assistant
                        st.markdown(f"""
                        <div class="chat-message assistant-message">
                            <strong>Assistant:</strong><br>
                            {message["content"]}
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # Show context sources if available (project-specific)
                        if "context" in message and message["context"]:
                            with st.expander(f"📄 Project Sources ({len(message['context'])})"):
                                for j, doc in enumerate(message["context"]):
                                    source_name = doc.get('metadata', {}).get('source', 'Unknown')
                                    source_filename = source_name.split('/')[-1] if source_name != 'Unknown' else 'Unknown'
                                    
                                    # Check if this source is from the current project
                                    is_project_source = any(project_doc.filename == source_filename 
                                                          for project_doc in project_documents)
                                    
                                    source_indicator = "✅" if is_project_source else "⚠️"
                                    
                                    # Get additional metadata for stored chunks
                                    relevance_score = doc.get('metadata', {}).get('relevance_score', 'N/A')
                                    chunk_length = doc.get('metadata', {}).get('chunk_length', len(doc.get('content', '')))
                                    
                                    st.markdown(f"""
                                    <div class="source-citation">
                                        <strong>{source_indicator} Source {j+1}:</strong> {source_filename}<br>
                                        <strong>Project:</strong> {current_project.name} {'(Verified)' if is_project_source else '(External)'}<br>
                                        <strong>Relevance Score:</strong> {relevance_score} | <strong>Chunk Size:</strong> {chunk_length} chars<br>
                                        <strong>Preview:</strong> {doc.get('content', '')[:200]}...
                                    </div>
                                    """, unsafe_allow_html=True)
                        
                        # Show debug info if available (for troubleshooting)
                        if "debug_info" in message and message["debug_info"]:
                            with st.expander("🔧 Debug Info (Technical Details)"):
                                debug_info = message["debug_info"]
                                
                                debug_col1, debug_col2 = st.columns(2)
                                with debug_col1:
                                    st.write(f"**Ready Documents:** {debug_info.get('ready_docs_count', 'Unknown')}")
                                    st.write(f"**Context Sources:** {debug_info.get('context_docs_count', 'Unknown')}")
                                    st.write(f"**Query Method:** {debug_info.get('query_method_used', 'Unknown')}")
                                    
                                    # Show chunk-specific info
                                    if debug_info.get('chunks_method') == 'stored_chunks':
                                        st.write(f"**Chunks Source:** {debug_info.get('chunks_searched', 'Unknown')}")
                                        if debug_info.get('avg_relevance_score'):
                                            st.write(f"**Avg Relevance:** {debug_info.get('avg_relevance_score', 0):.1f}")
                                        if debug_info.get('query_terms_used'):
                                            st.write(f"**Query Terms:** {', '.join(debug_info.get('query_terms_used', []))}")
                                
                                with debug_col2:
                                    st.write(f"**Project Query Available:** {'✅' if debug_info.get('has_project_query_method', False) else '❌'}")
                                    st.write(f"**Persistence Service:** {'✅' if debug_info.get('has_persistence_service', False) else '❌'}")
                                    
                                    # Show chunk-specific info
                                    if debug_info.get('chunks_method') == 'stored_chunks':
                                        st.write(f"**Using Stored Chunks:** ✅")
                                        if debug_info.get('max_relevance_score'):
                                            st.write(f"**Max Relevance:** {debug_info.get('max_relevance_score', 0)}")
                                
                                # Status indicators
                                if debug_info.get('chunks_method') == 'stored_chunks':
                                    if debug_info.get('context_docs_count', 0) > 0:
                                        st.success("✅ Using stored chunks directly from project database")
                                    else:
                                        st.warning("⚠️ No relevant chunks found in project storage")
                                elif debug_info.get('context_docs_count', 0) == 0:
                                    st.warning("⚠️ No context sources found - this may indicate a RAG system issue")
                                elif debug_info.get('ready_docs_count', 0) == 0:
                                    st.warning("⚠️ No ready documents found - upload and process documents first")
                                else:
                                    st.success("✅ System appears to be working correctly")
            
            # Show project documents available for chat BEFORE chat input
            ready_docs = [doc for doc in project_documents if doc.processing_status == "completed"]
            if 'project_documents' in locals() and project_documents:
                with st.expander(f"Available Documents ({len(project_documents)})", expanded=False):
                    total_chunks = sum(doc.chunks_count for doc in project_documents)
                    total_size = sum(doc.file_size for doc in project_documents) / 1024  # KB
                    
                    st.info(f"**Knowledge Base**: {len(project_documents)} documents, {total_chunks} text chunks, {total_size:.1f} KB")
                    
                    for doc in project_documents:
                        status_text = "Ready" if doc.processing_status == "completed" else "Processing..."
                        
                        col1, col2 = st.columns([3, 1])
                        with col1:
                            st.write(f"**{doc.filename}**")
                            st.caption(f"{doc.chunks_count} chunks • {doc.file_size / 1024:.1f} KB • {status_text}")
                        with col2:
                            if doc.processing_status == "completed":
                                st.success("Available")
                            else:
                                st.warning("Processing")

            # Chat input with project context
            if ready_docs:
                placeholder_text = f"Ask about {current_project.name} documents ({len(ready_docs)} available), ARCADIA methodology, or MBSE concepts..."
            else:
                placeholder_text = f"Upload documents to {current_project.name} first, then ask questions..."
            
            user_prompt = st.chat_input(
                placeholder_text, 
                key="integrated_project_chat_input",
                disabled=len(ready_docs) == 0
            )
            
            # Show chat scope indicator
            if ready_docs:
                st.caption(f"**Chat scope**: {len(ready_docs)} document(s) from {current_project.name} • {sum(doc.chunks_count for doc in ready_docs)} text chunks available")
            else:
                st.caption("**No documents available** - Upload and process documents to enable chat")
            
            if user_prompt:
                logger.info(f"New project chat message in {current_project.name}: {user_prompt[:100]}...")
                
                # Check if we have any processed documents
                if not ready_docs:
                    st.error("**No documents available for chat!** Please upload and process documents first.")
                    st.info("Upload documents in the section above, then wait for processing to complete.")
                    return
                
                # Update chat title if it's the first message
                if not current_chat["messages"]:
                    current_chat["title"] = user_prompt[:50] + ("..." if len(user_prompt) > 50 else "")
                
                # Add user message
                current_chat["messages"].append({
                    "role": "user", 
                    "content": user_prompt,
                    "timestamp": datetime.now().isoformat(),
                    "project_id": project_id,
                    "project_name": current_project.name
                })
                
                # Generate response with project-specific context
                with st.spinner(f"Analyzing {current_project.name} documents and generating response..."):
                    try:
                        response_data = _query_project_documents_robust(
                            rag_system, user_prompt, project_id, current_project.name, ready_docs
                        )
                        
                        # Extract response and context with better error handling
                        if isinstance(response_data, dict):
                            response = response_data.get('answer', '')
                            context_docs = response_data.get('sources', [])
                            
                            # Ensure we have a valid response
                            if not response or response.strip() == '':
                                response = f"I couldn't generate a response from the available documents in project '{current_project.name}'. Please try rephrasing your question or check if the documents contain relevant information."
                                logger.warning("Empty response from RAG system")
                        else:
                            response = str(response_data) if response_data else "No response generated from RAG system."
                            context_docs = []
                            logger.info(f"Non-dict response from RAG: {type(response_data)}")
                        
                        # Final safety check
                        if not response or response.strip() == '':
                            response = f"I'm having trouble generating a response. Please ensure that:\n1. Documents are uploaded to project '{current_project.name}'\n2. Documents are fully processed\n3. Your question relates to the document content"
                            logger.error("Final safety check: empty response")
                        
                        logger.info(f"Generated project-specific response with {len(context_docs)} context sources from project {current_project.name}")
                        
                        # Add assistant response with debug info
                        assistant_message = {
                            "role": "assistant",
                            "content": response,
                            "context": [{"content": doc.page_content, "metadata": doc.metadata} for doc in context_docs] if context_docs else [],
                            "timestamp": datetime.now().isoformat(),
                            "project_id": project_id,
                            "project_name": current_project.name
                        }
                        
                        # Add debug info for troubleshooting
                        debug_info = {
                            "ready_docs_count": len(ready_docs),
                            "context_docs_count": len(context_docs),
                            "has_project_query_method": hasattr(rag_system, 'query_project_documents'),
                            "has_persistence_service": hasattr(rag_system, 'persistence_service'),
                            "query_method_used": "project_documents" if hasattr(rag_system, 'query_project_documents') else "stored_chunks" if hasattr(rag_system, 'persistence_service') else "general"
                        }
                        
                        # Add chunk-specific info if using stored chunks
                        if hasattr(rag_system, 'persistence_service') and not hasattr(rag_system, 'query_project_documents'):
                            debug_info["chunks_method"] = "stored_chunks"
                            debug_info["chunks_searched"] = "direct_project_chunks"
                            debug_info["query_terms_used"] = user_prompt.split()
                            # Add relevance scores if available
                            if context_docs:
                                scores = [doc.metadata.get('relevance_score', 0) for doc in context_docs if isinstance(doc.metadata.get('relevance_score'), int)]
                                if scores:
                                    debug_info["avg_relevance_score"] = sum(scores) / len(scores)
                                    debug_info["max_relevance_score"] = max(scores)
                        
                        assistant_message["debug_info"] = debug_info
                        
                        current_chat["messages"].append(assistant_message)
                        
                        # Save project chats
                        st.session_state.project_chats[project_id] = current_project_chats
                        save_project_chats(current_project_chats, project_id)
                        
                        # Rerun to show new messages
                        st.rerun()
                        
                    except Exception as e:
                        logger.error(f"Error generating project chat response: {str(e)}")
                        error_message = f"I apologize, but I encountered an error: {str(e)}"
                        
                        current_chat["messages"].append({
                            "role": "assistant",
                            "content": error_message,
                            "timestamp": datetime.now().isoformat(),
                            "project_id": project_id,
                            "project_name": current_project.name
                        })
                        
                        # Save project chats
                        st.session_state.project_chats[project_id] = current_project_chats
                        save_project_chats(current_project_chats, project_id)
                        st.rerun()
        else:
            st.warning("Chat session not found. Please create a new chat or select an existing one.")
    
    elif has_project_management and current_project:
        # Project exists but no chat selected
        st.info("**Project Chat Available**")
        st.markdown(f"""
        **Ready to chat with {current_project.name} documents!**
        
        **Available in this project:**
        • {len(project_documents) if 'project_documents' in locals() else 0} processed documents
        
        **Get started:**
        1. Click "New Project Chat" to start
        2. Or select an existing chat from history
        3. Ask questions about your project documents
        """)
        
        # Show project documents available for chat
        if 'project_documents' in locals() and project_documents:
            with st.expander(f"📄 Available Documents for Chat ({len(project_documents)})"):
                total_chunks = sum(doc.chunks_count for doc in project_documents)
                total_size = sum(doc.file_size for doc in project_documents) / 1024  # KB
                
                st.info(f"📊 **Knowledge Base**: {len(project_documents)} documents, {total_chunks} text chunks, {total_size:.1f} KB")
                
                for doc in project_documents:
                    status_icon = "🟢" if doc.processing_status == "completed" else "🟡"
                    status_text = "Ready for chat" if doc.processing_status == "completed" else "Processing..."
                    
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.write(f"{status_icon} **{doc.filename}**")
                        st.caption(f"📊 {doc.chunks_count} chunks • {doc.file_size / 1024:.1f} KB • {status_text}")
                    with col2:
                        if doc.processing_status == "completed":
                            st.success("✅ Available")
                        else:
                            st.warning("⏳ Processing")
                
                # Chat readiness indicator
                ready_docs = [doc for doc in project_documents if doc.processing_status == "completed"]
                if ready_docs:
                    st.success(f"🎯 **Chat Ready**: {len(ready_docs)} document(s) processed and available for questions")
                else:
                    st.warning("⏳ **Processing**: Documents are being processed for chat availability")
        else:
            st.warning("📭 No documents in this project yet. Upload documents above to enable chat.")
            st.info("""
            **To enable project chat:**
            1. Upload documents using the uploader above
            2. Wait for processing to complete
            3. Documents will appear here when ready
            4. Start a new chat to ask questions about your documents
            """)
    
    elif has_project_management and not current_project:
        # No project selected
        st.warning("⚠️ **No project selected**")
        st.info("""
        **Project-specific chat requires an active project.**
        
        Please:
        1. Create or select a project in the sidebar
        2. Upload documents to the project
        3. Return here to chat with project documents
        
        💡 Each project maintains its own isolated chat history and document access.
        """)
    
    else:
        # Traditional mode without project management
        st.info("**Traditional Mode - Limited Chat**")
        st.markdown("""
        **Basic chat functionality available.**
        
        For full project-specific chat features:
        • Enable project management
        • Create dedicated projects
        • Isolated chat histories per project
        
        **Current limitations:** Global document access, no chat isolation
        """)

def requirements_analysis_tab(rag_system, eval_service, target_phase, req_types, export_format, 
                             enable_structured_analysis, enable_cross_phase_analysis, is_enhanced):
    """Combined Requirements & Analysis tab - Professional interface"""
    
    
    # Check project and documents status - Initialize variables first
    current_project = None
    project_documents = []
    has_project_management = hasattr(rag_system, 'get_current_project') and hasattr(rag_system, 'persistence_service')
    
    if has_project_management:
        current_project = rag_system.get_current_project() if hasattr(rag_system, 'get_current_project') else None
        if current_project:
            try:
                project_documents = rag_system.persistence_service.get_project_documents(current_project.id)
                project_documents = [doc for doc in project_documents if doc.processing_status == "completed"]
            except Exception as e:
                st.error(f"❌ Error loading project documents: {str(e)}")
                project_documents = []
        else:
            project_documents = []  # Ensure it's always defined
    else:
        project_documents = []  # Ensure it's always defined
    
    # Document-based input section
    st.markdown("#### 📄 Document Selection & Analysis")
    
    proposal_text = ""
    selected_documents = []
    
    if has_project_management and current_project and project_documents:
        # Project has documents - show document selector
        st.success(f"📋 **Project:** {current_project.name} ({len(project_documents)} documents available)")
        
        # Document selector
        st.markdown("##### Select Documents to Analyze")
        
        # Show documents with checkboxes
        col1, col2 = st.columns([2, 1])
        
        with col1:
            selected_doc_ids = []
            for doc in project_documents:
                # Create checkbox for each document
                is_selected = st.checkbox(
                    f"📄 **{doc.filename}**",
                    value=False,
                    help=f"Size: {doc.file_size / 1024:.1f} KB • Chunks: {doc.chunks_count}",
                    key=f"select_doc_{doc.id}"
                )
                
                if is_selected:
                    selected_doc_ids.append(doc.id)
                    selected_documents.append(doc)
                
                # Show document info
                if is_selected:
                    st.caption(f"✅ Selected - {doc.file_size / 1024:.1f} KB, {doc.chunks_count} text chunks")
        
        with col2:
            # Quick actions
            st.markdown("**Quick Actions**")
            
            if st.button("📤 Select All", use_container_width=True, key="select_all_docs"):
                st.info("💡 Use checkboxes to select documents")
            
            if st.button("📁 Upload More", use_container_width=True, key="upload_more_from_req"):
                st.info("💡 Go to 'Document Management' tab to upload more files")
            
            # Show selection summary
            if selected_documents:
                total_size = sum(doc.file_size for doc in selected_documents) / 1024
                total_chunks = sum(doc.chunks_count for doc in selected_documents)
                st.metric("Selected", f"{len(selected_documents)} docs")
                st.metric("Total Size", f"{total_size:.1f} KB")
                st.metric("Text Chunks", total_chunks)
        
        # Extract text from selected documents
        if selected_documents:
            try:
                # Get document chunks and combine them
                all_chunks = []
                for doc in selected_documents:
                    doc_chunks = rag_system.persistence_service.get_project_chunks(current_project.id)
                    # Filter chunks for this specific document
                    doc_specific_chunks = [chunk for chunk in doc_chunks 
                                         if chunk.get('metadata', {}).get('source_filename') == doc.filename]
                    all_chunks.extend(doc_specific_chunks)
                
                # Combine all chunk content
                proposal_text = "\n\n".join([chunk['content'] for chunk in all_chunks])
                
                if proposal_text:
                    st.success(f"✅ Combined content from {len(selected_documents)} document(s) ({len(proposal_text)} characters)")
                    
                    # Show preview
                    with st.expander(f"📋 Preview Combined Content ({len(all_chunks)} chunks)"):
                        preview_text = proposal_text[:1000] + "..." if len(proposal_text) > 1000 else proposal_text
                        st.text_area("", value=preview_text, disabled=True, height=200, key="combined_preview")
                else:
                    st.warning("⚠️ No content could be extracted from selected documents")
                    
            except Exception as e:
                st.error(f"❌ Error extracting document content: {str(e)}")
    
    elif has_project_management and current_project and not project_documents:
        # Project has no documents - guide user to upload
        st.warning("📭 **No documents found in your project**")
        
        st.markdown("""
        ### 🚀 Get Started with Requirements Analysis
        
        To generate requirements from your project documents, you need to:
        
        1. **📤 Upload documents** - Go to the 'Document Management' tab
        2. **⏳ Wait for processing** - Documents will be analyzed and chunked
        3. **🔄 Return here** - Select processed documents for analysis
        """)
        
        # Quick upload option
        with st.expander("📤 Quick Upload Documents", expanded=True):
            st.markdown("**Upload documents directly here:**")
            
            uploaded_files = st.file_uploader(
                "Select project documents",
                accept_multiple_files=True,
                type=['pdf', 'docx', 'txt', 'md', 'xml', 'json'],
                help="Upload project documents to analyze for requirements generation",
                key="req_analysis_quick_upload"
            )
            
            if uploaded_files:
                st.info(f"📄 {len(uploaded_files)} file(s) ready for upload")
                
                if st.button("🚀 Process Documents", type="primary", key="process_quick_upload"):
                    # Use the same duplicate detection logic as Document Management tab
                    process_documents_with_duplicate_detection(rag_system, current_project, uploaded_files)
                    
                    # Check if content was extracted for requirements generation
                    if st.session_state.get('extracted_document_content'):
                        st.success("🎯 **Content ready for requirements generation!**")
                        st.info("You can now scroll down to generate requirements using the uploaded document content.")
                        
                        # Automatically set the extracted content for requirements generation
                        extracted_content = st.session_state.get('extracted_document_content')
                        extracted_files_info = st.session_state.get('extracted_files_info', [])
                        
                        # Show preview of extracted content
                        st.markdown(f"**📋 Extracted Content Summary ({len(extracted_files_info)} files):**")
                        total_chars = sum(doc['size'] for doc in extracted_files_info)
                        st.write(f"**Total content:** {total_chars:,} characters from {len(extracted_files_info)} files")
                        
                        for doc_info in extracted_files_info:
                            st.write(f"• **{doc_info['filename']}**: {doc_info['size']:,} characters")
                        
                        # Show a brief preview
                        preview_text = extracted_content[:500] + "..." if len(extracted_content) > 500 else extracted_content
                        st.markdown("**Content Preview:**")
                        st.code(preview_text, language="text")
        
        # Alternative input methods for users without documents
        st.markdown("---")
        st.markdown("### 📝 Alternative: Manual Input")
        
        alt_input_method = st.radio(
            "Choose alternative input method",
            ["Paste Text", "Load Example"],
            key="alt_input_method",
            horizontal=True
        )
        
        if alt_input_method == "Paste Text":
            proposal_text = st.text_area(
                "Paste your project proposal text",
                height=200,
                help="Paste project text directly for analysis",
                key="manual_text_input"
            )
        
        else:  # Load Example
            example_choice = st.selectbox(
                "Choose Example Project",
                ["Transportation System", "Cybersecurity Infrastructure", "Industrial Automation"],
                key="manual_example"
            )
            
            if st.button("📋 Load Example", key="load_manual_example"):
                if example_choice == "Transportation System":
                    proposal_text = get_fallback_example("safe")
                elif example_choice == "Cybersecurity Infrastructure":
                    proposal_text = get_fallback_example("cyderco")
                else:
                    proposal_text = get_fallback_example("simple")
                
                if proposal_text:
                    st.success(f"✅ Example loaded: {example_choice}")
                    # Store in session state to persist
                    st.session_state['manual_proposal_text'] = proposal_text
        
        # Retrieve from session state if available
        if 'manual_proposal_text' in st.session_state and not proposal_text:
            proposal_text = st.session_state['manual_proposal_text']
        
        # Check for extracted content from duplicate documents
        if not proposal_text and st.session_state.get('extracted_document_content'):
            proposal_text = st.session_state['extracted_document_content']
            extracted_files_info = st.session_state.get('extracted_files_info', [])
            
            st.success(f"✅ **Using content from uploaded documents:** {', '.join([doc['filename'] for doc in extracted_files_info])}")
            st.info("Content extracted from duplicate files is ready for requirements analysis!")
    
    else:
        # Traditional mode without project management
        st.info("🔧 **Traditional Mode** - Project management not available")
        
        st.markdown("### 📝 Input Methods")
        
        input_method = st.radio(
            "Select Input Method",
            ["Paste Text", "Load Example"],
            key="traditional_input_method",
            horizontal=True
        )
        
        if input_method == "Paste Text":
            proposal_text = st.text_area(
                "Paste your project proposal text",
                height=200,
                help="Paste the text of your project proposal here",
                key="traditional_text_input"
            )
        
        else:  # Load Example
            example_choice = st.selectbox(
                "Choose Example",
                ["Transportation System", "Cybersecurity Infrastructure", "Industrial Automation"],
                key="traditional_example"
            )
            
            if st.button("Load Example", key="traditional_load_example"):
                if example_choice == "Transportation System":
                    proposal_text = get_fallback_example("safe")
                elif example_choice == "Cybersecurity Infrastructure":
                    proposal_text = get_fallback_example("cyderco")
                else:
                    proposal_text = get_fallback_example("simple")
                
                if proposal_text:
                    st.success(f"Example loaded: {example_choice}")
    
    # Configuration section before generation controls
    with st.expander("⚙️ Configuration Options", expanded=False):
        config_col1, config_col2, config_col3 = st.columns(3)
        
        with config_col1:
            st.markdown("#### ARCADIA Settings")
            # Phase selection
            target_phase_local = st.selectbox(
                "Target ARCADIA Phase",
                ["all"] + list(arcadia_config.ARCADIA_PHASES.keys()),
                format_func=lambda x: "All Phases" if x == "all" else arcadia_config.ARCADIA_PHASES[x]["name"],
                index=0 if target_phase == "all" else list(arcadia_config.ARCADIA_PHASES.keys()).index(target_phase) + 1,
                key="req_analysis_phase"
            )
            
            # Requirement types
            req_types_local = st.multiselect(
                "Requirement Types",
                ["functional", "non_functional", "stakeholder"],
                default=req_types,
                key="req_analysis_types"
            )
        
        with config_col2:
            st.markdown("#### Enhanced Analysis")
            if is_enhanced:
                enable_structured_analysis_local = st.checkbox(
                    "Enable Structured ARCADIA Analysis",
                    value=enable_structured_analysis,
                    help="Generate structured analysis with actors, capabilities, and cross-phase traceability",
                    key="req_analysis_structured"
                )
                enable_cross_phase_analysis_local = st.checkbox(
                    "Enable Cross-Phase Analysis", 
                    value=enable_cross_phase_analysis,
                    help="Perform traceability analysis, gap detection, and quality metrics across ARCADIA phases",
                    key="req_analysis_cross_phase"
                )
            else:
                enable_structured_analysis_local = False
                enable_cross_phase_analysis_local = False
                st.info("Enhanced analysis not available in basic mode")
        
        with config_col3:
            st.markdown("#### Generation Settings")
            max_requirements = st.slider("Max Requirements per Type", 5, 50, 20, key="req_analysis_max")
            include_rationale = st.checkbox("Include Rationale", True, key="req_analysis_rationale")
            include_verification = st.checkbox("Include Verification Methods", True, key="req_analysis_verification")
            quality_threshold = st.slider("Quality Threshold", 0.5, 1.0, 0.7, key="req_analysis_quality")
            
            # Export format
            export_formats = ["JSON", "CSV", "Excel", "ARCADIA_JSON", "Structured_Markdown"]
            export_format_local = st.selectbox(
                "Export Format",
                export_formats,
                index=export_formats.index(export_format) if export_format in export_formats else 0,
                key="req_analysis_export"
            )

    # Generation controls
    st.markdown("---")
    col1, col2 = st.columns([3, 1])
    
    with col1:
        if proposal_text:
            char_count = len(proposal_text)
            if selected_documents:
                st.success(f"✅ Ready to analyze {len(selected_documents)} document(s) ({char_count:,} characters)")
            else:
                st.success(f"✅ Input ready ({char_count:,} characters)")
        else:
            if has_project_management and current_project and project_documents:
                st.info("📄 Select documents above to begin analysis")
            else:
                st.info("📝 Please provide input to generate requirements")
    
    with col2:
        generate_btn = st.button(
            "🚀 Generate Requirements & Analysis", 
            type="primary", 
            disabled=not proposal_text,
            key="req_analysis_generate",
            use_container_width=True
        )
        
        if proposal_text and selected_documents:
            st.caption(f"Analyzing {len(selected_documents)} documents")
        elif proposal_text:
            st.caption("Ready for analysis")
    
    # Results display in split view
    if generate_btn and proposal_text:
        start_time = time.time()
        
        if selected_documents:
            logger.info(f"Starting requirements generation from {len(selected_documents)} project documents...")
            logger.info(f"Selected documents: {[doc.filename for doc in selected_documents]}")
        else:
            logger.info("Starting requirements generation from manual input...")
        
        with st.spinner("Generating requirements and performing analysis..."):
            try:
                if is_enhanced and hasattr(rag_system, 'generate_enhanced_requirements_from_proposal'):
                    # Use enhanced generation with structured analysis
                    results = rag_system.generate_enhanced_requirements_from_proposal(
                        proposal_text=proposal_text,
                        target_phase=target_phase_local,
                        requirement_types=req_types_local,
                        enable_structured_analysis=enable_structured_analysis_local,
                        enable_cross_phase_analysis=enable_cross_phase_analysis_local
                    )
                    # Store enhanced results
                    st.session_state['enhanced_results'] = results
                else:
                    # Use traditional generation
                    results = rag_system.generate_requirements_from_proposal(
                        proposal_text, target_phase_local, req_types_local
                    )
                
                end_time = time.time()
                generation_time = end_time - start_time
                
                # Log session activity
                if has_project_management and current_project:
                    try:
                        if selected_documents:
                            activity_desc = f"Generated requirements from {len(selected_documents)} documents: {', '.join([doc.filename for doc in selected_documents])}"
                        else:
                            activity_desc = "Generated requirements from manual input"
                        
                        rag_system.persistence_service.log_project_session(
                            current_project.id,
                            "requirements_generation",
                            activity_desc
                        )
                    except Exception as e:
                        logger.warning(f"Could not log session activity: {str(e)}")
                
                # Display results in full-width structured analysis layout
                st.markdown("#### 🏗️ Structured ARCADIA Analysis")
                
                if is_enhanced and 'enhanced_results' in st.session_state:
                    enhanced_results = st.session_state['enhanced_results']
                    
                    # Analysis overview with improved layout
                    enhancement_summary = enhanced_results.get('enhancement_summary', {})
                    
                    # Main metrics in a clean layout with improved styling
                    st.markdown("""
                    <div class="arcadia-overview">
                        <h3 style="margin-bottom: 1rem; color: #2c5aa0;">📊 Analysis Overview</h3>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    overview_col1, overview_col2, overview_col3, overview_col4 = st.columns(4)
                    with overview_col1:
                        actors_count = enhancement_summary.get('total_actors_identified', 0)
                        st.markdown(f"""
                        <div class="arcadia-metric-card">
                            <div style="font-size: 2rem; color: #28a745; margin-bottom: 0.5rem;">🎭</div>
                            <div style="font-size: 1.8rem; font-weight: 600; color: #2c5aa0;">{actors_count}</div>
                            <div style="color: #6c757d; font-size: 0.9rem;">Actors</div>
                        </div>
                        """, unsafe_allow_html=True)
                        
                    with overview_col2:
                        capabilities_count = enhancement_summary.get('total_capabilities_identified', 0)
                        st.markdown(f"""
                        <div class="arcadia-metric-card">
                            <div style="font-size: 2rem; color: #007bff; margin-bottom: 0.5rem;">🚀</div>
                            <div style="font-size: 1.8rem; font-weight: 600; color: #2c5aa0;">{capabilities_count}</div>
                            <div style="color: #6c757d; font-size: 0.9rem;">Capabilities</div>
                        </div>
                        """, unsafe_allow_html=True)
                        
                    with overview_col3:
                        components_count = enhancement_summary.get('total_components_identified', 0)
                        st.markdown(f"""
                        <div class="arcadia-metric-card">
                            <div style="font-size: 2rem; color: #ffc107; margin-bottom: 0.5rem;">🧩</div>
                            <div style="font-size: 1.8rem; font-weight: 600; color: #2c5aa0;">{components_count}</div>
                            <div style="color: #6c757d; font-size: 0.9rem;">Components</div>
                        </div>
                        """, unsafe_allow_html=True)
                        
                    with overview_col4:
                        # Calculate total requirements from enhanced results
                        total_enhanced_reqs = 0
                        if 'traditional_requirements' in enhanced_results:
                            for phase_reqs in enhanced_results['traditional_requirements'].get('requirements', {}).values():
                                for reqs in phase_reqs.values():
                                    total_enhanced_reqs += len(reqs)
                        st.markdown(f"""
                        <div class="arcadia-metric-card">
                            <div style="font-size: 2rem; color: #6f42c1; margin-bottom: 0.5rem;">📝</div>
                            <div style="font-size: 1.8rem; font-weight: 600; color: #2c5aa0;">{total_enhanced_reqs}</div>
                            <div style="color: #6c757d; font-size: 0.9rem;">Requirements</div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    # Phase tabs for detailed analysis with better spacing
                    st.markdown("---")
                    analysis_tab1, analysis_tab2, analysis_tab3, analysis_tab4 = st.tabs([
                        "🎭 Operational", 
                        "🏗️ System",
                        "🧩 Logical",
                        "🔧 Physical"
                    ])
                    
                    with analysis_tab1:
                        display_operational_analysis(enhanced_results)
                    with analysis_tab2:
                        display_system_analysis(enhanced_results)
                    with analysis_tab3:
                        display_logical_analysis(enhanced_results)
                    with analysis_tab4:
                        display_physical_analysis(enhanced_results)
                    
                    # Show generated requirements at the bottom in a clean format
                    if 'traditional_requirements' in enhanced_results and enhanced_results['traditional_requirements'].get('requirements'):
                        st.markdown("---")
                        st.markdown("#### 📝 Generated Requirements")
                        display_generation_results(enhanced_results['traditional_requirements'], export_format_local, rag_system)
                else:
                    st.info("Structured analysis not available. Enable enhanced analysis options in configuration.")
                    # Fallback to traditional requirements if enhanced is not available
                    st.markdown("---")
                    st.markdown("#### 📝 Generated Requirements")
                    display_generation_results(results, export_format_local, rag_system)
                
                # Quality metrics and evaluation inline
                st.markdown("#### 📊 Quality Metrics & Evaluation")
                
                eval_col1, eval_col2, eval_col3, eval_col4 = st.columns(4)
                
                with eval_col1:
                    st.metric("Generation Time", f"{generation_time:.1f}s")
                with eval_col2:
                    total_reqs = sum(len(reqs) for phase_reqs in results.get('requirements', {}).values() 
                                   for reqs in phase_reqs.values())
                    st.metric("Total Requirements", total_reqs)
                with eval_col3:
                    st.metric("Phases Covered", len(results.get('requirements', {})))
                with eval_col4:
                    if selected_documents:
                        st.metric("Source Documents", len(selected_documents))
                    else:
                        st.metric("Input Method", "Manual")
                
                # Show document source information
                if selected_documents:
                    with st.expander(f"📄 Source Documents ({len(selected_documents)})"):
                        for doc in selected_documents:
                            st.write(f"• **{doc.filename}** ({doc.file_size / 1024:.1f} KB, {doc.chunks_count} chunks)")
                elif proposal_text:
                    with st.expander("📝 Input Source"):
                        st.write(f"Manual input: {len(proposal_text):,} characters")
                
                # Quick evaluation
                if st.button("🔍 Run Quality Evaluation", key="req_analysis_eval"):
                    with st.spinner("Evaluating requirements quality..."):
                        try:
                            eval_results = eval_service.evaluate_requirements(
                                results.get('requirements', {}),
                                proposal_text
                            )
                            
                            # Display evaluation results
                            st.markdown("##### Evaluation Results")
                            
                            metrics_col1, metrics_col2, metrics_col3 = st.columns(3)
                            with metrics_col1:
                                st.metric("Clarity Score", f"{eval_results.get('clarity_score', 0):.2f}")
                            with metrics_col2:
                                st.metric("Completeness", f"{eval_results.get('completeness_score', 0):.2f}")
                            with metrics_col3:
                                st.metric("Consistency", f"{eval_results.get('consistency_score', 0):.2f}")
                            
                            if eval_results.get('recommendations'):
                                st.markdown("##### Recommendations")
                                for rec in eval_results['recommendations']:
                                    st.info(f"💡 {rec}")
                                    
                        except Exception as e:
                            st.error(f"Evaluation error: {str(e)}")
                
                # Export options at bottom
                st.markdown("#### Export Options")
                
                # Get the appropriate results for export
                enhanced_results = st.session_state.get('enhanced_results')
                export_results = enhanced_results.get('traditional_requirements', results) if enhanced_results else results
                
                export_col1, export_col2, export_col3 = st.columns(3)
                
                with export_col1:
                    st.markdown("**Requirements**")
                    try:
                        # Generate JSON export for requirements
                        req_export_content = export_requirements_to_json(export_results)
                        st.download_button(
                            "Download Requirements JSON",
                            req_export_content,
                            "requirements.json",
                            "application/json",
                            key="req_analysis_export_req_json"
                        )
                        
                        # Generate CSV export for requirements
                        req_csv_content = export_requirements_to_csv(export_results)
                        st.download_button(
                            "Download Requirements CSV",
                            req_csv_content,
                            "requirements.csv",
                            "text/csv",
                            key="req_analysis_export_req_csv"
                        )
                    except Exception as e:
                        st.error(f"Export error: {str(e)}")
                
                with export_col2:
                    st.markdown("**ARCADIA Analysis**")
                    if enhanced_results:
                        try:
                            # Export enhanced ARCADIA analysis
                            arcadia_export = {
                                "metadata": {
                                    "export_timestamp": datetime.now().isoformat(),
                                    "export_type": "ARCADIA_Analysis"
                                },
                                "traditional_requirements": enhanced_results.get('traditional_requirements', {}),
                                "structured_analysis": enhanced_results.get('structured_analysis', {}).__dict__ if hasattr(enhanced_results.get('structured_analysis', {}), '__dict__') else {},
                                "enhancement_summary": enhanced_results.get('enhancement_summary', {})
                            }
                            arcadia_content = json.dumps(arcadia_export, indent=2, ensure_ascii=False, default=str)
                            st.download_button(
                                "Download ARCADIA JSON",
                                arcadia_content,
                                "arcadia_analysis.json",
                                "application/json",
                                key="req_analysis_export_arcadia_json"
                            )
                        except Exception as e:
                            st.error(f"ARCADIA export error: {str(e)}")
                    else:
                        st.info("Enhanced analysis not available")
                
                with export_col3:
                    st.markdown("**Excel Format**")
                    try:
                        # Generate Excel-compatible CSV
                        excel_content = export_requirements_to_excel_csv(export_results)
                        st.download_button(
                            "Download Excel CSV",
                            excel_content,
                            "requirements_excel.csv",
                            "text/csv",
                            key="req_analysis_export_excel"
                        )
                        
                        # Show export summary
                        total_reqs = sum(len(reqs) for phase_reqs in export_results.get('requirements', {}).values() 
                                       for reqs in phase_reqs.values() if isinstance(reqs, list))
                        st.caption(f"{total_reqs} requirements ready")
                    except Exception as e:
                        st.error(f"Excel export error: {str(e)}")
                
            except Exception as e:
                st.error(f"❌ Generation error: {str(e)}")
                logger.error(f"Requirements generation error: {str(e)}")
    
    # Show any existing results if available
    elif 'enhanced_results' in st.session_state or 'requirements_results' in st.session_state:
        st.markdown("#### 📋 Previous Results")
        st.info("Previous generation results are available. Generate new requirements to update.")
        
        if st.button("🔄 Clear Previous Results", key="req_analysis_clear"):
            if 'enhanced_results' in st.session_state:
                del st.session_state['enhanced_results']
            if 'requirements_results' in st.session_state:
                del st.session_state['requirements_results']
            st.success("Previous results cleared!")
            st.rerun()

def project_insights_tab(rag_system, current_project, has_project_management):
    """Combined Project Insights tab - Professional interface"""
    st.markdown("### Project Insights")
    
    if not has_project_management:
        st.info("**Traditional Mode** - Limited insights available without project management")
        
        # Show basic system information
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("System Mode", "Traditional")
        with col2:
            st.metric("RAG System", "Active")
        with col3:
            st.metric("Persistence", "Disabled")
        
        st.markdown("#### Available Tools")
        st.markdown("• Basic requirements generation")
        st.markdown("• Document chat functionality") 
        st.markdown("• Simple evaluation metrics")
        st.markdown("\n**Enable project management for full insights dashboard**")
        return
    
    if not current_project:
        st.warning("No project selected. Please select or create a project to view insights.")
        return
    
    # Project dashboard with key metrics
    st.markdown(f"## 📋 {current_project.name} Dashboard")
    if current_project.description:
        st.info(current_project.description)
    
    # Key metrics header - Most important statistics only
    metric_col1, metric_col2, metric_col3 = st.columns(3)
    
    with metric_col1:
        st.metric("📄 Documents", current_project.documents_count)
    with metric_col2:
        st.metric("📝 Requirements", current_project.requirements_count)
    with metric_col3:
        days_active = (datetime.now() - current_project.created_at).days
        st.metric("📅 Days Active", days_active)
    
    # Load project data safely once for all sections
    project_data = _load_project_data_safely(rag_system, current_project.id)
    
    # Show any errors that occurred during data loading
    if project_data["errors"]:
        with st.expander("⚠️ Data Loading Issues", expanded=False):
            for error in project_data["errors"]:
                st.warning(f"• {error}")
            st.info("The application will continue with available data.")
    
    # Main insights sections
    insights_section1, insights_section2 = st.columns([2, 1])
    
    with insights_section1:
        # Requirements overview - Compact with safe loading
        st.markdown("#### 📊 Requirements Overview")
        
        # Use safely loaded project data
        requirements_data = project_data["requirements"]
        requirements = requirements_data.get("requirements", {})
        
        if requirements:
            # Calculate phase coverage
            phase_coverage = {}
            total_requirements = 0
            
            for phase, phase_reqs in requirements.items():
                phase_total = sum(len(reqs) for reqs in phase_reqs.values())
                phase_coverage[phase] = phase_total
                total_requirements += phase_total
            
            if phase_coverage:
                # Compact phase summary
                req_col1, req_col2 = st.columns(2)
                
                with req_col1:
                    st.markdown("**Phase Distribution:**")
                    for phase, count in phase_coverage.items():
                        percentage = (count / total_requirements) * 100 if total_requirements > 0 else 0
                        st.write(f"• **{phase.title()}**: {count} ({percentage:.0f}%)")
                
                with req_col2:
                    st.markdown("**Quality Status:**")
                    for phase, count in phase_coverage.items():
                        quality = "🟢 Good" if count >= 10 else "🟡 Moderate" if count >= 5 else "🔴 Low"
                        st.write(f"• **{phase.title()}**: {quality}")
        else:
            _handle_missing_data_gracefully("requirements", current_project.name)
    
    with insights_section2:
        # Project settings and management
        st.markdown("#### ⚙️ Project Settings")
        
        # Quick project info
        with st.expander("📋 Project Information", expanded=True):
            st.write(f"**Name:** {current_project.name}")
            st.write(f"**Created:** {current_project.created_at.strftime('%d/%m/%Y %H:%M')}")
            st.write(f"**Updated:** {current_project.updated_at.strftime('%d/%m/%Y %H:%M')}")
            
            if current_project.proposal_text:
                st.markdown("**📄 Proposal Text:**")
                proposal_preview = current_project.proposal_text[:500] + "..." if len(current_project.proposal_text) > 500 else current_project.proposal_text
                st.text_area("", value=proposal_preview, disabled=True, height=100, key="project_proposal_preview")
        
        # Compact project statistics with safe loading
        st.markdown("#### 📊 Key Stats")
        
        # Use safely loaded data (already loaded at the beginning)
        documents = project_data["documents"]
        stakeholders = project_data["stakeholders"]
        
        if documents:
            total_size = sum(doc.file_size for doc in documents) / (1024 * 1024)  # MB
            total_chunks = sum(doc.chunks_count for doc in documents)
            
            # Compact metrics in single line format
            st.write(f"📁 **Size**: {total_size:.1f} MB | 🧩 **Chunks**: {total_chunks} | 👥 **Stakeholders**: {len(stakeholders)}")
            
            # Document types in compact format
            doc_types = {}
            for doc in documents:
                ext = doc.filename.split('.')[-1].upper()
                doc_types[ext] = doc_types.get(ext, 0) + 1
            
            types_text = " | ".join([f"{dtype}: {count}" for dtype, count in doc_types.items()])
            st.caption(f"**Types**: {types_text}")
        else:
            _handle_missing_data_gracefully("documents", current_project.name)
        
        # Enhanced Project Settings with CRUD Operations
        st.markdown("#### Project Settings")
        
        # Get project manager for enhanced functionality
        project_manager = init_project_manager(rag_system)
        
        # Project settings sections
        settings_col1, settings_col2 = st.columns(2)
        
        with settings_col1:
            st.markdown("**Project Operations**")
            
            # Edit Project
            if st.button("Edit Project", use_container_width=True, key="edit_project_insights"):
                st.session_state.show_edit_project_insights_modal = True
            
            # Project Health Check
            if st.button("Health Check", use_container_width=True, key="health_check_insights"):
                if project_manager:
                    try:
                        health_check = project_manager.get_project_health_check(current_project)
                        st.session_state.show_health_check_modal = True
                        st.session_state.health_check_data = health_check
                    except Exception as e:
                        st.error(f"Health check error: {str(e)}")
                else:
                    st.error("Project manager not available")
        
        with settings_col2:
            st.markdown("**Data Operations**")
            
            # Enhanced Export with Project Manager
            if st.button("Export Project", use_container_width=True, key="export_project_insights"):
                if project_manager:
                    try:
                        export_data = project_manager.export_project_data(current_project)
                        st.session_state.show_export_modal = True
                        st.session_state.export_data = export_data
                    except Exception as e:
                        st.error(f"Export error: {str(e)}")
                else:
                    st.error("Project manager not available")
            
            # Refresh Data
            if st.button("Refresh Data", use_container_width=True, key="refresh_project_insights"):
                _handle_refresh_project_data(rag_system, current_project)
            
        # Danger Zone
        with st.expander("Danger Zone", expanded=False):
            st.warning("**Destructive Operations**")
            st.markdown("These operations cannot be undone. Use with caution.")
            
            if st.button("Delete Project", use_container_width=True, key="delete_project_insights", type="secondary"):
                st.session_state.show_delete_project_insights_modal = True
        
        # Handle all modals
        _handle_project_settings_modals(rag_system, current_project, project_manager)


def _handle_project_settings_modals(rag_system, current_project, project_manager):
    """Handle all project settings modals in the insights tab"""
    
    # Edit Project Modal
    if st.session_state.get('show_edit_project_insights_modal', False):
        with st.container():
            st.markdown("### Edit Project")
            
            with st.form("edit_project_insights_form"):
                # Pre-fill with current values
                new_name = st.text_input("Project Name", value=current_project.name, max_chars=100)
                new_description = st.text_area("Description", value=current_project.description, height=100, max_chars=1000)
                new_proposal = st.text_area("Proposal Text", value=current_project.proposal_text, height=150, max_chars=5000)
                
                col1, col2 = st.columns(2)
                
                with col1:
                    if st.form_submit_button("Save Changes", type="primary"):
                        try:
                            # Validate inputs
                            if project_manager:
                                is_valid, error_message = project_manager.validate_project_data(new_name, new_description, new_proposal)
                                if not is_valid:
                                    st.error(f"{error_message}")
                                    return
                            
                            # Update project
                            success = rag_system.persistence_service.update_project(
                                current_project.id,
                                name=new_name.strip() if new_name.strip() != current_project.name else None,
                                description=new_description.strip() if new_description.strip() != current_project.description else None,
                                proposal_text=new_proposal.strip() if new_proposal.strip() != current_project.proposal_text else None
                            )
                            
                            if success:
                                st.success("Project updated successfully!")
                                st.session_state.show_edit_project_insights_modal = False
                                st.rerun()
                            else:
                                st.error("Failed to update project")
                                
                        except Exception as e:
                            st.error(f"Update error: {str(e)}")
                
                with col2:
                    if st.form_submit_button("Cancel"):
                        st.session_state.show_edit_project_insights_modal = False
                        st.rerun()
    
    # Health Check Modal
    if st.session_state.get('show_health_check_modal', False):
        with st.container():
            st.markdown("### Project Health Check")
            
            health_data = st.session_state.get('health_check_data', {})
            
            if health_data:
                # Health status with appropriate color
                status = health_data.get('status', 'unknown')
                score = health_data.get('score', 0)
                
                if status == 'healthy':
                    st.success(f"**Status**: {status.title()} (Score: {score}/100)")
                elif status == 'warning':
                    st.warning(f"**Status**: {status.title()} (Score: {score}/100)")
                else:
                    st.error(f"**Status**: {status.title()} (Score: {score}/100)")
                
                # Issues
                issues = health_data.get('issues', [])
                if issues:
                    st.markdown("**Issues Found:**")
                    for issue in issues:
                        st.markdown(f"• {issue}")
                
                # Recommendations
                recommendations = health_data.get('recommendations', [])
                if recommendations:
                    st.markdown("**Recommendations:**")
                    for rec in recommendations:
                        st.markdown(f"• {rec}")
                
                if not issues and not recommendations:
                    st.success("No issues found! Your project is in excellent health.")
            
            if st.button("Close Health Check", use_container_width=True):
                st.session_state.show_health_check_modal = False
                st.rerun()
    
    # Export Modal
    if st.session_state.get('show_export_modal', False):
        with st.container():
            st.markdown("### Export Project Data")
            
            export_data = st.session_state.get('export_data', {})
            
            if export_data and not export_data.get('error'):
                # Create downloadable JSON
                export_json = json.dumps(export_data, indent=2, ensure_ascii=False)
                
                st.download_button(
                    "Download Project Export",
                    export_json,
                    f"{current_project.name}_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    "application/json",
                    use_container_width=True,
                    key="download_project_export"
                )
                
                # Show export summary
                st.success("Export prepared successfully!")
                
                # Display export contents
                project_info = export_data.get('project_info', {})
                documents = export_data.get('documents', [])
                requirements = export_data.get('requirements', {})
                stakeholders = export_data.get('stakeholders', [])
                
                st.markdown("**Export Contents:**")
                st.markdown(f"• **Project**: {project_info.get('name', 'N/A')}")
                st.markdown(f"• **Documents**: {len(documents)} files")
                st.markdown(f"• **Requirements**: {len(requirements.get('requirements', {}))} phases")
                st.markdown(f"• **Stakeholders**: {len(stakeholders)} records")
                
            elif export_data.get('error'):
                st.error(f"Export failed: {export_data['error']}")
            
            if st.button("Close Export", use_container_width=True):
                st.session_state.show_export_modal = False
                st.rerun()
    
    # Delete Project Modal
    if st.session_state.get('show_delete_project_insights_modal', False):
        with st.container():
            st.markdown("### Delete Project")
            st.error("**DANGER: This action cannot be undone!**")
            
            st.markdown(f"""
            **Project to delete:** {current_project.name}
            
            This will permanently delete:
            • All project documents and chunks
            • All requirements and analyses
            • All stakeholder information
            • All project history
            """)
            
            with st.form("delete_project_insights_form"):
                # Confirmation inputs
                confirm_delete = st.checkbox("I understand this action is irreversible")
                confirmation_text = st.text_input(
                    "Type the project name to confirm deletion",
                    placeholder=current_project.name
                )
                
                col1, col2 = st.columns(2)
                
                with col1:
                    if st.form_submit_button("DELETE PROJECT", type="secondary"):
                        if not confirm_delete:
                            st.error("Please confirm you understand this action is irreversible")
                            return
                        
                        if confirmation_text != current_project.name:
                            st.error("Project name doesn't match. Please type the exact project name.")
                            return
                        
                        try:
                            # Delete project
                            success = rag_system.persistence_service.delete_project(current_project.id)
                            
                            if success:
                                st.success("Project deleted successfully!")
                                st.session_state.show_delete_project_insights_modal = False
                                st.info("Redirecting to project selection...")
                                st.rerun()
                            else:
                                st.error("Failed to delete project")
                                
                        except Exception as e:
                            st.error(f"Delete error: {str(e)}")
                
                with col2:
                    if st.form_submit_button("Cancel"):
                        st.session_state.show_delete_project_insights_modal = False
                        st.rerun()


if __name__ == "__main__":
    main()